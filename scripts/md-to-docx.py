#!/usr/bin/env python3
"""
Converts the project's Markdown documents into polished .docx files for
client submission, and renders the API list as a Word table.

Usage:
    python scripts/md-to-docx.py

Outputs (docs/downloads/):
    Pixous_HR_Requirements_v1.0.docx
    Pixous_HR_Unit_Testing_v1.0.docx
    Pixous_HR_API_List_v1.0.docx
    Pixous_HR_API_List.json
"""
import os
import re
import json
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), "..")
DOCS = os.path.join(ROOT, "docs")
OUT = os.path.join(DOCS, "downloads")

BRAND = RGBColor(0x4F, 0x46, 0xE5)      # product indigo
DARK = RGBColor(0x1E, 0x29, 0x3B)       # slate-900
GRAY = RGBColor(0x64, 0x74, 0x8B)       # slate-500
CODE_FONT = "Consolas"


def style_document(doc):
    """Set a consistent professional look: brand headings, tight tables."""
    for section in doc.sections:
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 20, BRAND, 18, 8),
        ("Heading 2", 15, DARK, 14, 6),
        ("Heading 3", 12, DARK, 10, 4),
        ("Heading 4", 11, GRAY, 8, 3),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_runs(paragraph, text):
    """Render **bold** and `code` inline markup into styled runs."""
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(9.5)
            r.font.color.rgb = BRAND
        else:
            paragraph.add_run(tok)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_markdown_table(doc, rows):
    ncols = len(rows[0])
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, cell_text in enumerate(rows[0]):
        p = hdr[i].paragraphs[0]
        add_runs(p, cell_text)
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "4F46E5")
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(ncols):
            txt = row[i] if i < len(row) else ""
            p = cells[i].paragraphs[0]
            add_runs(p, txt)
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(9)
    return table


def md_to_docx(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style_document(doc)

    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()

        # fenced code blocks
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = CODE_FONT
            r.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line.strip()):
            i += 1
            continue

        # table block
        if line.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.fullmatch(r":?-{2,}:?", cells[0]) or len(cells) == 1:
                    block.append(cells)
                i += 1
            if len(block) > 1:
                add_markdown_table(doc, block)
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            p = doc.add_heading(level=level)
            add_runs(p, m.group(2).strip())
            i += 1
            continue

        # list items
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # blockquote
        m = re.match(r"^>\s?(.*)", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(m.group(1))
            r.italic = True
            r.font.color.rgb = GRAY
            i += 1
            continue

        # blank line
        if not line.strip():
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


def api_list_to_docx():
    with open(os.path.join(DOCS, "api-list.json"), encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    style_document(doc)
    h = doc.add_heading(level=1)
    add_runs(h, "Pixous HR Portal — API List")
    p = doc.add_paragraph()
    add_runs(p, f"Complete REST API inventory · {data['count']} endpoints · generated {data['generated']}")
    p2 = doc.add_paragraph()
    add_runs(p2, "Base URL: " + data["baseUrl"])

    rows = [["#", "Method", "Path", "Controller", "Authorization"]]
    for idx, e in enumerate(data["endpoints"], 1):
        rows.append([
            str(idx),
            e["method"],
            e["path"],
            e["controller"].replace(".java", ""),
            e["authorization"] or "— (auth or public per controller)",
        ])
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(rows[0]):
        r = hdr[i].paragraphs[0].add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "4F46E5")
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(5):
            r = cells[i].paragraphs[0].add_run(row[i])
            r.font.size = Pt(8)

    out = os.path.join(OUT, "Pixous_HR_API_List_v1.0.docx")
    os.makedirs(OUT, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")

    # Convenience copy of the raw JSON alongside the Word documents.
    dst = os.path.join(OUT, "Pixous_HR_API_List.json")
    shutil.copyfile(os.path.join(DOCS, "api-list.json"), dst)
    print(f"Wrote {dst}")


def main():
    os.makedirs(OUT, exist_ok=True)
    md_to_docx(os.path.join(DOCS, "REQUIREMENTS.md"),
               os.path.join(OUT, "Pixous_HR_Requirements_v1.0.docx"))
    md_to_docx(os.path.join(DOCS, "UNIT-TESTING.md"),
               os.path.join(OUT, "Pixous_HR_Unit_Testing_v1.0.docx"))
    api_list_to_docx()


if __name__ == "__main__":
    main()
