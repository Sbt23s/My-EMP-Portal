"""
Shared house style for the PIXOUS HR Portal testing documents.

There were four documents to produce and one look to keep, so the cover, the
headings and the tables live here rather than being pasted into each generator
and drifting apart. The style follows the Performance Testing document that was
already in this repository: Calibri, navy headings, the logo centred on the
cover, and a metadata block underneath it.

Every figure these documents quote is measured, not estimated. The measurement
commands are recorded next to the numbers so a reader can re-run them.
"""
from __future__ import annotations

import datetime
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = RGBColor(0, 51, 102)
TEAL = RGBColor(0, 102, 153)
HEADER_FILL = '003366'

# The generators run from work/docgen; the logo sits one level up beside the
# older generators that already used it.
LOGO = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                    'pixous_logo.png')


def new_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    normal = doc.styles['Normal'].font
    normal.name = 'Calibri'
    normal.size = Pt(10)
    return doc


def add_logo(doc: Document, width_inches: float = 1.9) -> None:
    """The logo, or nothing at all if the file has moved.

    A missing image must not stop a document being produced: the content is
    the deliverable and the mark on the cover is not worth failing over.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO):
        p.add_run().add_picture(LOGO, width=Inches(width_inches))
    else:
        run = p.add_run('PIXOUS TECHNOLOGIES')
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY


def cover(doc: Document, doc_title: str, subtitle: str, version: str = '1.0',
          extra_rows: list[tuple[str, str]] | None = None) -> None:
    add_logo(doc)
    doc.add_paragraph()

    title = doc.add_heading('PIXOUS HR Portal', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_heading(doc_title, level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = TEAL

    line = doc.add_paragraph(subtitle)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in line.runs:
        run.font.size = Pt(10.5)
        run.font.italic = True

    doc.add_paragraph()
    rows = [
        ('Document Type', doc_title),
        ('Version', version),
        ('Date', datetime.date.today().strftime('%d %B %Y')),
        ('Application', 'PIXOUS HR Portal — web, mobile and backend'),
        ('Environment', 'Production — https://pixoushrportal.pixous.info'),
        ('Prepared By', 'PIXOUS Technologies — Engineering'),
        ('Classification', 'Confidential — Internal / Client Deliverable'),
    ] + (extra_rows or [])

    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for run in t.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
        for run in t.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(9.5)
    doc.add_page_break()


def h1(doc: Document, text: str):
    head = doc.add_heading(text, level=1)
    for run in head.runs:
        run.font.color.rgb = NAVY
    return head


def h2(doc: Document, text: str):
    head = doc.add_heading(text, level=2)
    for run in head.runs:
        run.font.color.rgb = TEAL
    return head


def para(doc: Document, text: str, *, italic: bool = False, size: float = 10):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.italic = italic
    return p


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(9.5)


def table(doc: Document, headers: list[str], rows: list[list],
          widths: list[float] | None = None):
    """A header-shaded table. Column widths in inches, when given.

    Widths are applied per cell rather than per column: Word ignores column
    widths on a table whose autofit it has already decided, and applying them
    to every cell is the form it honours.
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = str(header)
        shading = cell._element.get_or_add_tcPr()
        shading.append(shading.makeelement(
            qn('w:shd'), {qn('w:fill'): HEADER_FILL, qn('w:val'): 'clear'}))
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(255, 255, 255)

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = '' if value is None else str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)

    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                if j < len(row.cells):
                    row.cells[j].width = Inches(w)
    return t


def note(doc: Document, text: str) -> None:
    """A short indented remark, for the reasoning behind a result."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(70, 70, 70)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def save(doc: Document, path: str) -> str:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path
