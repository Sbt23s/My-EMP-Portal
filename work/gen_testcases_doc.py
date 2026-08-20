"""Generate PIXOUS HR Portal - Test Cases Document"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

def add_logo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('pixous_logo.png', width=Inches(2.0))

def add_styled_table(doc, headers, rows_data):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): '003366', qn('w:val'): 'clear'})
        shading.append(shd)
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    return table

# ═══ COVER PAGE ═══
add_logo(doc)
doc.add_paragraph()
title = doc.add_heading('PIXOUS HR Portal', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_heading('Comprehensive Test Cases Document', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
t = doc.add_table(rows=5, cols=2)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ('Document Type', 'Test Cases Document (QA / UAT)'),
    ('Version', '1.0'),
    ('Date', datetime.date.today().strftime('%B %d, %Y')),
    ('Prepared By', 'PIXOUS Technologies QA Team'),
    ('Classification', 'Confidential - Client Deliverable'),
]
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
doc.add_page_break()

# ═══ TOC ═══
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Test Strategy Overview',
    '2. Test Environment',
    '3. Authentication & Login Test Cases',
    '4. Employee Role - Test Cases',
    '5. HR Manager Role - Test Cases',
    '6. Team Lead / Manager Role - Test Cases',
    '7. System Admin Role - Test Cases',
    '8. CTO / Executive Role - Test Cases',
    '9. Attendance Module Test Cases',
    '10. Leave Management Test Cases',
    '11. Payroll Module Test Cases',
    '12. Asset Management Test Cases',
    '13. Helpdesk / Ticketing Test Cases',
    '14. Community & Chat Test Cases',
    '15. Task Management Test Cases',
    '16. Dashboard & Reports Test Cases',
    '17. Onboarding Test Cases',
    '18. Cross-Browser & Responsive Test Cases',
    '19. Regression Test Summary',
    '20. Defect Tracking Template',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)
doc.add_page_break()

# ═══ 1. TEST STRATEGY ═══
doc.add_heading('1. Test Strategy Overview', level=1)
doc.add_paragraph(
    'This document defines comprehensive manual and automated test cases for the PIXOUS HR Portal. '
    'Test cases cover all user roles (Employee, HR Manager, Team Lead, System Admin, CTO/Executive) '
    'across all functional modules. Each test case includes step-by-step instructions, expected results, '
    'pre-conditions, and pass/fail criteria.'
)

doc.add_paragraph(
    'Each test case includes a unique identifier, description, priority level, pre-conditions, '
    'detailed steps, and expected results. The document is organized by role and module for easy '
    'navigation and execution tracking.'
)

doc.add_paragraph(
    'This document is intended for QA engineers, test leads, and project managers involved in '
    'the quality assurance process for the PIXOUS HR Portal. It should be reviewed and updated '
    'as new features are added or existing features are modified.'
)

doc.add_heading('Test Types Covered', level=2)
types = [
    ['Functional Testing', 'Verify each feature works as specified per business requirements'],
    ['Role-Based Access Testing', 'Verify each role can access only permitted features'],
    ['Integration Testing', 'Verify cross-module workflows (e.g., leave -> payroll deduction)'],
    ['UI/UX Testing', 'Verify interface elements, responsiveness, and user experience'],
    ['API Testing', 'Verify REST API endpoints return correct status codes and data'],
    ['Negative Testing', 'Verify system handles invalid inputs and unauthorized access gracefully'],
    ['Regression Testing', 'Verify existing features remain intact after changes'],
]
add_styled_table(doc, ['Test Type', 'Description'], types)

doc.add_heading('Test Case ID Format', level=2)
doc.add_paragraph('TC-[MODULE]-[ROLE]-[SEQ] | Example: TC-AUTH-EMP-001')
doc.add_paragraph('Modules: AUTH, ATT, LEA, PAY, AST, HD, CHAT, TASK, DASH, ONB, PERF, CAL, EMP')
doc.add_paragraph('Roles: EMP (Employee), HR (HR Manager), TL (Team Lead), ADM (Admin), CTO (Executive)')
doc.add_paragraph('Priority: High (must pass for release), Medium (should pass), Low (nice to have)')

doc.add_page_break()

# ═══ 2. TEST ENVIRONMENT ═══
doc.add_heading('2. Test Environment', level=1)
env = [
    ['Application URL', 'https://hr.pixoustech.com (Production) / http://localhost:5173 (Dev)'],
    ['API Base URL', 'https://api.pixoustech.com/api/ / http://localhost:8080/api/'],
    ['Web Browser', 'Chrome 120+, Firefox 120+, Safari 17+, Edge 120+'],
    ['Mobile', 'Android 12+ (React Native), iOS 16+ (Flutter)'],
    ['OS', 'Windows 11, macOS 14, Ubuntu 22.04'],
    ['Database', 'MySQL 8.4 (port 3307 dev, 3306 prod)'],
    ['Test Data', 'Pre-seeded via Flyway migrations V1-V101'],
    ['Test Accounts', 'Employee: testemp / HR: testhr / TL: testtl / Admin: testadmin'],
]
add_styled_table(doc, ['Component', 'Details'], env)

doc.add_page_break()

# ═══ 3. AUTHENTICATION TEST CASES ═══
doc.add_heading('3. Authentication & Login Test Cases', level=1)

auth_tests = [
    ['TC-AUTH-EMP-001', 'Valid Login', 'All Roles', 'High',
     '1. Navigate to login page\n2. Enter valid username\n3. Enter valid password\n4. Click "Login" button',
     'User is authenticated and redirected to dashboard. JWT tokens stored.'],
    ['TC-AUTH-EMP-002', 'Invalid Password', 'All Roles', 'High',
     '1. Navigate to login page\n2. Enter valid username\n3. Enter INVALID password\n4. Click "Login" button',
     'Error message displayed: "Invalid username or password". User remains on login page.'],
    ['TC-AUTH-EMP-003', 'Empty Fields Validation', 'All Roles', 'Medium',
     '1. Navigate to login page\n2. Leave username and password empty\n3. Click "Login" button',
     'Form validation errors shown for both fields. No API call made.'],
    ['TC-AUTH-EMP-004', 'Account Lockout (5 failures)', 'All Roles', 'High',
     '1. Navigate to login page\n2. Enter valid username\n3. Enter wrong password 5 times consecutively',
     'Account locked after 5th failure. Message: "Account locked. Contact administrator."'],
    ['TC-AUTH-EMP-005', 'Token Refresh on 401', 'All Roles', 'High',
     '1. Login successfully\n2. Wait for access token to expire (4hr)\n3. Perform any action',
     'Axios interceptor automatically refreshes token. User session continues seamlessly.'],
    ['TC-AUTH-EMP-006', 'Logout', 'All Roles', 'High',
     '1. Login successfully\n2. Click user profile menu\n3. Click "Logout"',
     'All tokens revoked. User redirected to login page. Subsequent API calls return 401.'],
    ['TC-AUTH-EMP-007', 'Password Change', 'All Roles', 'Medium',
     '1. Login successfully\n2. Navigate to change password\n3. Enter current and new password\n4. Submit',
     'Password updated. Old password no longer works. New password works on next login.'],
    ['TC-AUTH-EMP-008', 'Username Availability Check', 'Registration', 'Low',
     '1. Navigate to signup page\n2. Enter existing username\n3. Check availability indicator',
     'System shows "Username not available" in real-time.'],
    ['TC-AUTH-EMP-009', 'Phone Number Validation', 'Registration', 'Low',
     '1. Navigate to signup page\n2. Enter already registered phone number\n3. Submit form',
     'Error: "Phone number already registered". Form not submitted.'],
    ['TC-AUTH-EMP-010', 'Session Persistence', 'All Roles', 'Medium',
     '1. Login successfully\n2. Close browser tab\n3. Open new tab and navigate to app URL',
     'User remains logged in (refresh token used). Dashboard loads automatically.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Role', 'Priority', 'Steps', 'Expected Result'], auth_tests)

doc.add_page_break()

# ═══ 4. EMPLOYEE ROLE TEST CASES ═══
doc.add_heading('4. Employee Role - Test Cases', level=1)

doc.add_heading('4.1 Dashboard & Navigation', level=2)
emp_dash = [
    ['TC-EMP-DASH-001', 'View Personal Dashboard', 'High',
     '1. Login as Employee\n2. Observe dashboard widgets',
     'Dashboard shows attendance summary, pending tasks, notifications, leave balance, celebrations.'],
    ['TC-EMP-DASH-002', 'Navigation Menu Access', 'High',
     '1. Login as Employee\n2. Check sidebar navigation items',
     'Only permitted menu items visible: Dashboard, Attendance, Leave, Payslip, Helpdesk, Chat, Tasks.'],
    ['TC-EMP-DASH-003', 'Notification Bell', 'Medium',
     '1. Login as Employee\n2. Click notification bell icon\n3. View notification list',
     'Notification dropdown shows recent notifications with timestamps. Mark as read works.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], emp_dash)

doc.add_heading('4.2 Attendance Self-Service', level=2)
emp_att = [
    ['TC-EMP-ATT-001', 'GPS Punch In', 'High',
     '1. Login as Employee\n2. Navigate to Attendance\n3. Click "Punch In" button\n4. Allow GPS access',
     'Attendance recorded with GPS coordinates. Time stamp shown. "Punch In" changes to "Punch Out" button.'],
    ['TC-EMP-ATT-002', 'GPS Punch Out', 'High',
     '1. Login as Employee (already punched in)\n2. Navigate to Attendance\n3. Click "Punch Out" button',
     'Attendance out recorded. Duration calculated and displayed. Work hours shown.'],
    ['TC-EMP-ATT-003', 'Face Recognition Punch', 'High',
     '1. Login as Employee\n2. Navigate to Attendance\n3. Click "Face Punch" button\n4. Position face in camera frame',
     'Face verified against stored embedding. Attendance recorded with face verification flag.'],
    ['TC-EMP-ATT-004', 'Today Attendance Status', 'Medium',
     '1. Login as Employee\n2. Navigate to Attendance\n3. View today\'s status',
     'Shows punch-in time, punch-out time, total hours, and attendance status (Present/Absent/Late).'],
    ['TC-EMP-ATT-005', 'Monthly Attendance Calendar', 'Medium',
     '1. Login as Employee\n2. Navigate to Attendance Calendar\n3. Select month',
     'Calendar shows daily attendance status with color coding (green=present, red=absent, yellow=leave).'],
    ['TC-EMP-ATT-006', 'Monthly Summary', 'Medium',
     '1. Login as Employee\n2. Navigate to Attendance Summary\n3. View monthly stats',
     'Shows total present days, absent days, leave days, overtime hours, and LOP days.'],
    ['TC-EMP-ATT-007', 'Attendance Insights', 'Low',
     '1. Login as Employee\n2. Navigate to Attendance Insights',
     'Shows attendance patterns, anomalies (late arrivals, early departures), and suggestions.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], emp_att)

doc.add_heading('4.3 Leave Self-Service', level=2)
emp_leave = [
    ['TC-EMP-LEA-001', 'View Leave Balances', 'High',
     '1. Login as Employee\n2. Navigate to Leave > Balances',
     'Leave balances shown for each type (Casual, Sick, Earned) with remaining and total counts.'],
    ['TC-EMP-LEA-002', 'Apply for Leave', 'High',
     '1. Login as Employee\n2. Navigate to Leave > Apply\n3. Select leave type\n4. Select start and end dates\n5. Enter reason\n6. Click "Submit"',
     'Leave request created with PENDING status. Approver notified via WebSocket.'],
    ['TC-EMP-LEA-003', 'View Leave History', 'Medium',
     '1. Login as Employee\n2. Navigate to Leave > My Requests',
     'Shows all leave requests with status (Pending/Approved/Rejected/Cancelled) and dates.'],
    ['TC-EMP-LEA-004', 'Cancel Pending Leave', 'Medium',
     '1. Login as Employee\n2. Navigate to Leave > My Requests\n3. Find pending request\n4. Click "Cancel"',
     'Leave request cancelled. Status changed to CANCELLED. Leave balance restored.'],
    ['TC-EMP-LEA-005', 'Cancel Approved Leave', 'Medium',
     '1. Login as Employee\n2. Navigate to Leave > My Requests\n3. Find approved request\n4. Click "Cancel"',
     'Leave request cancelled. Leave balance restored. Approver notified.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], emp_leave)

doc.add_heading('4.4 Payslip Access', level=2)
emp_pay = [
    ['TC-EMP-PAY-001', 'View My Payslips', 'High',
     '1. Login as Employee\n2. Navigate to Payslips\n3. View list of payslips',
     'List shows all generated payslips with month, year, and net pay amount.'],
    ['TC-EMP-PAY-002', 'Download Payslip PDF', 'High',
     '1. Login as Employee\n2. Navigate to Payslips\n3. Click "Download" on a payslip',
     'PDF file downloaded with correct payslip details, company name, and PIXOUS logo.'],
    ['TC-EMP-PAY-003', 'View Salary Structure', 'Medium',
     '1. Login as Employee\n2. Navigate to Salary section',
     'Shows own salary breakdown (Basic, HRA, Conveyance, Deductions, Net Pay).'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], emp_pay)

doc.add_heading('4.5 Helpdesk Self-Service', level=2)
emp_help = [
    ['TC-EMP-HD-001', 'Raise Ticket', 'High',
     '1. Login as Employee\n2. Navigate to Helpdesk\n3. Click "New Ticket"\n4. Enter subject and description\n5. Select category\n6. Click "Submit"',
     'Ticket created with OPEN status. Ticket ID assigned. Agent notified.'],
    ['TC-EMP-HD-002', 'Upload Attachment', 'Medium',
     '1. Login as Employee\n2. Create or open a ticket\n3. Click "Attach File"\n4. Select file\n5. Upload',
     'File uploaded and attached to ticket. File size and name shown.'],
    ['TC-EMP-HD-003', 'View My Tickets', 'Medium',
     '1. Login as Employee\n2. Navigate to Helpdesk\n3. View ticket list',
     'All tickets listed with status, date, priority, and last update.'],
    ['TC-EMP-HD-004', 'Add Comment to Ticket', 'Medium',
     '1. Login as Employee\n2. Open a ticket\n3. Type comment\n4. Click "Add Comment"',
     'Comment added to ticket thread. Timestamp and author shown.'],
    ['TC-EMP-HD-005', 'Rate Resolved Ticket', 'Low',
     '1. Login as Employee\n2. Open a RESOLVED ticket\n3. Click "Rate"\n4. Select rating (1-5)\n5. Submit',
     'Rating recorded. Ticket shows rating stars.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], emp_help)

doc.add_page_break()

# ═══ 5. HR MANAGER ROLE TEST CASES ═══
doc.add_heading('5. HR Manager Role - Test Cases', level=1)

doc.add_heading('5.1 Employee Management', level=2)
hr_emp = [
    ['TC-HR-EMP-001', 'Create New Employee', 'High',
     '1. Login as HR Manager\n2. Navigate to Employees > Add New\n3. Fill all required fields (name, email, phone, department, designation)\n4. Assign role\n5. Click "Create"',
     'Employee created successfully. Success message shown. Employee appears in directory.'],
    ['TC-HR-EMP-002', 'Bulk Employee Import', 'High',
     '1. Login as HR Manager\n2. Navigate to Employees > Bulk Import\n3. Upload Excel file with employee data\n4. Preview imported data\n5. Confirm import',
     'All employees from Excel imported. Success count shown. Undo option available.'],
    ['TC-HR-EMP-003', 'Undo Bulk Import', 'Medium',
     '1. Login as HR Manager\n2. Navigate to Employees > Import History\n3. Select a past import\n4. Click "Undo"',
     'All employees from that import removed. Import marked as undone.'],
    ['TC-HR-EMP-004', 'Edit Employee Profile', 'High',
     '1. Login as HR Manager\n2. Navigate to Employee Directory\n3. Select an employee\n4. Click "Edit"\n5. Modify details\n6. Save',
     'Employee profile updated. Changes reflected immediately.'],
    ['TC-HR-EMP-005', 'Deactivate Employee', 'High',
     '1. Login as HR Manager\n2. Navigate to Employee Directory\n3. Select an employee\n4. Click "Deactivate"\n5. Confirm',
     'Employee account deactivated. Employee cannot login. Offboarding record created.'],
    ['TC-HR-EMP-006', 'View Employee Profile', 'Medium',
     '1. Login as HR Manager\n2. Navigate to Employee Directory\n3. Click on employee name',
     'Full employee profile shown: personal info, bank details, documents, attendance, leave, payslips.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], hr_emp)

doc.add_heading('5.2 Leave Approval', level=2)
hr_leave = [
    ['TC-HR-LEA-001', 'View Pending Leave Requests', 'High',
     '1. Login as HR Manager\n2. Navigate to Leave > Pending Approvals',
     'All pending leave requests from employees listed with details (employee name, dates, type, reason).'],
    ['TC-HR-LEA-002', 'Approve Leave Request', 'High',
     '1. Login as HR Manager\n2. Navigate to Leave > Pending Approvals\n3. Select a request\n4. Click "Approve"\n5. Add optional comment',
     'Leave status changed to APPROVED. Employee notified. Leave balance deducted.'],
    ['TC-HR-LEA-003', 'Reject Leave Request', 'High',
     '1. Login as HR Manager\n2. Navigate to Leave > Pending Approvals\n3. Select a request\n4. Click "Reject"\n5. Enter rejection reason',
     'Leave status changed to REJECTED. Employee notified with rejection reason.'],
    ['TC-HR-LEA-004', 'Bulk Approve/Reject', 'Medium',
     '1. Login as HR Manager\n2. Navigate to Leave > Pending Approvals\n3. Select multiple requests\n4. Click "Bulk Action"\n5. Choose Approve or Reject',
     'Multiple leave requests processed in bulk. Individual notifications sent.'],
    ['TC-HR-LEA-005', 'View Leave Calendar', 'Medium',
     '1. Login as HR Manager\n2. Navigate to Leave > Calendar',
     'Calendar shows all team leaves by date. Color-coded by leave type.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], hr_leave)

doc.add_heading('5.3 Payroll Processing', level=2)
hr_pay = [
    ['TC-HR-PAY-001', 'Set Salary Structure', 'High',
     '1. Login as HR Manager\n2. Navigate to Payroll > Salary Structure\n3. Select employee\n4. Enter Basic, HRA, Conveyance, etc.\n5. Save',
     'Salary structure saved. Components shown in employee profile.'],
    ['TC-HR-PAY-002', 'Generate Payslip', 'High',
     '1. Login as HR Manager\n2. Navigate to Payroll > Generate Payslip\n3. Select employee and month\n4. Click "Generate"',
     'Payslip generated with all earnings and deductions. PDF available for download.'],
    ['TC-HR-PAY-003', 'Start Payroll Run', 'High',
     '1. Login as HR Manager\n2. Navigate to Payroll > Payroll Runs\n3. Select month\n4. Click "Start Run"',
     'Payroll run initiated. Status shows "IN_PROGRESS". All eligible employees included.'],
    ['TC-HR-PAY-004', 'Confirm Payroll Run', 'High',
     '1. Login as HR Manager\n2. Navigate to Payroll > Payroll Runs\n3. Select IN_PROGRESS run\n4. Review totals\n5. Click "Confirm"',
     'Payroll run confirmed. Status changes to CONFIRMED. Finance approval request sent.'],
    ['TC-HR-PAY-005', 'View Employee Payslips', 'Medium',
     '1. Login as HR Manager\n2. Navigate to Payroll > Payslips\n3. Select employee',
     'All payslips for the selected employee listed with month, year, net pay.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], hr_pay)

doc.add_page_break()

# ═══ 6. TEAM LEAD ROLE TEST CASES ═══
doc.add_heading('6. Team Lead / Manager Role - Test Cases', level=1)

doc.add_heading('6.1 Team Attendance', level=2)
tl_att = [
    ['TC-TL-ATT-001', 'View Team Attendance Today', 'High',
     '1. Login as Team Lead\n2. Navigate to Attendance > Team Today',
     'Shows team members with today\'s attendance status (present, absent, on leave).'],
    ['TC-TL-ATT-002', 'View Team Attendance Range', 'High',
     '1. Login as Team Lead\n2. Navigate to Attendance > Team\n3. Select date range\n4. View results',
     'Table shows team attendance for selected range with daily status.'],
    ['TC-TL-ATT-003', 'Only Own Team Visible', 'High',
     '1. Login as Team Lead\n2. Navigate to Attendance > Team\n3. Verify only team members shown',
     'Only direct reports visible. Other department employees NOT shown.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], tl_att)

doc.add_heading('6.2 Leave Approval', level=2)
tl_leave = [
    ['TC-TL-LEA-001', 'View My Approval Queue', 'High',
     '1. Login as Team Lead\n2. Navigate to Leave > My Queue',
     'Leave requests from direct reports shown in approval queue.'],
    ['TC-TL-LEA-002', 'Approve Team Leave', 'High',
     '1. Login as Team Lead\n2. Navigate to Leave > My Queue\n3. Select a request\n4. Click "Approve"',
     'Leave approved. Employee notified. Leave balance deducted. Calendar updated.'],
    ['TC-TL-LEA-003', 'Reject Team Leave', 'High',
     '1. Login as Team Lead\n2. Navigate to Leave > My Queue\n3. Select a request\n4. Click "Reject"\n5. Enter reason',
     'Leave rejected. Employee notified with reason.'],
    ['TC-TL-LEA-004', 'Cannot Approve Other Teams', 'High',
     '1. Login as Team Lead\n2. Attempt to access leave requests from other teams',
     'Other team leave requests NOT visible. Access denied.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], tl_leave)

doc.add_heading('6.3 Task Management', level=2)
tl_task = [
    ['TC-TL-TASK-001', 'Create Task for Team', 'High',
     '1. Login as Team Lead\n2. Navigate to Tasks > Create\n3. Enter title, description, assignee, due date\n4. Click "Create"',
     'Task created and assigned. Assignee notified via WebSocket.'],
    ['TC-TL-TASK-002', 'View Team Workload', 'Medium',
     '1. Login as Team Lead\n2. Navigate to Tasks > Team Workload',
     'Shows task count per team member with status breakdown (pending, in-progress, completed).'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], tl_task)

doc.add_page_break()

# ═══ 7. SYSTEM ADMIN ROLE TEST CASES ═══
doc.add_heading('7. System Admin Role - Test Cases', level=1)

doc.add_heading('7.1 Organization Management', level=2)
adm_org = [
    ['TC-ADM-ORG-001', 'Create Department', 'High',
     '1. Login as Admin\n2. Navigate to Organization > Departments\n3. Click "Add Department"\n4. Enter name and description\n5. Save',
     'Department created. Appears in dropdown menus for employee assignment.'],
    ['TC-ADM-ORG-002', 'Create Designation', 'High',
     '1. Login as Admin\n2. Navigate to Organization > Designations\n3. Click "Add Designation"\n4. Enter title and level\n5. Save',
     'Designation created. Available for employee role assignment.'],
    ['TC-ADM-ORG-003', 'Manage Office Locations', 'Medium',
     '1. Login as Admin\n2. Navigate to Organization > Locations\n3. Add/Edit/Delete locations',
     'Office locations managed. Geofence boundaries configurable.'],
    ['TC-ADM-ORG-004', 'Configure Holidays', 'Medium',
     '1. Login as Admin\n2. Navigate to Organization > Holidays\n3. Add holidays for the year\n4. Save',
     'Holidays configured. Employees see holidays in calendar. Attendance not required on holidays.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], adm_org)

doc.add_heading('7.2 System Settings', level=2)
adm_settings = [
    ['TC-ADM-SET-001', 'Module Management', 'High',
     '1. Login as Admin\n2. Navigate to Settings > Modules\n3. Toggle modules ON/OFF\n4. Save',
     'Modules enabled/disabled. Disabled modules hidden from all users.'],
    ['TC-ADM-SET-002', 'Branding Configuration', 'Medium',
     '1. Login as Admin\n2. Navigate to Settings > Branding\n3. Upload logo, set colors\n4. Save',
     'Application branding updated. Logo and colors reflected across UI.'],
    ['TC-ADM-SET-003', 'Global Announcements', 'Medium',
     '1. Login as Admin\n2. Navigate to Announcements\n3. Create announcement with entrance effect\n4. Publish',
     'Announcement shown to all users on login with configured animation.'],
    ['TC-ADM-SET-004', 'Audit Log Review', 'High',
     '1. Login as Admin\n2. Navigate to Audit > Logs\n3. Apply filters (date, user, action)\n4. View results',
     'Audit trail shown with timestamp, user, action, IP address, and details.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], adm_settings)

doc.add_heading('7.3 Cache & Data Management', level=2)
adm_cache = [
    ['TC-ADM-CACHE-001', 'Clear Redis Cache', 'Medium',
     '1. Login as Admin\n2. Navigate to Settings > Cache\n3. Click "Clear Cache"\n4. Confirm',
     'Redis cache cleared. All cached permissions and data refreshed on next request.'],
    ['TC-ADM-CACHE-002', 'View Login History', 'Medium',
     '1. Login as Admin\n2. Navigate to Audit > Login History',
     'All login attempts shown with timestamp, IP, browser, success/failure status.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], adm_cache)

doc.add_page_break()

# ═══ 8. CTO / EXECUTIVE ROLE TEST CASES ═══
doc.add_heading('8. CTO / Executive Role - Test Cases', level=1)

cto_tests = [
    ['TC-CTO-DASH-001', 'View Executive Dashboard', 'High',
     '1. Login as CEO/CTO\n2. Navigate to Executive Dashboard',
     'Dashboard shows KPIs: headcount, attrition rate, attendance %, payroll summary, ticket resolution.'],
    ['TC-CTO-DASH-002', 'View Org Insights', 'High',
     '1. Login as CEO/CTO\n2. Navigate to Org Insights',
     'Organization overview: department distribution, gender ratio, location spread, age distribution.'],
    ['TC-CTO-DASH-003', 'View Celebrations', 'Low',
     '1. Login as CEO/CTO\n2. Navigate to Celebrations',
     'Birthdays and work anniversaries shown with employee photos and dates.'],
    ['TC-CTO-DASH-004', 'Download Reports', 'Medium',
     '1. Login as CEO/CTO\n2. Navigate to Reports\n3. Select attendance report\n4. Click "Download Excel"',
     'Excel report downloaded with attendance data for all employees.'],
    ['TC-CTO-ACC-001', 'Cannot Access Payroll Operations', 'High',
     '1. Login as CEO/CTO\n2. Attempt to access Payroll > Generate Payslip',
     'Access denied. Payroll generation not available for CEO role.'],
    ['TC-CTO-ACC-002', 'Cannot Create Employees', 'High',
     '1. Login as CEO/CTO\n2. Attempt to access Employees > Add New',
     'Employee creation form not accessible. Access denied.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], cto_tests)

doc.add_page_break()

# ═══ 9-18 MODULE TEST CASES ═══
doc.add_heading('9. Attendance Module - Cross-Role Test Cases', level=1)
att_cross = [
    ['TC-ATT-XROLE-001', 'Geofence Validation', 'High',
     '1. Login as any user\n2. Attempt punch-in from outside geofence area',
     'Punch-in rejected. Warning: "Outside designated work area."'],
    ['TC-ATT-XROLE-002', 'Double Punch Prevention', 'High',
     '1. Login as any user\n2. Punch in\n3. Attempt to punch in again without punching out',
     'Second punch-in rejected. Message: "Already punched in. Punch out first."'],
    ['TC-ATT-XROLE-003', 'WebSocket Attendance Update', 'Medium',
     '1. Login as Employee (punch in)\n2. Login as Team Lead simultaneously\n3. Check team attendance',
     'Team Lead sees real-time attendance update via WebSocket.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], att_cross)

doc.add_heading('10. Leave Management - Cross-Role Test Cases', level=1)
leave_cross = [
    ['TC-LEA-XROLE-001', 'Leave Balance Auto-Deduction', 'High',
     '1. Employee applies leave (3 days)\n2. Manager approves\n3. Check leave balance',
     'Leave balance automatically reduced by 3 days.'],
    ['TC-LEA-XROLE-002', 'Leave Balance Restore on Cancel', 'High',
     '1. Employee applies leave\n2. Manager approves\n3. Employee cancels\n4. Check balance',
     'Leave balance restored to original count.'],
    ['TC-LEA-XROLE-003', 'LOP Calculation for Payroll', 'High',
     '1. Employee exceeds leave quota\n2. HR runs payroll\n3. Check payslip',
     'LOP (Loss of Pay) days deducted from salary. LOP amount shown on payslip.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], leave_cross)

doc.add_heading('11. Payroll Module - Cross-Role Test Cases', level=1)
pay_cross = [
    ['TC-PAY-XROLE-001', 'Payroll Run Lifecycle', 'High',
     '1. HR sets salary for all employees\n2. HR generates payslips\n3. HR starts payroll run\n4. HR confirms\n5. Finance approves',
     'Full payroll lifecycle completes. All payslips generated and accessible.'],
    ['TC-PAY-XROLE-002', 'Payslip PDF Generation', 'High',
     '1. HR generates payslip\n2. Employee downloads PDF',
     'PDF contains correct company name (PIXOUS), employee details, earnings, deductions, net pay.'],
    ['TC-PAY-XROLE-003', 'Salary Confidentiality', 'High',
     '1. Login as Employee A\n2. Attempt to view Employee B\'s salary',
     'Employee A can only see own salary. Other salaries not accessible.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], pay_cross)

doc.add_heading('12. Asset Management - Test Cases', level=1)
ast_tests = [
    ['TC-AST-ADM-001', 'Create Asset', 'High',
     '1. Login as Asset Manager\n2. Navigate to Assets > Add New\n3. Enter asset details (name, type, serial number)\n4. Save',
     'Asset created with unique code. QR code generated.'],
    ['TC-AST-ADM-002', 'Allocate Asset to Employee', 'High',
     '1. Login as Asset Manager\n2. Select asset\n3. Click "Allocate"\n4. Select employee\n5. Confirm',
     'Asset allocated. Employee notified. Asset status changed to ALLOCATED.'],
    ['TC-AST-EMP-001', 'Acknowledge Asset Receipt', 'Medium',
     '1. Login as Employee\n2. Navigate to My Assets\n3. Click "Acknowledge" on allocated asset',
     'Asset receipt acknowledged. Status updated. Allocation record confirmed.'],
    ['TC-AST-ADM-003', 'Return Asset', 'High',
     '1. Login as Asset Manager\n2. Select allocated asset\n3. Click "Return"\n4. Confirm',
     'Asset returned. Status changed to AVAILABLE. Employee notified.'],
    ['TC-AST-EMP-002', 'Generate QR Code', 'Low',
     '1. Login as any user\n2. Navigate to asset detail\n3. Click "QR Code"',
     'QR code PNG generated and downloadable with asset details embedded.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], ast_tests)

doc.add_heading('13. Helpdesk / Ticketing - Test Cases', level=1)
hd_tests = [
    ['TC-HD-EMP-001', 'Raise Ticket with Attachment', 'High',
     '1. Login as Employee\n2. Create new ticket\n3. Enter subject and description\n4. Upload attachment\n5. Submit',
     'Ticket created with attachment. Agent notified.'],
    ['TC-HD-AGENT-001', 'Pick Up Ticket', 'High',
     '1. Login as Admin/Agent\n2. Navigate to Helpdesk > My Queue\n3. Select a ticket',
     'Ticket details loaded. Comments and history visible.'],
    ['TC-HD-AGENT-002', 'Change Ticket Status', 'High',
     '1. Login as Admin/Agent\n2. Open ticket\n3. Change status to IN_PROGRESS\n4. Save',
     'Status updated. Employee notified of status change.'],
    ['TC-HD-AGENT-003', 'Resolve Ticket', 'High',
     '1. Login as Admin/Agent\n2. Open ticket\n3. Change status to RESOLVED\n4. Add resolution comment\n5. Save',
     'Ticket resolved. Employee notified. Rating prompt shown.'],
    ['TC-HD-EMP-002', 'Rate Resolved Ticket', 'Medium',
     '1. Login as Employee\n2. Open resolved ticket\n3. Click "Rate"\n4. Select 4 stars\n5. Submit',
     'Rating saved. Ticket shows 4-star rating.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], hd_tests)

doc.add_heading('14. Community & Chat - Test Cases', level=1)
chat_tests = [
    ['TC-CHAT-001', 'Send Text Message', 'High',
     '1. Login as any user\n2. Open a chat group\n3. Type message\n4. Click Send',
     'Message sent. Timestamp and sender name shown. Other members see message in real-time.'],
    ['TC-CHAT-002', 'Send Voice Message', 'Medium',
     '1. Login as any user\n2. Open a chat group\n3. Click microphone icon\n4. Record voice\n5. Send',
     'Voice message sent. Playable by other members. Duration shown.'],
    ['TC-CHAT-003', 'Send File Attachment', 'Medium',
     '1. Login as any user\n2. Open a chat group\n3. Click attach icon\n4. Select file\n5. Send',
     'File uploaded and displayed in chat. Downloadable by other members.'],
    ['TC-CHAT-004', 'React to Message', 'Low',
     '1. Login as any user\n2. Hover over a message\n3. Click reaction emoji',
     'Reaction added. Reaction count shown under message.'],
    ['TC-CHAT-005', 'Pin Message', 'Medium',
     '1. Login as Admin/HR\n2. Select message\n3. Click "Pin"',
     'Message pinned. Pinned messages section shows pinned message.'],
    ['TC-CHAT-006', 'Direct Message (1:1)', 'High',
     '1. Login as any user\n2. Navigate to Contacts\n3. Click on a contact\n4. Send message',
     '1:1 chat opened. Messages exchanged in private.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], chat_tests)

doc.add_heading('15. Task Management - Test Cases', level=1)
task_tests = [
    ['TC-TASK-001', 'Create Task', 'High',
     '1. Login as Manager/HR\n2. Navigate to Tasks > Create\n3. Enter title, description, assignee, due date\n4. Save',
     'Task created. Assignee notified. Task appears in assignee\'s task list.'],
    ['TC-TASK-002', 'Update Task Status', 'High',
     '1. Login as assigned employee\n2. Open task\n3. Change status to IN_PROGRESS\n4. Save',
     'Task status updated. Manager sees status change.'],
    ['TC-TASK-003', 'Complete Task', 'High',
     '1. Login as assigned employee\n2. Open task\n3. Change status to COMPLETED\n4. Add completion notes\n5. Save',
     'Task marked complete. Manager notified. Work report can reference task.'],
    ['TC-TASK-004', 'View Team Workload', 'Medium',
     '1. Login as Manager\n2. Navigate to Tasks > Team Workload',
     'Chart/table shows tasks per team member by status.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], task_tests)

doc.add_heading('16. Dashboard & Reports - Test Cases', level=1)
dash_tests = [
    ['TC-DASH-EMP-001', 'Personal Dashboard', 'High',
     '1. Login as Employee\n2. View dashboard',
     'Shows: attendance summary, leave balance, pending tasks, notifications, celebrations.'],
    ['TC-DASH-EXEC-001', 'Executive Dashboard', 'High',
     '1. Login as CEO/Admin\n2. Navigate to Executive Dashboard',
     'Shows: headcount, attrition, attendance %, payroll cost, ticket resolution rate.'],
    ['TC-DASH-ORG-001', 'Org Insights', 'Medium',
     '1. Login as HR/Admin\n2. Navigate to Org Insights',
     'Charts showing: department distribution, gender ratio, location spread, age demographics.'],
    ['TC-RPT-001', 'Attendance Report Excel', 'Medium',
     '1. Login as HR/Admin\n2. Navigate to Reports > Attendance\n3. Select date range\n4. Click "Export Excel"',
     'Excel file downloaded with attendance data for all employees in selected range.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], dash_tests)

doc.add_heading('17. Onboarding - Test Cases', level=1)
onb_tests = [
    ['TC-ONB-001', 'Create Onboarding Checklist', 'High',
     '1. Login as HR\n2. Navigate to Onboarding\n3. Create checklist for new hire\n4. Add tasks\n5. Save',
     'Checklist created with tasks. Assignee notified.'],
    ['TC-ONB-002', 'Track Onboarding Progress', 'Medium',
     '1. Login as HR\n2. Navigate to Onboarding > Progress',
     'Progress bar and task completion status shown for each new hire.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], onb_tests)

doc.add_page_break()

# ═══ 18. CROSS-BROWSER ═══
doc.add_heading('18. Cross-Browser & Responsive Test Cases', level=1)
browser_tests = [
    ['TC-BRW-001', 'Chrome Desktop', 'High', '1. Open application in Chrome 120+\n2. Test all major features', 'All features work correctly. UI renders properly.'],
    ['TC-BRW-002', 'Firefox Desktop', 'High', '1. Open application in Firefox 120+\n2. Test all major features', 'All features work correctly. No browser-specific bugs.'],
    ['TC-BRW-003', 'Safari Desktop', 'Medium', '1. Open application in Safari 17+\n2. Test all major features', 'All features work correctly.'],
    ['TC-BRW-004', 'Edge Desktop', 'Medium', '1. Open application in Edge 120+\n2. Test all major features', 'All features work correctly.'],
    ['TC-BRW-005', 'Mobile Responsive', 'High', '1. Open application on mobile browser\n2. Test navigation, forms, tables', 'Responsive layout adapts. Touch targets adequate. Forms usable.'],
    ['TC-BRW-006', 'PWA Installation', 'Medium', '1. Open application in Chrome\n2. Click "Install App"\n3. Verify PWA launches', 'PWA installed. Opens in standalone window. Offline indicators work.'],
]
add_styled_table(doc, ['TC ID', 'Test Case', 'Priority', 'Steps', 'Expected Result'], browser_tests)

doc.add_page_break()

# ═══ 19. REGRESSION SUMMARY ═══
doc.add_heading('19. Regression Test Summary', level=1)
doc.add_paragraph(
    'The following regression test matrix defines the critical paths that must be re-tested '
    'after every major release or hotfix deployment.'
)
regression = [
    ['REG-001', 'Login & Logout', 'AUTH', 'All roles can login and logout successfully'],
    ['REG-002', 'Attendance Punch In/Out', 'ATT', 'GPS and face punch work for all roles'],
    ['REG-003', 'Leave Apply & Approve', 'LEA', 'Full leave lifecycle from apply to balance deduction'],
    ['REG-004', 'Payroll Generation', 'PAY', 'Salary, payslip generation, PDF download'],
    ['REG-005', 'Helpdesk Ticket Lifecycle', 'HD', 'Raise -> Assign -> Resolve -> Rate'],
    ['REG-006', 'Chat Messaging', 'CHAT', 'Send text, voice, file in groups and DMs'],
    ['REG-007', 'Role-Based Access', 'RBAC', 'Each role can only access permitted features'],
    ['REG-008', 'Dashboard Data Accuracy', 'DASH', 'Dashboard numbers match actual data'],
    ['REG-009', 'WebSocket Notifications', 'WS', 'Real-time push notifications delivered'],
    ['REG-010', 'API Authentication', 'API', 'JWT validation, token refresh, 401 handling'],
]
add_styled_table(doc, ['ID', 'Test Path', 'Module', 'Description'], regression)

doc.add_page_break()

# ═══ 20. DEFECT TRACKING TEMPLATE ═══
doc.add_heading('20. Defect Tracking Template', level=1)
doc.add_paragraph(
    'Use the following template for logging defects found during testing:'
)
defect_template = [
    ['Defect ID', 'Auto-generated (BUG-XXX)'],
    ['Title', 'Brief description of the defect'],
    ['Module', 'Affected module (Auth, Attendance, Leave, etc.)'],
    ['Severity', 'Critical / High / Medium / Low'],
    ['Priority', 'P1 (Immediate) / P2 (High) / P3 (Medium) / P4 (Low)'],
    ['Steps to Reproduce', 'Numbered steps to reproduce the defect'],
    ['Expected Result', 'What should have happened'],
    ['Actual Result', 'What actually happened'],
    ['Environment', 'Browser, OS, device, URL'],
    ['Screenshots', 'Attached screenshots or screen recordings'],
    ['Reported By', 'Tester name and date'],
    ['Assigned To', 'Developer name'],
    ['Status', 'Open / In Progress / Fixed / Verified / Closed / Reopened'],
    ['Fix Version', 'Version where the fix is applied'],
    ['Regression Status', 'Pass / Fail after fix verification'],
]
add_styled_table(doc, ['Field', 'Description'], defect_template)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PIXOUS Technologies Pvt. Ltd. | Confidential')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# ═══ APPENDIX A: TEST EXECUTION CHECKLIST ═══
doc.add_heading('Appendix A: Test Execution Checklist', level=1)
doc.add_paragraph(
    'Before each release, the QA team must complete the following checklist to ensure all critical '
    'areas have been tested and documented. This checklist serves as the final gate before deployment.'
)
doc.add_paragraph(
    'All test cases must be executed in the staging environment with production-like data. '
    'Any failures must be documented with screenshots and reproduction steps before the release '
    'can proceed to production deployment.'
)

doc.add_heading('Pre-Release Checklist', level=2)
checklist = [
    ['1', 'All P1 and P2 test cases passed', 'Mandatory'],
    ['2', 'No open Critical/High severity defects', 'Mandatory'],
    ['3', 'Cross-browser testing completed (Chrome, Firefox, Safari, Edge)', 'Mandatory'],
    ['4', 'Mobile responsive testing completed', 'Mandatory'],
    ['5', 'API regression suite passed', 'Mandatory'],
    ['6', 'Database migration verified (Flyway)', 'Mandatory'],
    ['7', 'WebSocket real-time notifications verified', 'Mandatory'],
    ['8', 'JWT authentication flow verified', 'Mandatory'],
    ['9', 'Role-based access control verified for all roles', 'Mandatory'],
    ['10', 'Payslip PDF generation and download verified', 'Mandatory'],
    ['11', 'Face recognition attendance verified', 'Conditional'],
    ['12', 'GPS geofence validation verified', 'Conditional'],
    ['13', 'Kafka chat streaming verified', 'Conditional'],
    ['14', 'Performance test results within thresholds', 'Recommended'],
    ['15', 'Security scan completed with no critical findings', 'Recommended'],
]
add_styled_table(doc, ['#', 'Checklist Item', 'Status'], checklist)

doc.add_heading('Test Data Requirements', level=2)
doc.add_paragraph(
    'The following test data must be prepared and maintained in the staging environment before '
    'test execution begins. Test data should represent realistic scenarios while ensuring no '
    'real employee or financial data is used.'
)
test_data = [
    ['Employee Accounts', 'Minimum 20 test employees across all roles (HR, TL, Employee, Finance)'],
    ['Departments', 'Minimum 5 departments (Engineering, HR, Finance, Operations, Marketing)'],
    ['Designations', 'Minimum 8 designations (Intern, Junior, Senior, Lead, Manager, Director, VP, C-Level)'],
    ['Leave Types', 'All configured leave types (Casual, Sick, Earned, Maternity, Paternity, LOP)'],
    ['Assets', 'Minimum 10 test assets with various statuses (Available, Allocated, Returned)'],
    ['Helpdesk Tickets', 'Minimum 15 tickets in various states (Open, In Progress, Resolved, Closed)'],
    ['Chat Messages', 'Test messages in at least 3 community groups and 2 direct message threads'],
    ['Tasks', 'Minimum 10 tasks assigned across team members in various statuses'],
    ['Payroll Data', 'Test salary structures for all roles with at least 2 months of payslip history'],
    ['Attendance Records', 'Minimum 30 days of attendance data with various patterns (present, absent, late, overtime)'],
]
add_styled_table(doc, ['Data Type', 'Requirement'], test_data)

doc.add_heading('Defect Severity Definitions', level=2)
doc.add_paragraph(
    'Understanding defect severity is critical for proper prioritization and resolution. '
    'The following definitions guide the QA team in categorizing defects found during testing.'
)
severity = [
    ['Critical', 'System crash, data loss, security breach, or complete feature failure. No workaround available. Blocks release.'],
    ['High', 'Major feature not working as designed, significant data incorrect, or workflow blocked. Workaround may exist but is impractical.'],
    ['Medium', 'Feature partially working, cosmetic issues affecting usability, or minor data discrepancies. Workaround available.'],
    ['Low', 'Typo, cosmetic alignment issue, minor UI inconsistency, or enhancement request. Does not affect functionality.'],
]
add_styled_table(doc, ['Severity', 'Definition & Impact'], severity)

doc.add_heading('Priority Definitions', level=2)
doc.add_paragraph(
    'Priority determines the order in which defects should be fixed. Priority is independent '
    'of severity — a low-severity bug in a critical path may have high priority.'
)
priority_def = [
    ['P1 - Immediate', 'Fix within 24 hours. Blocks release or critical business process.'],
    ['P2 - High', 'Fix within 1 sprint (2 weeks). Major feature impacted, workaround exists but is painful.'],
    ['P3 - Medium', 'Fix within 2 sprints. Minor feature impacted, workaround is acceptable.'],
    ['P4 - Low', 'Fix when convenient. Cosmetic or enhancement, no functional impact.'],
]
add_styled_table(doc, ['Priority', 'Definition & Timeline'], priority_def)

# ═══ APPENDIX C: TEST EXECUTION LOG TEMPLATE ═══
doc.add_heading('Appendix C: Test Execution Log Template', level=1)
doc.add_paragraph(
    'Use this template to log each test execution session. Record the date, tester name, '
    'environment details, and results for each test case executed during the session.'
)
doc.add_paragraph(
    'The execution log serves as an audit trail for test coverage and helps track progress '
    'toward the release readiness criteria. All logs should be stored in the project documentation repository.'
)

exec_log = [
    ['Execution ID', 'Auto-generated (EXEC-XXX)'],
    ['Execution Date', 'Date and time of test execution'],
    ['Tester Name', 'Name of the QA engineer executing tests'],
    ['Environment', 'Staging / Dev / Production'],
    ['Browser & Version', 'Browser used for execution'],
    ['OS', 'Operating system used'],
    ['Test Suite', 'Module or feature area being tested'],
    ['Total Test Cases', 'Number of test cases in the suite'],
    ['Passed', 'Number of test cases that passed'],
    ['Failed', 'Number of test cases that failed'],
    ['Blocked', 'Number of test cases that could not be executed'],
    ['Pass Rate', 'Percentage of passed test cases'],
    ['Defects Found', 'Number of new defects identified'],
    ['Execution Time', 'Total time taken for the test suite'],
    ['Notes', 'Any additional observations or comments'],
]
add_styled_table(doc, ['Field', 'Description'], exec_log)

doc.add_heading('Test Coverage Metrics', level=2)
doc.add_paragraph(
    'Track the following metrics to ensure adequate test coverage across the application. '
    'These metrics should be reported to project stakeholders at the end of each test cycle.'
)
metrics = [
    ['Requirements Coverage', 'Percentage of requirements covered by test cases', 'Target: 100%'],
    ['Code Coverage', 'Percentage of code paths covered by automated tests', 'Target: >= 80%'],
    ['Defect Detection Rate', 'Percentage of defects found before production', 'Target: >= 90%'],
    ['Test Case Execution Rate', 'Percentage of planned test cases executed', 'Target: 100%'],
    ['Pass Rate', 'Percentage of executed test cases that passed', 'Target: >= 95%'],
    ['Defect Recurrence Rate', 'Percentage of defects that reappear after fix', 'Target: <= 5%'],
    ['Mean Time to Fix', 'Average time from defect report to resolution', 'Target: <= 3 days'],
    ['Test Environment Availability', 'Percentage of time test environment is accessible', 'Target: >= 95%'],
]
add_styled_table(doc, ['Metric', 'Definition', 'Target'], metrics)

# ═══ APPENDIX B: TEST ENVIRONMENT SETUP ═══
doc.add_heading('Appendix B: Test Environment Setup Guide', level=1)
doc.add_paragraph(
    'This appendix provides instructions for setting up the test environment to replicate '
    'production conditions for comprehensive testing.'
)

doc.add_heading('Local Development Setup', level=2)
doc.add_paragraph('Step 1: Clone the repository from GitHub.')
doc.add_paragraph('Step 2: Install Docker and Docker Compose on the test machine.')
doc.add_paragraph('Step 3: Run docker-compose.yml to start MySQL 8.4, Redis 7, and Kafka 3.9.')
doc.add_paragraph('Step 4: Start the Spring Boot backend with the dev profile enabled.')
doc.add_paragraph('Step 5: Start the React frontend with npm run dev or vite dev server.')
doc.add_paragraph('Step 6: Verify all services are running and accessible on their respective ports.')

doc.add_heading('Test Account Credentials', level=2)
doc.add_paragraph(
    'The following test accounts are pre-configured in the staging environment. '
    'These credentials should not be used in production and must be rotated after each test cycle.'
)
accounts = [
    ['System Admin', 'testadmin', 'admin123', 'SUPER_ADMIN', 'Full system access'],
    ['HR Manager', 'testhr', 'hr123', 'IT_HR', 'HR and payroll operations'],
    ['Team Lead', 'testtl', 'tl123', 'IT_MGR', 'Team approvals and management'],
    ['Finance Officer', 'testfin', 'fin123', 'IT_FIN', 'Payroll finance approval'],
    ['CEO/Executive', 'testceo', 'ceo123', 'IT_CEO', 'Executive dashboard access'],
    ['Asset Manager', 'testast', 'ast123', 'IT_AST', 'Asset lifecycle management'],
    ['Employee 1', 'testemp1', 'emp123', 'IT_EMP', 'Standard employee self-service'],
    ['Employee 2', 'testemp2', 'emp123', 'IT_EMP', 'Standard employee self-service'],
]
add_styled_table(doc, ['Role', 'Username', 'Password', 'Role Code', 'Access Level'], accounts)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PIXOUS Technologies Pvt. Ltd. | Confidential')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save('PIXOUS_HR_Portal_Test_Cases_Document.docx')
print("Document 2 generated successfully: PIXOUS_HR_Portal_Test_Cases_Document.docx")
