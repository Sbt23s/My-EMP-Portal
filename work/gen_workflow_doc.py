"""Generate PIXOUS HR Portal - Full Workflow Document"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper: Add Logo ──
def add_logo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('pixous_logo.png', width=Inches(2.0))
    return p

# ── Helper: Add Cover Title ──
def add_cover(doc):
    add_logo(doc)
    doc.add_paragraph()
    title = doc.add_heading('PIXOUS HR Portal', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading('Full Application Workflow Document', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(0, 102, 153)
    
    doc.add_paragraph()
    
    # Info table
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ('Document Type', 'Full Workflow Document'),
        ('Version', '1.0'),
        ('Date', datetime.date.today().strftime('%B %d, %Y')),
        ('Prepared By', 'PIXOUS Technologies QA Team'),
        ('Classification', 'Confidential - Client Deliverable'),
    ]
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for cell in t.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_page_break()

# ── Helper: Styled Table ──
def add_styled_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear'
        })
        shading.append(shd)
    # Rows
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
add_cover(doc)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. System Overview & Architecture',
    '3. Technology Stack',
    '4. User Roles & Permissions Matrix',
    '5. HR Manager - Full Workflow',
    '6. CTO / Executive - Full Workflow',
    '7. Team Lead / Manager - Full Workflow',
    '8. System Admin / Company Admin - Full Workflow',
    '9. Employee - Full Workflow',
    '10. Module-wise Workflow Details',
    '  10.1 Authentication & Login Flow',
    '  10.2 Attendance Management',
    '  10.3 Leave Management',
    '  10.4 Payroll Management',
    '  10.5 Asset Management',
    '  10.6 Helpdesk / Ticketing',
    '  10.7 Community & Chat',
    '  10.8 Task Management',
    '  10.9 Dashboard & Reports',
    '  10.10 Onboarding & Offboarding',
    '  10.11 Performance Management',
    '  10.12 Announcements & Calendar',
    '11. Cross-Role Workflow Diagrams',
    '12. Data Flow Architecture',
    '13. API Workflow Summary',
    '14. Document Revision History',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal is a comprehensive, enterprise-grade Human Resource Management System '
    'designed to streamline and automate all HR operations for organizations across IT and Civil/Construction '
    'industries. The platform serves as a unified solution covering employee lifecycle management, '
    'attendance tracking, leave management, payroll processing, asset management, helpdesk support, '
    'internal communication, task management, performance reviews, and executive reporting.'
)
doc.add_paragraph(
    'This document provides a complete workflow analysis of the application, detailing every module, '
    'every user role, every interaction path, and every business process supported by the system. '
    'It serves as the definitive reference for QA testing, client demonstration, and ongoing '
    'maintenance of the PIXOUS HR Portal.'
)

doc.add_heading('Key Highlights', level=2)
highlights = [
    'Multi-industry support: IT and Civil/Construction industry role templates',
    '5+ distinct user roles with granular RBAC permissions (14+ permission codes)',
    '25+ functional modules covering the complete employee lifecycle',
    'Real-time WebSocket notifications and chat capabilities',
    'Face recognition and GPS geofenced attendance',
    'Multi-platform: Web (React), Mobile (React Native + Flutter), REST API (Spring Boot)',
    '101 database migrations with Flyway versioning',
    'JWT-based authentication with token rotation and account lockout',
    'Docker-based deployment with production-ready configurations',
    'SaaS multi-tenant architecture with module-level toggle per company',
]
for h in highlights:
    doc.add_paragraph(h, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. SYSTEM OVERVIEW & ARCHITECTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. System Overview & Architecture', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal follows a modern microservices-inspired architecture with a monolithic '
    'Spring Boot backend, React SPA frontend, React Native and Flutter mobile apps, and a Python-based '
    'analytics service for face recognition and OCR processing.'
)

doc.add_heading('Architecture Layers', level=2)
add_styled_table(doc,
    ['Layer', 'Technology', 'Responsibility'],
    [
        ['Presentation (Web)', 'React 19 + TypeScript + Vite', 'Single-page application with Tailwind CSS'],
        ['Presentation (Mobile)', 'React Native / Flutter', 'Native mobile apps with GPS & face auth'],
        ['API Gateway', 'Spring Boot 3.5 + Spring Security', 'REST API, JWT auth, RBAC enforcement'],
        ['Business Logic', 'Java 17 + Spring Data JPA', '25+ module controllers & services'],
        ['Data Layer', 'MySQL 8.4 + Flyway Migrations', 'Relational persistence, schema versioning'],
        ['Cache Layer', 'Redis 7 (optional)', 'Session cache, permission caching'],
        ['Message Queue', 'Apache Kafka 3.9 (KRaft)', 'Chat event streaming'],
        ['Real-time', 'WebSocket (STOMP/SockJS)', 'Push notifications, live presence'],
        ['Analytics', 'Python FastAPI + face_recognition', 'Face verification, OCR, analytics'],
        ['Infrastructure', 'Docker + Docker Compose', 'Container orchestration, production deploy'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Technology Stack', level=1)

doc.add_heading('3.1 Backend Stack', level=2)
add_styled_table(doc,
    ['Component', 'Technology', 'Version'],
    [
        ['Framework', 'Spring Boot', '3.5'],
        ['Language', 'Java', '17'],
        ['Build Tool', 'Maven', '3.x'],
        ['Database', 'MySQL', '8.4'],
        ['ORM', 'Hibernate / Spring Data JPA', 'JPA 3.1'],
        ['Migrations', 'Flyway', '101 migrations (V1-V101)'],
        ['Security', 'Spring Security + JWT (jjwt)', '0.12.6'],
        ['Cache', 'Redis (Lettuce client)', '7.x'],
        ['Messaging', 'Apache Kafka', '3.9 (KRaft)'],
        ['WebSocket', 'Spring WebSocket (STOMP)', '-'],
        ['PDF Generation', 'OpenPDF', '2.0.3'],
        ['Excel Import/Export', 'Apache POI', '5.3.0'],
        ['QR Codes', 'ZXing', '3.5.3'],
        ['Image Processing', 'Thumbnailator', '0.4.20'],
        ['SMS', 'Twilio + Fast2SMS', 'Dual provider'],
        ['API Docs', 'SpringDoc OpenAPI', '2.8.9'],
        ['Code Gen', 'Lombok + MapStruct', '1.18.38 / 1.6.3'],
        ['Container', 'Docker (multi-stage)', 'Eclipse Temurin 17 JRE'],
    ]
)

doc.add_heading('3.2 Frontend Stack', level=2)
add_styled_table(doc,
    ['Component', 'Technology', 'Version'],
    [
        ['Framework', 'React', '19'],
        ['Language', 'TypeScript', '-'],
        ['Build Tool', 'Vite', '6.0.5'],
        ['Routing', 'React Router', 'v7'],
        ['State Management', 'TanStack Query', 'v5'],
        ['Table', 'TanStack Table', 'v8'],
        ['HTTP Client', 'Axios', '1.7'],
        ['Forms', 'React Hook Form + Zod', '-'],
        ['Styling', 'Tailwind CSS', '3.4'],
        ['Charts', 'Recharts', '2.15'],
        ['Animations', 'Framer Motion', '11'],
        ['WebSocket', 'STOMP (stompjs)', '-'],
        ['PWA', 'vite-plugin-pwa', '-'],
        ['Testing', 'Vitest + Testing Library', '-'],
        ['Web Server', 'Nginx', '1.27'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. USER ROLES & PERMISSIONS MATRIX
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. User Roles & Permissions Matrix', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal implements a comprehensive Role-Based Access Control (RBAC) system with '
    '14+ distinct permission codes and 12+ pre-configured role templates. The system supports both '
    'IT and Civil/Construction industry role configurations.'
)

doc.add_heading('4.1 System Roles', level=2)
add_styled_table(doc,
    ['Role Code', 'Display Name', 'Industry', 'Description'],
    [
        ['SUPER_ADMIN', 'Platform Super Admin', 'Both', 'Full configuration and system access'],
        ['COMPANY_ADMIN', 'Company Admin', 'Both', 'Company-level full access'],
        ['IT_HR', 'HR Manager', 'IT', 'HR & payroll administration'],
        ['IT_MGR', 'IT Manager / Team Lead', 'IT', 'Team approvals & appraisals'],
        ['IT_FIN', 'Finance Officer', 'IT', 'Payroll & expense approvals'],
        ['IT_CEO', 'CEO / Executive', 'IT', 'Executive dashboards & policy'],
        ['IT_AST', 'IT Asset Manager', 'IT', 'IT asset lifecycle management'],
        ['IT_EMP', 'IT Employee', 'IT', 'Self-service employee portal'],
        ['CV_HR', 'Civil HR Manager', 'Civil', 'Labour compliance & workforce'],
        ['CV_SUP', 'Civil Supervisor', 'Civil', 'Site approvals & safety'],
        ['CV_ADM', 'Facilities Admin', 'Civil', 'Sites & facilities management'],
        ['CV_AST', 'Civil Asset Manager', 'Civil', 'Machinery & materials'],
        ['CV_EMP', 'Civil Site Employee', 'Civil', 'Site self-service portal'],
        ['BOARD_ADMIN', 'Board Admin', 'Both', 'Board-level oversight'],
    ]
)

doc.add_heading('4.2 Permission Codes', level=2)
add_styled_table(doc,
    ['Permission Code', 'Description', 'Applies To'],
    [
        ['USER_MANAGE', 'Create, edit, deactivate employees', 'HR, Admin'],
        ['LEAVE_APPLY', 'Apply for leave (self)', 'All employees'],
        ['LEAVE_APPROVE', 'Approve/reject leave requests', 'Manager, HR, Admin'],
        ['ATTENDANCE_SELF', 'Mark own attendance (punch in/out)', 'All employees'],
        ['ATTENDANCE_TEAM', 'View team attendance records', 'Manager, HR, Admin'],
        ['PAYROLL_VIEW', 'View payslips and salary data', 'Employee (own), HR, Admin'],
        ['PAYROLL_RUN', 'Generate payslips and run payroll', 'HR, Admin'],
        ['PAYROLL_APPROVE', 'Finance approval for payroll runs', 'Finance, Admin'],
        ['ASSET_MANAGE', 'Create, allocate, return assets', 'Asset Manager, Admin'],
        ['HELPDESK_RAISE', 'Create support tickets', 'All employees'],
        ['HELPDESK_AGENT', 'Resolve support tickets', 'Asset Manager, Admin'],
        ['REPORT_VIEW', 'View attendance and other reports', 'Manager, HR, Admin'],
        ['DASHBOARD_EXEC', 'View executive KPI dashboard', 'CEO, Admin'],
        ['ORG_MANAGE', 'Manage departments, designations, etc.', 'HR, Admin'],
        ['COMMUNITY_MANAGE', 'Manage chat groups and communities', 'HR, Admin'],
    ]
)

doc.add_heading('4.3 Role-Feature Access Matrix', level=2)
add_styled_table(doc,
    ['Feature', 'Employee', 'Team Lead', 'HR Manager', 'Finance', 'CEO', 'Admin'],
    [
        ['Personal Dashboard', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Executive Dashboard', 'No', 'No', 'No', 'No', 'Yes', 'Yes'],
        ['Org Insights', 'No', 'No', 'Yes', 'No', 'No', 'Yes'],
        ['Attendance (Self)', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Attendance (Team)', 'No', 'Own team', 'All', 'No', 'No', 'All'],
        ['Face Recognition', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['GPS Geofence', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Leave (Apply)', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Leave (Approve)', 'No', 'Own team', 'All', 'No', 'No', 'All'],
        ['Leave Policy Mgmt', 'No', 'No', 'Yes', 'No', 'No', 'Yes'],
        ['Payslips (View)', 'Own', 'Own', 'All', 'Own', 'No', 'All'],
        ['Payslip Generation', 'No', 'No', 'Yes', 'No', 'No', 'Yes'],
        ['Payroll Runs', 'No', 'No', 'Yes', 'Approve', 'No', 'Yes'],
        ['Salary Structure', 'No', 'No', 'Yes', 'View', 'No', 'Yes'],
        ['Assets (Manage)', 'No', 'No', 'No', 'No', 'No', 'Yes'],
        ['Assets (View Own)', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Helpdesk (Raise)', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Helpdesk (Resolve)', 'No', 'No', 'No', 'No', 'No', 'Yes'],
        ['Employee Directory', 'Limited', 'Own team', 'All', 'Limited', 'All', 'All'],
        ['Create Employee', 'No', 'No', 'Yes', 'No', 'No', 'Yes'],
        ['Bulk Import', 'No', 'No', 'Yes', 'No', 'No', 'Yes'],
        ['Chat / Communities', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Task Management', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Work Reports', 'Own', 'Own team', 'All', 'Own', 'All', 'All'],
        ['TA Expenses', 'Own', 'Own', 'Own+Approve', 'Own', 'All', 'All'],
        ['Safety Incidents', 'Report', 'Report', 'View', 'No', 'View', 'All'],
        ['Onboarding', 'No', 'No', 'Yes', 'No', 'No', 'Yes'],
        ['Calendar / Events', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Audit Log', 'No', 'No', 'No', 'No', 'No', 'Yes'],
        ['Module Management', 'No', 'No', 'No', 'No', 'No', 'Yes'],
        ['Branding / Theme', 'No', 'No', 'No', 'No', 'No', 'Yes'],
        ['Global Announcements', 'No', 'No', 'No', 'No', 'No', 'Yes'],
        ['Data Reset', 'No', 'No', 'No', 'No', 'No', 'Yes'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. HR MANAGER - FULL WORKFLOW
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. HR Manager - Full Workflow', level=1)
doc.add_paragraph(
    'The HR Manager is the primary operational role in the PIXOUS HR Portal, responsible for day-to-day '
    'HR administration, employee lifecycle management, payroll processing, and organizational compliance.'
)

doc.add_heading('5.1 Daily Workflow', level=2)
hr_daily = [
    ['1', 'Login & Authentication', 'Access portal via username/password, receive JWT tokens', '/api/auth/login'],
    ['2', 'Dashboard Review', 'Review personal dashboard, pending approvals, notifications', '/api/dashboard/me'],
    ['3', 'Check Absent Today', 'View employees absent today for follow-up', '/api/attendance/absent-today'],
    ['4', 'Process Leave Queue', 'Review pending leave requests, approve/reject', '/api/leave/pending'],
    ['5', 'Attendance Team View', 'Monitor team attendance for current date', '/api/attendance/team'],
    ['6', 'Payroll Processing', 'Generate payslips, run batch payroll', '/api/payroll/runs'],
    ['7', 'New Employee Onboarding', 'Create employee accounts, assign roles', '/api/auth/employees'],
    ['8', 'Helpdesk Monitoring', 'Review and assign support tickets', '/api/tickets/all'],
    ['9', 'Report Generation', 'Generate attendance and payroll reports', '/api/reports/'],
    ['10', 'Chat & Communication', 'Manage community groups, send announcements', '/api/communities/'],
]
add_styled_table(doc, ['Step', 'Activity', 'Description', 'API Endpoint'], hr_daily)

doc.add_heading('5.2 Employee Lifecycle Management', level=2)
doc.add_paragraph('The HR Manager manages the complete employee lifecycle through the following workflow steps:')

lifecycle = [
    ['Employee Creation', 'HR creates new employee via /api/auth/employees with role, department, designation, salary structure assignment', 'USER_MANAGE'],
    ['Bulk Import', 'HR imports multiple employees via Excel upload through /api/auth/employees/bulk with preview and undo capability', 'USER_MANAGE'],
    ['Profile Management', 'HR updates employee profiles including personal info, bank details, documents, and profile photos', 'USER_MANAGE'],
    ['Role Assignment', 'HR assigns roles (Employee, Manager, HR, etc.) with corresponding permission sets', 'USER_MANAGE'],
    ['Onboarding Tasks', 'HR creates and tracks onboarding checklists for new employees via /api/onboarding/', 'USER_MANAGE'],
    ['Leave Policy Setup', 'HR configures leave types, default balances, and approval workflows via /api/leave/types', 'ORG_MANAGE'],
    ['Salary Configuration', 'HR sets salary structure (basic, HRA, deductions) via /api/payroll/salary', 'PAYROLL_RUN'],
    ['Offboarding', 'HR processes employee exit including asset return, final settlement, and account deactivation', 'USER_MANAGE'],
]
add_styled_table(doc, ['Step', 'Workflow Detail', 'Permission Required'], lifecycle)

doc.add_heading('5.3 Payroll Workflow', level=2)
payroll_steps = [
    ['1', 'Set Salary Structure', 'HR defines basic salary, HRA, conveyance, special allowances, deductions per employee', 'PAYROLL_RUN'],
    ['2', 'Record Monthly Salary', 'HR records monthly basic salary adjustments via /api/payroll/salary-months', 'PAYROLL_RUN'],
    ['3', 'Attendance Review', 'HR reviews attendance data including LOP (Loss of Pay) days for the month', 'PAYROLL_VIEW'],
    ['4', 'Leave Balance Review', 'HR checks leave balances and LOP preview via /api/leave/lop-preview', 'PAYROLL_VIEW'],
    ['5', 'Generate Payslips', 'HR triggers payslip generation via /api/payroll/payslip/generate for selected employees', 'PAYROLL_RUN'],
    ['6', 'Start Payroll Run', 'HR initiates batch payroll via /api/payroll/runs with employee selection', 'PAYROLL_RUN'],
    ['7', 'Confirm Payroll Run', 'HR confirms the payroll run via /api/payroll/runs/{id}/confirm', 'PAYROLL_RUN'],
    ['8', 'Finance Approval', 'Finance officer approves via /api/payroll/runs/{id}/finance-approve', 'PAYROLL_APPROVE'],
    ['9', 'PDF Generation', 'System generates payslip PDFs using OpenPDF for download/print', 'Automated'],
    ['10', 'Employee Access', 'Employees download payslips via /api/payroll/payslip/{id}/pdf', 'PAYROLL_VIEW'],
]
add_styled_table(doc, ['Step', 'Action', 'Description', 'Permission'], payroll_steps)

doc.add_heading('5.4 Leave Management Workflow', level=2)
leave_steps = [
    ['1', 'Configure Leave Types', 'HR creates leave types (Casual, Sick, Earned, etc.) with annual quotas via /api/leave/types', 'ORG_MANAGE'],
    ['2', 'Allocate Defaults', 'HR applies default leave balances via /api/leave/allocations/apply-defaults', 'ORG_MANAGE'],
    ['3', 'Employee Applies', 'Employee submits leave request with dates and reason via /api/leave/apply', 'LEAVE_APPLY'],
    ['4', 'Approval Notification', 'System sends WebSocket notification to approver via /topic/notifications/{userId}', 'Automated'],
    ['5', 'Manager Reviews', 'Team lead/HR reviews request in approval queue via /api/leave/my-queue', 'LEAVE_APPROVE'],
    ['6', 'Decision', 'Approver approves or rejects via /api/leave/{id}/decision with comments', 'LEAVE_APPROVE'],
    ['7', 'Bulk Operations', 'HR processes multiple requests via /api/leave/bulk-decision', 'LEAVE_APPROVE'],
    ['8', 'Calendar Update', 'Approved leave reflected in team calendar via /api/leave/calendar', 'Automated'],
    ['9', 'Balance Update', 'Leave balance auto-deducted, LOP calculated if applicable', 'Automated'],
    ['10', 'Cancellation', 'Employee can cancel pending/approved leave via /api/leave/{id}/cancel', 'LEAVE_APPLY'],
]
add_styled_table(doc, ['Step', 'Action', 'Description', 'Permission'], leave_steps)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. CTO / EXECUTIVE - FULL WORKFLOW
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. CTO / Executive - Full Workflow', level=1)
doc.add_paragraph(
    'The CTO/Executive role (mapped to IT_CEO) has a read-heavy workflow focused on strategic oversight, '
    'executive dashboards, organizational KPIs, and policy decisions. The role has access to executive-level '
    'dashboards and organization-wide reports without operational transaction capabilities.'
)

doc.add_heading('6.1 Executive Dashboard Workflow', level=2)
cto_workflow = [
    ['1', 'Executive Login', 'Access portal via executive credentials with DASHBOARD_EXEC permission', '/api/auth/login'],
    ['2', 'Executive Dashboard', 'View KPIs: headcount, attrition, attendance %, payroll summary, ticket resolution rates', '/api/dashboard/executive'],
    ['3', 'Org Insights', 'Review organization-wide insights: department distribution, gender ratio, location spread', '/api/dashboard/org-insights'],
    ['4', 'Celebrations', 'View birthdays and work anniversaries for employee engagement', '/api/dashboard/celebrations'],
    ['5', 'Attendance Overview', 'View attendance trends and patterns across the organization', '/api/attendance/team-range'],
    ['6', 'Payroll Summary', 'Review payroll summary by department, month, and status', '/api/payroll/salary-months'],
    ['7', 'Leave Analytics', 'Review leave calendar and absence patterns', '/api/leave/calendar'],
    ['8', 'Work Reports', 'Review work reports submitted across teams', '/api/work-reports/'],
    ['9', 'Report Downloads', 'Download attendance and payroll Excel reports', '/api/reports/attendance/excel'],
    ['10', 'Announcements', 'Review global announcements and company events', '/api/global-announcements/'],
]
add_styled_table(doc, ['Step', 'Activity', 'Description', 'API Endpoint'], cto_workflow)

doc.add_heading('6.2 Strategic Decision Support', level=2)
strategic = [
    ['Headcount Analysis', 'Real-time employee count by department, designation, location, and status'],
    ['Attrition Tracking', 'Monthly and quarterly attrition rates with trend analysis'],
    ['Cost Analysis', 'Total payroll cost by department with month-over-month comparison'],
    ['Attendance Analytics', 'Organization-wide attendance percentage, overtime, and leave utilization'],
    ['Ticket Resolution', 'Helpdesk SLA compliance, average resolution time, satisfaction ratings'],
    ['Performance Metrics', 'Goal completion rates, review scores, and team performance rankings'],
]
add_styled_table(doc, ['Metric', 'Description'], strategic)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. TEAM LEAD / MANAGER - FULL WORKFLOW
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Team Lead / Manager - Full Workflow', level=1)
doc.add_paragraph(
    'The Team Lead/Manager role (IT_MGR) focuses on team-level operations: approving leaves, '
    'monitoring team attendance, assigning tasks, reviewing work reports, and managing team performance.'
)

doc.add_heading('7.1 Team Management Workflow', level=2)
tl_workflow = [
    ['1', 'Login & Dashboard', 'View personal dashboard with team summary', '/api/dashboard/me'],
    ['2', 'Team Attendance', 'Check today\'s team attendance via /api/attendance/my-team-today', '/api/attendance/my-team-today'],
    ['3', 'Leave Approvals', 'Review and approve/reject team leave requests via /api/leave/my-queue', '/api/leave/my-queue'],
    ['4', 'Task Assignment', 'Create and assign tasks to team members via /api/tasks/', '/api/tasks/'],
    ['5', 'Work Report Review', 'Review and provide feedback on team work reports', '/api/work-reports/'],
    ['6', 'Team Chat', 'Communicate with team via team chat room', '/api/communities/team'],
    ['7', 'Performance Reviews', 'Conduct team member appraisals and goal setting', '/api/performance/'],
    ['8', 'Onboarding Support', 'Guide new team members through onboarding checklist', '/api/onboarding/'],
    ['9', 'Helpdesk Escalation', 'Escalate unresolved tickets to admin', '/api/tickets/'],
    ['10', 'Weekly Reports', 'Generate and review team weekly attendance and performance reports', '/api/reports/'],
]
add_styled_table(doc, ['Step', 'Activity', 'Description', 'API Endpoint'], tl_workflow)

doc.add_heading('7.2 Team Approval Matrix', level=2)
approval_matrix = [
    ['Leave Request', 'Own team members only', '/api/leave/{id}/decision', 'LEAVE_APPROVE'],
    ['Overtime Request', 'Own team members only', 'Via attendance module', 'ATTENDANCE_TEAM'],
    ['Work Report', 'Own team submissions', '/api/work-reports/', 'REPORT_VIEW'],
    ['Task Completion', 'Assigned tasks only', '/api/tasks/', 'ATTENDANCE_SELF'],
    ['Expense Claim', 'Own team members', '/api/ta-expenses/', 'REPORT_VIEW'],
]
add_styled_table(doc, ['Request Type', 'Scope', 'Endpoint', 'Permission'], approval_matrix)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. SYSTEM ADMIN / COMPANY ADMIN - FULL WORKFLOW
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. System Admin / Company Admin - Full Workflow', level=1)
doc.add_paragraph(
    'The System Admin / Company Admin (SUPER_ADMIN / COMPANY_ADMIN) has unrestricted access to all '
    'modules, including system configuration, audit logging, data management, branding, module toggling, '
    'and technical administration.'
)

doc.add_heading('8.1 System Administration Workflow', level=2)
admin_workflow = [
    ['1', 'System Login', 'Full admin access with all permission codes', '/api/auth/login'],
    ['2', 'Dashboard Overview', 'View executive dashboard with all KPIs', '/api/dashboard/executive'],
    ['3', 'Organization Setup', 'Configure company, departments, designations, locations via /api/org/', '/api/org/'],
    ['4', 'Module Management', 'Enable/disable application modules per company via /api/my-modules/', '/api/my-modules/'],
    ['5', 'Role Configuration', 'Create custom roles with specific permission combinations', '/api/auth/'],
    ['6', 'Employee Management', 'Full CRUD on employees including bulk import/export', '/api/auth/employees'],
    ['7', 'Leave Policy Admin', 'Configure leave types, quotas, approval chains, carry-forward rules', '/api/leave/types'],
    ['8', 'Payroll Admin', 'Run payroll, generate payslips, manage salary structures', '/api/payroll/'],
    ['9', 'Asset Administration', 'Manage IT assets, allocate, return, track via QR codes', '/api/assets/'],
    ['10', 'Helpdesk Administration', 'Monitor all tickets, assign agents, configure SLAs', '/api/tickets/all'],
    ['11', 'Chat Management', 'Create/manage community groups, set retention policies', '/api/communities/'],
    ['12', 'Branding & Theme', 'Customize application branding, colors, logo', '/api/settings/'],
    ['13', 'Audit Log Review', 'Review all system audit logs and technical audit logs', '/api/audit/'],
    ['14', 'Data Reset', 'Reset employee data, re-run migrations, cache management', '/api/cache/'],
    ['15', 'Announcements', 'Create global login announcements with entrance effects', '/api/global-announcements/'],
    ['16', 'Calendar Management', 'Create company events and holidays', '/api/calendar/'],
    ['17', 'System Settings', 'Configure system-wide settings via /api/settings/', '/api/settings/'],
    ['18', 'Cache Management', 'Clear Redis cache, manage cache configurations', '/api/cache/'],
]
add_styled_table(doc, ['Step', 'Activity', 'Description', 'API Endpoint'], admin_workflow)

doc.add_heading('8.2 Technical Admin (SaaS)', level=2)
doc.add_paragraph(
    'The Technical Admin operates through a separate login endpoint (/api/technical-admin/) and manages '
    'multi-tenant SaaS operations including company creation, module assignment, usage tracking, and '
    'platform-level configurations.'
)

saas_workflow = [
    ['Company Onboarding', 'Create new company tenant with schema and seed data'],
    ['Module Assignment', 'Enable specific modules per company (e.g., disable payroll for free tier)'],
    ['Usage Tracking', 'Monitor API calls, storage usage, active users per company'],
    ['Role Templates', 'Deploy industry-specific role templates (IT/Civil)'],
    ['License Management', 'Manage subscription tiers and feature gates'],
    ['Data Migration', 'Import/export company data across tenants'],
]
add_styled_table(doc, ['Function', 'Description'], saas_workflow)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. EMPLOYEE - FULL WORKFLOW
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Employee - Full Workflow', level=1)
doc.add_paragraph(
    'The Employee role (IT_EMP / CV_EMP) is the most common role in the system, providing self-service '
    'capabilities for attendance, leave, payslips, helpdesk, chat, and work reports.'
)

doc.add_heading('9.1 Employee Daily Workflow', level=2)
emp_workflow = [
    ['1', 'Login', 'Access portal via username/password or face recognition', '/api/auth/login'],
    ['2', 'Dashboard', 'View personal dashboard with attendance summary, pending tasks, notifications', '/api/dashboard/me'],
    ['3', 'Punch In', 'Mark attendance via GPS punch or face recognition', '/api/attendance/punch-in'],
    ['4', 'Check Attendance', 'View today\'s attendance status and monthly summary', '/api/attendance/today'],
    ['5', 'Apply Leave', 'Submit leave request with dates and reason', '/api/leave/apply'],
    ['6', 'View Leave Balance', 'Check remaining leave balances by type', '/api/leave/balances'],
    ['7', 'View Payslip', 'Download monthly payslip as PDF', '/api/payroll/payslip/{id}/pdf'],
    ['8', 'Raise Ticket', 'Create helpdesk ticket for IT/HR issues', '/api/tickets/'],
    ['9', 'Chat', 'Send messages to team or community groups', '/api/communities/{id}/messages'],
    ['10', 'Work Report', 'Submit daily/weekly work report', '/api/work-reports/'],
    ['11', 'View Tasks', 'Check assigned tasks and update status', '/api/tasks/'],
    ['12', 'View Assets', 'Check allocated assets and acknowledge receipt', '/api/assets/my-assets'],
    ['13', 'Punch Out', 'Mark attendance out at end of day', '/api/attendance/punch-out'],
]
add_styled_table(doc, ['Step', 'Activity', 'Description', 'API Endpoint'], emp_workflow)

doc.add_heading('9.2 Employee Self-Service Features', level=2)
self_service = [
    ['Profile Management', 'Update personal info, profile photo, bank details'],
    ['Password Change', 'Change password via /api/auth/change-password'],
    ['Leave Application', 'Apply, view, and cancel leave requests'],
    ['Payslip Download', 'View and download monthly payslips as PDF'],
    ['Attendance History', 'View personal attendance calendar and summaries'],
    ['Asset Acknowledgment', 'Acknowledge receipt of allocated IT assets'],
    ['Helpdesk Tickets', 'Raise and track support tickets with SLA'],
    ['Work Reports', 'Submit daily/weekly work reports'],
    ['Chat & Messaging', 'Send messages, voice notes, file attachments'],
    ['Notifications', 'Receive real-time WebSocket push notifications'],
    ['Calendar Events', 'View company events and holidays'],
    ['TA Expenses', 'Submit travel/allowance expense claims'],
]
add_styled_table(doc, ['Feature', 'Description'], self_service)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. MODULE-WISE WORKFLOW DETAILS
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Module-wise Workflow Details', level=1)

# 10.1 Auth
doc.add_heading('10.1 Authentication & Login Flow', level=2)
doc.add_paragraph('The authentication system implements JWT-based stateless authentication with token rotation, '
    'account lockout, and refresh token management.')
auth_flow = [
    ['Login Request', 'POST /api/auth/login', 'Public', 'Username + password submitted'],
    ['Rate Limiting', 'Server-side', '-', '5 failed attempts triggers account lockout'],
    ['JWT Generation', 'Server', '-', 'Access token (4hr TTL) + refresh token generated'],
    ['Token Storage', 'Client (localStorage)', '-', 'Access token stored for API authorization'],
    ['Auto Refresh', 'Axios interceptor', '-', 'Automatic token refresh on 401 response'],
    ['Logout', 'POST /api/auth/logout', 'Authenticated', 'All refresh tokens revoked'],
    ['Password Change', 'POST /api/auth/change-password', 'Authenticated', 'BCrypt hash verification'],
    ['Phone Validation', 'POST /api/auth/validate-phone', 'Public', 'Check phone number uniqueness'],
    ['Username Check', 'GET /api/auth/check-username', 'Public', 'Check username availability'],
]
add_styled_table(doc, ['Step', 'Endpoint', 'Auth Required', 'Description'], auth_flow)

# 10.2 Attendance
doc.add_heading('10.2 Attendance Management', level=2)
doc.add_paragraph('Supports multiple attendance methods: GPS punch, face recognition, and manual entry.')
att_flow = [
    ['Punch In', 'POST /api/attendance/punch-in', 'Self', 'GPS coordinates recorded, geofence validated'],
    ['Punch Out', 'POST /api/attendance/punch-out', 'Self', 'Duration calculated, overtime flagged'],
    ['Face Punch', 'POST /api/attendance/face-punch', 'Self', 'Face verified against stored embedding'],
    ['Today Status', 'GET /api/attendance/today', 'Self', 'Current day attendance record'],
    ['Calendar View', 'GET /api/attendance/me', 'Self', 'Monthly attendance calendar'],
    ['Monthly Summary', 'GET /api/attendance/me/summary', 'Self', 'Present/absent/overtime summary'],
    ['Insights', 'GET /api/attendance/insights', 'Self', 'Anomaly and pattern detection'],
    ['Absent Today', 'GET /api/attendance/absent-today', 'All', 'List of absent employees'],
    ['Team Today', 'GET /api/attendance/my-team-today', 'Self', 'Team presence today'],
    ['Team Attendance', 'GET /api/attendance/team', 'ATTENDANCE_TEAM', 'Team attendance for date range'],
    ['Team Range', 'GET /api/attendance/team-range', 'ATTENDANCE_TEAM', 'Team attendance over date range'],
]
add_styled_table(doc, ['Action', 'Endpoint', 'Permission', 'Description'], att_flow)

# 10.3 Leave
doc.add_heading('10.3 Leave Management', level=2)
leave_flow = [
    ['List Types', 'GET /api/leave/types', 'Authenticated', 'Available leave types'],
    ['Create Type', 'POST /api/leave/types', 'ORG_MANAGE', 'Define new leave type with quota'],
    ['Balances', 'GET /api/leave/balances', 'Self', 'Current leave balances by type'],
    ['Apply', 'POST /api/leave/apply', 'LEAVE_APPLY', 'Submit leave with dates and reason'],
    ['My Requests', 'GET /api/leave/me', 'Self', 'View my leave request history'],
    ['Approval Queue', 'GET /api/leave/my-queue', 'Self', 'Pending approvals for current user'],
    ['Cancel', 'POST /api/leave/{id}/cancel', 'Self', 'Cancel a pending/approved request'],
    ['Pending All', 'GET /api/leave/pending', 'LEAVE_APPROVE', 'All pending approvals'],
    ['Decision', 'POST /api/leave/{id}/decision', 'LEAVE_APPROVE', 'Approve or reject with comment'],
    ['Bulk Decision', 'POST /api/leave/bulk-decision', 'LEAVE_APPROVE', 'Process multiple requests'],
    ['Calendar', 'GET /api/leave/calendar', 'LEAVE_APPROVE', 'Leave calendar view'],
    ['On Leave Today', 'GET /api/leave/on-leave', 'All', 'Employees on leave today'],
    ['LOP Preview', 'GET /api/leave/lop-preview', 'PAYROLL_RUN', 'LOP calculation for payslip'],
    ['Reset', 'POST /api/leave/reset/{userId}', 'USER_MANAGE', 'Reset user leave balances'],
]
add_styled_table(doc, ['Action', 'Endpoint', 'Permission', 'Description'], leave_flow)

# 10.4 Payroll
doc.add_heading('10.4 Payroll Management', level=2)
payroll_flow = [
    ['Set Salary', 'POST /api/payroll/salary', 'PAYROLL_RUN', 'Define salary structure'],
    ['Get Salary', 'GET /api/payroll/salary/{userId}', 'PAYROLL_VIEW', 'View salary details'],
    ['List Salaries', 'GET /api/payroll/salaries', 'PAYROLL_VIEW', 'All salary records'],
    ['Monthly Salary', 'POST /api/payroll/salary-months', 'PAYROLL_RUN', 'Record monthly basic'],
    ['My Monthly', 'GET /api/payroll/salary-months/me', 'Self', 'My monthly salary data'],
    ['Generate Payslip', 'POST /api/payroll/payslip/generate', 'PAYROLL_RUN', 'Generate individual payslip'],
    ['My Payslips', 'GET /api/payroll/payslip/list', 'Self', 'My payslip history'],
    ['Employee Payslips', 'GET /api/payroll/payslip/list/{userId}', 'PAYROLL_VIEW', 'View employee payslips'],
    ['Get Payslip', 'GET /api/payroll/payslip/{id}', 'Self', 'View specific payslip'],
    ['Download PDF', 'GET /api/payroll/payslip/{id}/pdf', 'Self', 'Download payslip PDF'],
    ['Start Run', 'POST /api/payroll/runs', 'PAYROLL_RUN', 'Start batch payroll run'],
    ['Confirm Run', 'POST /api/payroll/runs/{id}/confirm', 'PAYROLL_RUN', 'Confirm payroll run'],
    ['Finance Approve', 'POST /api/payroll/runs/{id}/finance-approve', 'PAYROLL_APPROVE', 'Finance approval'],
    ['List Runs', 'GET /api/payroll/runs', 'PAYROLL_RUN', 'List all payroll runs'],
    ['Request Payslip', 'POST /api/payroll/requests', 'Self', 'Request payslip generation'],
    ['Approve Request', 'POST /api/payroll/requests/{id}/approve', 'PAYROLL_RUN', 'Approve payslip request'],
]
add_styled_table(doc, ['Action', 'Endpoint', 'Permission', 'Description'], payroll_flow)

doc.add_page_break()

# 10.5 Assets
doc.add_heading('10.5 Asset Management', level=2)
asset_flow = [
    ['List Assets', 'GET /api/assets/', 'ASSET_MANAGE', 'All assets with status'],
    ['My Assets', 'GET /api/assets/my-assets', 'Self', 'My allocated assets'],
    ['Lookup', 'GET /api/assets/lookup', 'Self', 'Find asset by code'],
    ['Create Asset', 'POST /api/assets/', 'ASSET_MANAGE', 'Register new asset'],
    ['QR Code', 'GET /api/assets/{id}/qr', 'Self', 'Generate asset QR code'],
    ['Allocate', 'POST /api/assets/{id}/allocate', 'ASSET_MANAGE', 'Allocate to employee'],
    ['Acknowledge', 'POST /api/assets/{id}/acknowledge', 'Self', 'Acknowledge receipt'],
    ['Return', 'POST /api/assets/{id}/return', 'ASSET_MANAGE', 'Return asset'],
    ['Delete', 'DELETE /api/assets/{id}', 'ASSET_MANAGE', 'Remove asset record'],
]
add_styled_table(doc, ['Action', 'Endpoint', 'Permission', 'Description'], asset_flow)

# 10.6 Helpdesk
doc.add_heading('10.6 Helpdesk / Ticketing', level=2)
helpdesk_flow = [
    ['My Tickets', 'GET /api/tickets/', 'Self', 'View my submitted tickets'],
    ['Raise Ticket', 'POST /api/tickets/', 'HELPDESK_RAISE', 'Create new support ticket'],
    ['Upload File', 'POST /api/tickets/upload', 'Self', 'Attach file to ticket'],
    ['Agents List', 'GET /api/tickets/agents', 'Self', 'Available support agents'],
    ['My Queue', 'GET /api/tickets/assigned-to-me', 'HELPDESK_AGENT', 'Agent ticket queue'],
    ['All Tickets', 'GET /api/tickets/all', 'HELPDESK_AGENT', 'All system tickets'],
    ['Update', 'PUT /api/tickets/{id}', 'Self', 'Update own ticket'],
    ['Get Ticket', 'GET /api/tickets/{id}', 'Self', 'View ticket details'],
    ['Add Comment', 'POST /api/tickets/{id}/comments', 'Self', 'Add ticket comment'],
    ['Change Status', 'POST /api/tickets/{id}/status', 'HELPDESK_AGENT', 'Update ticket status'],
    ['Rate', 'POST /api/tickets/{id}/rating', 'Self', 'Rate resolved ticket'],
]
add_styled_table(doc, ['Action', 'Endpoint', 'Permission', 'Description'], helpdesk_flow)

# 10.7 Community & Chat
doc.add_heading('10.7 Community & Chat', level=2)
chat_flow = [
    ['Create Group', 'POST /api/communities/', 'Admin/HR', 'Create chat group'],
    ['All Groups', 'GET /api/communities/', 'Admin', 'List all groups'],
    ['My Groups', 'GET /api/communities/me', 'Self', 'My chat groups'],
    ['Contacts', 'GET /api/communities/contacts', 'Self', 'Chat contacts list'],
    ['Direct Message', 'POST /api/communities/direct/{userId}', 'Self', 'Open 1:1 DM'],
    ['Team Room', 'POST /api/communities/team', 'Self', 'Open team chat room'],
    ['Members', 'GET /api/communities/{id}/members', 'Member', 'View group members'],
    ['Add Member', 'POST /api/communities/{id}/members', 'Admin', 'Add member to group'],
    ['Messages', 'GET /api/communities/{id}/messages', 'Member', 'Message history'],
    ['Send Message', 'POST /api/communities/{id}/messages', 'Member', 'Send text message'],
    ['Search', 'GET /api/communities/{id}/messages/search', 'Member', 'Search messages'],
    ['Pin Message', 'POST /api/communities/messages/{id}/pin', 'Admin', 'Pin/unpin message'],
    ['Reaction', 'POST /api/communities/messages/{id}/reactions', 'Member', 'Toggle reaction'],
    ['Mark Read', 'POST /api/communities/messages/{id}/read', 'Member', 'Mark as read'],
    ['Read Receipts', 'GET /api/communities/messages/{id}/receipts', 'Admin', 'View read receipts'],
    ['Vote', 'POST /api/communities/messages/{id}/vote', 'Member', 'Vote in poll'],
    ['Voice Message', 'POST /api/communities/{id}/voice', 'Member', 'Send voice message'],
    ['Attachment', 'POST /api/communities/{id}/attachments', 'Member', 'Send file'],
    ['Retention', 'GET/PUT /api/communities/retention', 'Admin', 'Chat retention policy'],
]
add_styled_table(doc, ['Action', 'Endpoint', 'Permission', 'Description'], chat_flow)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. CROSS-ROLE WORKFLOW DIAGRAMS
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. Cross-Role Workflow Diagrams', level=1)

doc.add_heading('11.1 Employee Onboarding Workflow', level=2)
onboarding = [
    ['1', 'HR creates employee account', 'HR Manager', '/api/auth/employees', 'Employee profile created in system'],
    ['2', 'HR assigns role and department', 'HR Manager', '/api/auth/employees', 'Role permissions auto-assigned'],
    ['3', 'HR configures salary structure', 'HR Manager', '/api/payroll/salary', 'Salary components defined'],
    ['4', 'HR sets default leave balances', 'HR Manager', '/api/leave/allocations/apply-defaults', 'Leave quotas allocated'],
    ['5', 'HR creates onboarding checklist', 'HR Manager', '/api/onboarding/', 'Tasks defined for new hire'],
    ['6', 'Employee receives welcome email', 'System', 'Email/SMS', 'Login credentials shared'],
    ['7', 'Employee logs in and updates profile', 'Employee', '/api/auth/me + profile update', 'Personal info completed'],
    ['8', 'Employee acknowledges asset receipt', 'Employee', '/api/assets/{id}/acknowledge', 'Asset allocation confirmed'],
    ['9', 'Team lead assigns initial tasks', 'Team Lead', '/api/tasks/', 'First assignments created'],
    ['10', 'Employee submits work report', 'Employee', '/api/work-reports/', 'First work report submitted'],
    ['11', 'Manager reviews and provides feedback', 'Team Lead', '/api/work-reports/', 'Performance baseline set'],
    ['12', 'HR monitors onboarding completion', 'HR Manager', '/api/onboarding/', 'Checklist progress tracked'],
]
add_styled_table(doc, ['Step', 'Action', 'Actor', 'API Endpoint', 'Result'], onboarding)

doc.add_heading('11.2 Payroll Processing Workflow (End-to-End)', level=2)
payroll_e2e = [
    ['1', 'HR sets salary structure for employees', 'HR Manager', '/api/payroll/salary', 'Salary components defined'],
    ['2', 'HR records monthly basic salary', 'HR Manager', '/api/payroll/salary-months', 'Monthly salary set'],
    ['3', 'System captures attendance data', 'System (auto)', '/api/attendance/*', 'Working days, overtime, LOP'],
    ['4', 'System calculates leave deductions', 'System (auto)', '/api/leave/lop-preview', 'LOP days computed'],
    ['5', 'HR generates individual payslips', 'HR Manager', '/api/payroll/payslip/generate', 'Payslip records created'],
    ['6', 'HR starts batch payroll run', 'HR Manager', '/api/payroll/runs', 'Payroll run initiated'],
    ['7', 'System calculates earnings & deductions', 'System (auto)', 'Payroll engine', 'Net pay computed'],
    ['8', 'HR confirms the payroll run', 'HR Manager', '/api/payroll/runs/{id}/confirm', 'Run locked for review'],
    ['9', 'Finance officer reviews and approves', 'Finance Officer', '/api/payroll/runs/{id}/finance-approve', 'Financial approval'],
    ['10', 'System generates payslip PDFs', 'System (auto)', 'OpenPDF engine', 'PDF payslips created'],
    ['11', 'Employees download their payslips', 'Employee', '/api/payroll/payslip/{id}/pdf', 'Payslip accessible'],
    ['12', 'CEO reviews payroll summary', 'CEO/CTO', '/api/dashboard/executive', 'Executive overview'],
]
add_styled_table(doc, ['Step', 'Action', 'Actor', 'API/Component', 'Result'], payroll_e2e)

doc.add_heading('11.3 Helpdesk Ticket Lifecycle', level=2)
ticket_lifecycle = [
    ['1', 'Employee raises ticket', 'Employee', '/api/tickets/', 'Ticket created with OPEN status'],
    ['2', 'System notifies agents', 'System (WebSocket)', '/topic/notifications', 'Real-time push to agents'],
    ['3', 'Agent picks up ticket', 'Admin/Agent', '/api/tickets/assigned-to-me', 'Ticket in agent queue'],
    ['4', 'Agent adds comments', 'Admin/Agent', '/api/tickets/{id}/comments', 'Communication thread'],
    ['5', 'Agent changes status to IN_PROGRESS', 'Admin/Agent', '/api/tickets/{id}/status', 'Status updated'],
    ['6', 'Agent resolves ticket', 'Admin/Agent', '/api/tickets/{id}/status', 'Status -> RESOLVED'],
    ['7', 'System notifies employee', 'System (WebSocket)', '/topic/notifications', 'Resolution notification'],
    ['8', 'Employee rates the resolution', 'Employee', '/api/tickets/{id}/rating', 'Feedback captured'],
    ['9', 'Ticket archived', 'System', 'Database', 'Ticket history preserved'],
]
add_styled_table(doc, ['Step', 'Action', 'Actor', 'Endpoint/System', 'Result'], ticket_lifecycle)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. DATA FLOW ARCHITECTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. Data Flow Architecture', level=1)

doc.add_heading('12.1 Request Flow', level=2)
doc.add_paragraph(
    'All client requests follow a standardized flow through the application layers:'
)
request_flow = [
    ['Client (Browser/Mobile)', 'HTTP/HTTPS Request + JWT Bearer Token'],
    ['Nginx Reverse Proxy', 'SSL termination, static file serving, load balancing'],
    ['Spring Security Filter Chain', 'JWT validation, RBAC enforcement, rate limiting'],
    ['Controller Layer', 'Request validation, DTO mapping (MapStruct)'],
    ['Service Layer', 'Business logic, transaction management'],
    ['Repository Layer (JPA)', 'Database queries with multi-tenant filtering'],
    ['MySQL 8.4', 'Data persistence with Flyway-managed schema'],
    ['Response DTO', 'MapStruct mapping, JSON serialization'],
    ['Client', 'TanStack Query cache update, UI render'],
]
add_styled_table(doc, ['Layer', 'Description'], request_flow)

doc.add_heading('12.2 Real-time Data Flow', level=2)
realtime_flow = [
    ['WebSocket Connection', 'Client connects via STOMP/SockJS to /ws endpoint'],
    ['Authentication', 'JWT validated during WebSocket handshake'],
    ['Subscription', 'Client subscribes to /topic/notifications/{userId}'],
    ['Event Trigger', 'Server-side event (leave approval, ticket update, etc.)'],
    ['Message Dispatch', 'SimpMessagingTemplate sends to subscribed topic'],
    ['Client Receive', 'STOMP client processes message, updates UI'],
    ['Cache Invalidation', 'TanStack Query invalidates relevant queries'],
]
add_styled_table(doc, ['Step', 'Description'], realtime_flow)

doc.add_heading('12.3 Authentication Data Flow', level=2)
auth_data = [
    ['Login', 'Client POSTs credentials -> Server validates -> Returns JWT pair'],
    ['API Call', 'Client adds Authorization: Bearer <token> header'],
    ['Token Validation', 'Server extracts and validates JWT signature and expiry'],
    ['Permission Check', '@PreAuthorize checks required permission codes'],
    ['Token Refresh', 'On 401, client sends refresh token -> Gets new pair'],
    ['Logout', 'Client calls logout -> All refresh tokens revoked'],
]
add_styled_table(doc, ['Stage', 'Description'], auth_data)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. API WORKFLOW SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. API Workflow Summary', level=1)

doc.add_heading('13.1 API Endpoint Count by Module', level=2)
api_count = [
    ['Authentication', '/api/auth/*', '10+', 'Login, signup, token refresh, password change'],
    ['Attendance', '/api/attendance/*', '12+', 'Punch in/out, face auth, team views, insights'],
    ['Leave', '/api/leave/*', '14+', 'Types, balances, apply, approve, calendar, LOP'],
    ['Payroll', '/api/payroll/*', '16+', 'Salary, payslips, runs, requests, PDF generation'],
    ['Assets', '/api/assets/*', '9+', 'CRUD, allocation, QR codes, acknowledge, return'],
    ['Helpdesk', '/api/tickets/*', '11+', 'Tickets, comments, status, ratings, uploads'],
    ['Community', '/api/communities/*', '19+', 'Groups, messages, reactions, polls, retention'],
    ['Dashboard', '/api/dashboard/*', '4+', 'Personal, executive, celebrations, org insights'],
    ['Notifications', '/api/notifications/*', '4+', 'WebSocket push, list, read, mark-all'],
    ['Tasks', '/api/tasks/*', '5+', 'CRUD, assign, workload, task chat'],
    ['Onboarding', '/api/onboarding/*', '3+', 'Checklist, tasks, progress'],
    ['Performance', '/api/performance/*', '3+', 'Goals, reviews'],
    ['TA Expenses', '/api/ta-expenses/*', '5+', 'Submit, approve, reject, categories'],
    ['Complaints', '/api/complaints/*', '3+', 'Raise, review, respond'],
    ['Safety', '/api/safety/*', '3+', 'Report, resolve, list'],
    ['Work Reports', '/api/work-reports/*', '5+', 'CRUD, team views, reminders'],
    ['Calendar', '/api/calendar/*', '2+', 'Events CRUD'],
    ['Chatbot', '/api/chatbot/*', '3+', 'Knowledge, org context, query'],
    ['Announcements', '/api/global-announcements/*', '3+', 'Login announcements with effects'],
    ['Audit', '/api/audit/*', '3+', 'Audit log viewing and filtering'],
    ['Org', '/api/org/*', '10+', 'Departments, designations, locations, holidays'],
    ['Settings', '/api/settings/*', '3+', 'System settings'],
    ['Cache', '/api/cache/*', '2+', 'Cache management'],
    ['Reports', '/api/reports/*', '3+', 'Attendance reports, Excel export'],
    ['Tech Admin', '/api/technical-admin/*', '10+', 'SaaS management'],
]
add_styled_table(doc, ['Module', 'Base Path', 'Endpoints', 'Key Functions'], api_count)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. REVISION HISTORY
# ══════════════════════════════════════════════════════════════
doc.add_heading('14. Document Revision History', level=1)
revisions = [
    ['1.0', datetime.date.today().strftime('%Y-%m-%d'), 'PIXOUS Technologies QA Team', 'Initial comprehensive workflow document'],
]
add_styled_table(doc, ['Version', 'Date', 'Author', 'Description'], revisions)

# ── Footer ──
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PIXOUS Technologies Pvt. Ltd. | Confidential')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save('PIXOUS_HR_Portal_Full_Workflow_Document.docx')
print("Document 1 generated successfully: PIXOUS_HR_Portal_Full_Workflow_Document.docx")
