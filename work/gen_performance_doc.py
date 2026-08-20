"""Generate PIXOUS HR Portal - Performance Testing Document"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

def add_logo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('pixous_logo.png', width=Inches(2.0))

def add_styled_table(doc, headers, rows_data):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): '003366', qn('w:val'): 'clear'})
        shading.append(shd)
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ═══ COVER PAGE ═══
add_logo(doc)
doc.add_paragraph()
title = doc.add_heading('PIXOUS HR Portal', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_heading('Performance Testing Document', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
t = doc.add_table(rows=5, cols=2)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ('Document Type', 'Performance Testing Document'),
    ('Version', '1.0'),
    ('Date', datetime.date.today().strftime('%B %d, %Y')),
    ('Prepared By', 'PIXOUS Technologies QA Team'),
    ('Classification', 'Confidential - Client Deliverable'),
]
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
doc.add_page_break()

# ═══ TOC ═══
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Executive Summary',
    '2. Test Objectives',
    '3. Test Environment',
    '4. Performance Test Types',
    '5. Test Scenarios - Authentication Module',
    '6. Test Scenarios - Attendance Module',
    '7. Test Scenarios - Leave Management Module',
    '8. Test Scenarios - Payroll Module',
    '9. Test Scenarios - Helpdesk Module',
    '10. Test Scenarios - Chat & Community Module',
    '11. Test Scenarios - Dashboard & Reports',
    '12. Test Scenarios - File Operations',
    '13. Performance Metrics & KPIs',
    '14. Role-wise Performance Impact Analysis',
    '15. Performance Test Results Template',
    '16. Bottleneck Analysis Framework',
    '17. Capacity Planning Guidelines',
    '18. Recommendations',
    '19. Appendix: Sample JMeter Configuration',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)
doc.add_page_break()

# ═══ 1. EXECUTIVE SUMMARY ═══
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document defines the performance testing strategy, test scenarios, metrics, and acceptance '
    'criteria for the PIXOUS HR Portal. Performance testing ensures the application can handle '
    'expected user loads while maintaining acceptable response times and system stability.'
)
doc.add_paragraph(
    'The PIXOUS HR Portal is a multi-platform enterprise application serving HR operations across '
    'IT and Civil/Construction industries. The system must support concurrent users across web, '
    'mobile (React Native + Flutter), and API clients simultaneously without degradation.'
)
doc.add_paragraph(
    'Performance testing covers all critical modules including authentication, attendance, leave '
    'management, payroll processing, helpdesk, chat, dashboards, and file operations. Each module '
    'has specific performance requirements based on expected usage patterns and business criticality.'
)
doc.add_paragraph(
    'The testing approach includes load testing, stress testing, spike testing, endurance testing, '
    'and scalability testing to validate the system under various conditions. Results are measured '
    'against defined KPIs and thresholds to determine release readiness.'
)

# ═══ 2. TEST OBJECTIVES ═══
doc.add_heading('2. Test Objectives', level=1)
doc.add_paragraph(
    'The primary objectives of performance testing for the PIXOUS HR Portal are to validate system '
    'behavior under expected and extreme conditions, identify bottlenecks, and ensure the application '
    'meets its performance requirements before production deployment.'
)
objectives = [
    ['OBJ-001', 'Validate response times meet SLA requirements', 'All critical API endpoints must respond within defined thresholds'],
    ['OBJ-002', 'Determine maximum concurrent user capacity', 'System must support minimum 200 concurrent users without degradation'],
    ['OBJ-003', 'Identify performance bottlenecks', 'Pinpoint slow queries, memory leaks, connection pool exhaustion'],
    ['OBJ-004', 'Validate database performance', 'MySQL queries must execute within acceptable time under load'],
    ['OBJ-005', 'Test WebSocket scalability', 'Real-time notifications must reach all connected clients reliably'],
    ['OBJ-006', 'Validate file upload/download performance', 'Payslip PDFs and attachments must transfer within thresholds'],
    ['OBJ-007', 'Test system recovery after peak load', 'System must stabilize within 60 seconds after load removal'],
    ['OBJ-008', 'Validate caching effectiveness', 'Redis cache must reduce database load by at least 60%'],
    ['OBJ-009', 'Test batch operations performance', 'Payroll runs and bulk imports must complete within acceptable time'],
    ['OBJ-010', 'Establish performance baseline', 'Document baseline metrics for future comparison'],
]
add_styled_table(doc, ['ID', 'Objective', 'Description'], objectives)

# ═══ 3. TEST ENVIRONMENT ═══
doc.add_heading('3. Test Environment', level=1)
doc.add_paragraph(
    'The performance test environment is designed to replicate production conditions as closely as '
    'possible while being isolated from development and staging environments to avoid interference.'
)
env = [
    ['Application Server', '4 vCPU, 8GB RAM, Ubuntu 22.04', 'Spring Boot 3.5 on Java 17'],
    ['Database Server', '4 vCPU, 16GB RAM, MySQL 8.4', 'HikariCP connection pool (max 50)'],
    ['Cache Server', '2 vCPU, 4GB RAM, Redis 7', 'Lettuce client, optional fallback'],
    ['Message Queue', '2 vCPU, 4GB RAM, Kafka 3.9', 'KRaft mode, 3 partitions'],
    ['Analytics Service', '2 vCPU, 4GB RAM, Python 3.11', 'FastAPI + face_recognition'],
    ['Load Generator', '8 vCPU, 16GB RAM', 'Apache JMeter 5.6 / Gatling'],
    ['Network', '1 Gbps internal, 100 Mbps external', 'Simulated production latency'],
    ['Monitoring', 'Spring Boot Actuator + Micrometer', 'Prometheus + Grafana dashboards'],
]
add_styled_table(doc, ['Component', 'Specification', 'Details'], env)

doc.add_heading('Test Data Volume', level=2)
data_vol = [
    ['Employees', '500', 'Across 5 departments and 8 designations'],
    ['Attendance Records', '15,000', '30 days x 500 employees'],
    ['Leave Records', '2,000', 'Various statuses (pending, approved, rejected)'],
    ['Payslips', '1,000', '2 months x 500 employees'],
    ['Helpdesk Tickets', '500', 'Various statuses and priorities'],
    ['Chat Messages', '10,000', 'Across 20 community groups'],
    ['Assets', '200', 'Various statuses (available, allocated, returned)'],
    ['Tasks', '300', 'Various statuses and assignees'],
]
add_styled_table(doc, ['Data Type', 'Volume', 'Notes'], data_vol)

doc.add_page_break()

# ═══ 4. PERFORMANCE TEST TYPES ═══
doc.add_heading('4. Performance Test Types', level=1)

doc.add_heading('4.1 Load Testing', level=2)
doc.add_paragraph(
    'Load testing simulates expected production traffic to validate that the system meets performance '
    'requirements under normal conditions. This is the primary performance test conducted before each release.'
)
doc.add_paragraph(
    'The load profile gradually increases from 0 to 200 concurrent users over 10 minutes, maintains '
    'the peak load for 30 minutes, then gradually decreases over 10 minutes. This mimics a typical '
    'workday pattern where users login in the morning, perform operations throughout the day, and '
    'logout in the evening.'
)
load_config = [
    ['Ramp-Up Duration', '10 minutes', 'Gradual user increase from 0 to 200'],
    ['Peak Load Duration', '30 minutes', 'Sustained load at maximum concurrent users'],
    ['Ramp-Down Duration', '10 minutes', 'Gradual user decrease from 200 to 0'],
    ['Total Test Duration', '50 minutes', 'Complete load test cycle'],
    ['Target Concurrent Users', '200', 'Maximum simultaneous users'],
    ['Requests per Second Target', '500', 'Aggregate throughput requirement'],
    ['Error Rate Threshold', '<1%', 'Maximum acceptable error percentage'],
]
add_styled_table(doc, ['Parameter', 'Value', 'Description'], load_config)

doc.add_heading('4.2 Stress Testing', level=2)
doc.add_paragraph(
    'Stress testing pushes the system beyond its normal capacity to identify breaking points and '
    'observe system behavior under extreme conditions. This helps determine the maximum capacity '
    'and understand how the system degrades gracefully.'
)
doc.add_paragraph(
    'The stress test starts with normal load and progressively increases to 500 concurrent users '
    '(2.5x expected capacity). The system should either handle the load or fail gracefully with '
    'appropriate error messages rather than crashing.'
)
stress_config = [
    ['Starting Load', '50 users', 'Normal baseline load'],
    ['Peak Stress Load', '500 users', '2.5x expected capacity'],
    ['Ramp-Up per Stage', '5 minutes', 'Add 50 users every 5 minutes'],
    ['Hold at Peak', '15 minutes', 'Sustain maximum stress load'],
    ['Recovery Test', '10 minutes', 'Monitor system recovery after load removal'],
    ['Breaking Point Target', 'Identify', 'Find the maximum sustainable capacity'],
]
add_styled_table(doc, ['Parameter', 'Value', 'Description'], stress_config)

doc.add_heading('4.3 Spike Testing', level=2)
doc.add_paragraph(
    'Spike testing simulates sudden, sharp increases in traffic to validate system resilience. '
    'This is critical for scenarios like Monday morning login rush or payroll day when many '
    'employees access the system simultaneously.'
)
doc.add_paragraph(
    'The test rapidly increases from 20 to 300 users within 1 minute, holds for 5 minutes, '
    'then drops back to 20 users. The system should absorb the spike without crashing and '
    'recover quickly once the spike subsides.'
)
spike_config = [
    ['Base Load', '20 users', 'Normal background traffic'],
    ['Spike Load', '300 users', '15x base load'],
    ['Spike Duration', '1 minute', 'Rapid increase from base to spike'],
    ['Hold at Spike', '5 minutes', 'Sustain spike load'],
    ['Recovery Time Target', '<60 seconds', 'System must stabilize within 60 seconds'],
]
add_styled_table(doc, ['Parameter', 'Value', 'Description'], spike_config)

doc.add_heading('4.4 Endurance (Soak) Testing', level=2)
doc.add_paragraph(
    'Endurance testing runs the system under sustained load for an extended period to detect '
    'memory leaks, connection pool exhaustion, disk space issues, and other problems that '
    'only manifest over time.'
)
doc.add_paragraph(
    'The test runs at 100 concurrent users for 8 hours (a full workday) with realistic '
    'request patterns. Memory usage, CPU utilization, database connections, and disk I/O '
    'are monitored throughout for any upward trends indicating resource leaks.'
)
endurance_config = [
    ['Concurrent Users', '100', 'Moderate sustained load'],
    ['Duration', '8 hours', 'Full workday simulation'],
    ['Request Pattern', 'Realistic mix', 'Based on production usage analytics'],
    ['Memory Monitoring', 'Every 5 minutes', 'Detect memory leaks or growth trends'],
    ['Connection Pool Monitoring', 'Every 1 minute', 'Detect connection leaks'],
    ['Disk Space Monitoring', 'Every 15 minutes', 'Detect log/file accumulation'],
]
add_styled_table(doc, ['Parameter', 'Value', 'Description'], endurance_config)

doc.add_heading('4.5 Scalability Testing', level=2)
doc.add_paragraph(
    'Scalability testing validates whether the system can handle increasing load by adding resources '
    '(horizontal scaling) or upgrading resources (vertical scaling). This is critical for planning '
    'production infrastructure and future capacity.'
)
doc.add_paragraph(
    'The test measures performance at different resource levels: single server, 2-server cluster, '
    'and 4-server cluster. Results help determine the optimal scaling strategy and cost-effectiveness '
    'of adding more servers versus upgrading existing ones.'
)
scale_config = [
    ['Single Server', '4 vCPU, 8GB RAM', 'Baseline performance measurement'],
    ['2-Server Cluster', '2x (4 vCPU, 8GB RAM)', 'Horizontal scaling validation'],
    ['4-Server Cluster', '4x (4 vCPU, 8GB RAM)', 'Maximum horizontal scaling'],
    ['Vertical Scale', '8 vCPU, 16GB RAM', 'Vertical scaling comparison'],
    ['Load per Test', '200 users', 'Standard load for comparison'],
    ['Measurement', 'Response time, throughput', 'Compare metrics across configurations'],
]
add_styled_table(doc, ['Configuration', 'Resources', 'Purpose'], scale_config)

doc.add_page_break()

# ═══ 5-12. TEST SCENARIOS BY MODULE ═══
doc.add_heading('5. Test Scenarios - Authentication Module', level=1)
doc.add_paragraph(
    'Authentication is the most critical module for performance as every user action begins with '
    'a valid JWT token. Login performance directly impacts user experience and perceived system speed.'
)
auth_perf = [
    ['PERF-AUTH-001', 'Login API', '200 concurrent logins', '<2 seconds', '99%', 'POST /api/auth/login'],
    ['PERF-AUTH-002', 'Token Refresh', '200 concurrent refreshes', '<1 second', '99%', 'POST /api/auth/refresh'],
    ['PERF-AUTH-003', 'Get Current User', '200 concurrent requests', '<500ms', '99%', 'GET /api/auth/me'],
    ['PERF-AUTH-004', 'Logout', '200 concurrent logouts', '<1 second', '99%', 'POST /api/auth/logout'],
    ['PERF-AUTH-005', 'Password Change', '50 concurrent changes', '<2 seconds', '99%', 'POST /api/auth/change-password'],
    ['PERF-AUTH-006', 'Bulk Login Burst', '300 users in 1 minute', '<3 seconds', '95%', 'Spike scenario'],
    ['PERF-AUTH-007', 'JWT Validation', '500 token validations/sec', '<50ms', '99%', 'Middleware validation'],
    ['PERF-AUTH-008', 'Rate Limiting', '10 rapid failed attempts', 'Lockout triggered', '100%', 'Account lockout'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], auth_perf)

doc.add_heading('6. Test Scenarios - Attendance Module', level=1)
doc.add_paragraph(
    'Attendance operations are time-sensitive with peak loads during morning punch-in (9:00-9:30 AM) '
    'and evening punch-out (5:30-6:00 PM). GPS and face recognition add processing overhead.'
)
att_perf = [
    ['PERF-ATT-001', 'GPS Punch In', '200 concurrent punches', '<1 second', '99%', 'POST /api/attendance/punch-in'],
    ['PERF-ATT-002', 'GPS Punch Out', '200 concurrent punches', '<1 second', '99%', 'POST /api/attendance/punch-out'],
    ['PERF-ATT-003', 'Face Recognition Punch', '100 concurrent face punches', '<3 seconds', '95%', 'POST /api/attendance/face-punch'],
    ['PERF-ATT-004', 'Today Attendance', '200 concurrent requests', '<500ms', '99%', 'GET /api/attendance/today'],
    ['PERF-ATT-005', 'Monthly Calendar', '100 concurrent requests', '<1 second', '99%', 'GET /api/attendance/me'],
    ['PERF-ATT-006', 'Team Attendance', '50 concurrent requests', '<2 seconds', '99%', 'GET /api/attendance/team'],
    ['PERF-ATT-007', 'Attendance Insights', '100 concurrent requests', '<2 seconds', '99%', 'GET /api/attendance/insights'],
    ['PERF-ATT-008', 'Absent Today List', '200 concurrent requests', '<1 second', '99%', 'GET /api/attendance/absent-today'],
    ['PERF-ATT-009', 'Morning Rush Simulation', '300 users in 5 minutes', '<2 seconds', '95%', 'Combined punch-in scenarios'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], att_perf)

doc.add_heading('7. Test Scenarios - Leave Management Module', level=1)
doc.add_paragraph(
    'Leave operations involve balance calculations and approval workflows that require database '
    'transactions. Performance testing ensures leave operations complete within acceptable time.'
)
leave_perf = [
    ['PERF-LEA-001', 'Apply for Leave', '50 concurrent applications', '<2 seconds', '99%', 'POST /api/leave/apply'],
    ['PERF-LEA-002', 'View Leave Balances', '200 concurrent requests', '<500ms', '99%', 'GET /api/leave/balances'],
    ['PERF-LEA-003', 'Approve Leave', '50 concurrent approvals', '<2 seconds', '99%', 'POST /api/leave/{id}/decision'],
    ['PERF-LEA-004', 'Bulk Leave Decision', '10 bulk operations', '<10 seconds', '99%', 'POST /api/leave/bulk-decision'],
    ['PERF-LEA-005', 'Leave Calendar', '100 concurrent requests', '<2 seconds', '99%', 'GET /api/leave/calendar'],
    ['PERF-LEA-006', 'My Leave Requests', '200 concurrent requests', '<1 second', '99%', 'GET /api/leave/me'],
    ['PERF-LEA-007', 'LOP Preview', '50 concurrent requests', '<3 seconds', '99%', 'GET /api/leave/lop-preview'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], leave_perf)

doc.add_heading('8. Test Scenarios - Payroll Module', level=1)
doc.add_paragraph(
    'Payroll operations are batch-intensive with PDF generation being the most resource-heavy operation. '
    'Testing focuses on batch processing performance and PDF generation throughput.'
)
payroll_perf = [
    ['PERF-PAY-001', 'Generate Payslip', '10 concurrent generations', '<5 seconds', '95%', 'POST /api/payroll/payslip/generate'],
    ['PERF-PAY-002', 'Download Payslip PDF', '50 concurrent downloads', '<3 seconds', '99%', 'GET /api/payroll/payslip/{id}/pdf'],
    ['PERF-PAY-003', 'Start Payroll Run', '5 concurrent runs', '<30 seconds', '95%', 'POST /api/payroll/runs'],
    ['PERF-PAY-004', 'Confirm Payroll Run', '5 concurrent confirms', '<10 seconds', '95%', 'POST /api/payroll/runs/{id}/confirm'],
    ['PERF-PAY-005', 'View Payslips List', '200 concurrent requests', '<2 seconds', '99%', 'GET /api/payroll/payslip/list'],
    ['PERF-PAY-006', 'Salary Structure', '50 concurrent requests', '<1 second', '99%', 'GET /api/payroll/salary/{userId}'],
    ['PERF-PAY-007', 'Batch Payslip Generation', '100 employees in batch', '<5 minutes', '95%', 'Full payroll run scenario'],
    ['PERF-PAY-008', 'Salary Months View', '100 concurrent requests', '<2 seconds', '99%', 'GET /api/payroll/salary-months'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], payroll_perf)

doc.add_heading('9. Test Scenarios - Helpdesk Module', level=1)
doc.add_paragraph(
    'Helpdesk operations involve ticket CRUD with file uploads and status updates. Performance '
    'testing ensures ticket operations complete quickly even under high volume.'
)
hd_perf = [
    ['PERF-HD-001', 'Raise Ticket', '100 concurrent tickets', '<2 seconds', '99%', 'POST /api/tickets/'],
    ['PERF-HD-002', 'Upload Attachment', '50 concurrent uploads', '<5 seconds', '95%', 'POST /api/tickets/upload'],
    ['PERF-HD-003', 'View My Tickets', '200 concurrent requests', '<1 second', '99%', 'GET /api/tickets/'],
    ['PERF-HD-004', 'Add Comment', '100 concurrent comments', '<2 seconds', '99%', 'POST /api/tickets/{id}/comments'],
    ['PERF-HD-005', 'Change Status', '50 concurrent updates', '<1 second', '99%', 'POST /api/tickets/{id}/status'],
    ['PERF-HD-006', 'View All Tickets', '20 concurrent requests', '<3 seconds', '99%', 'GET /api/tickets/all'],
    ['PERF-HD-007', 'Rate Ticket', '50 concurrent ratings', '<1 second', '99%', 'POST /api/tickets/{id}/rating'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], hd_perf)

doc.add_heading('10. Test Scenarios - Chat & Community Module', level=1)
doc.add_paragraph(
    'Chat is the highest-concurrency module with real-time WebSocket connections. Performance '
    'testing focuses on message delivery latency and WebSocket connection handling.'
)
chat_perf = [
    ['PERF-CHAT-001', 'Send Text Message', '500 concurrent messages', '<500ms', '99%', 'POST /api/communities/{id}/messages'],
    ['PERF-CHAT-002', 'Send Voice Message', '100 concurrent uploads', '<5 seconds', '95%', 'POST /api/communities/{id}/voice'],
    ['PERF-CHAT-003', 'Send File Attachment', '50 concurrent uploads', '<5 seconds', '95%', 'POST /api/communities/{id}/attachments'],
    ['PERF-CHAT-004', 'Load Message History', '200 concurrent requests', '<2 seconds', '99%', 'GET /api/communities/{id}/messages'],
    ['PERF-CHAT-005', 'WebSocket Connections', '500 simultaneous connections', 'Stable', '99%', 'STOMP WebSocket endpoint'],
    ['PERF-CHAT-006', 'Message Broadcasting', '500 clients in group', '<1 second', '99%', 'Topic message delivery'],
    ['PERF-CHAT-007', 'Message Search', '100 concurrent searches', '<3 seconds', '99%', 'GET /api/communities/{id}/messages/search'],
    ['PERF-CHAT-008', 'Reaction Toggle', '200 concurrent reactions', '<500ms', '99%', 'POST /api/communities/messages/{id}/reactions'],
    ['PERF-CHAT-009', 'Poll Voting', '100 concurrent votes', '<1 second', '99%', 'POST /api/communities/messages/{id}/vote'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], chat_perf)

doc.add_heading('11. Test Scenarios - Dashboard & Reports', level=1)
doc.add_paragraph(
    'Dashboard loading involves aggregating data from multiple tables and modules. Performance '
    'testing ensures dashboards load quickly despite complex queries.'
)
dash_perf = [
    ['PERF-DASH-001', 'Personal Dashboard', '200 concurrent requests', '<3 seconds', '99%', 'GET /api/dashboard/me'],
    ['PERF-DASH-002', 'Executive Dashboard', '20 concurrent requests', '<5 seconds', '95%', 'GET /api/dashboard/executive'],
    ['PERF-DASH-003', 'Org Insights', '20 concurrent requests', '<5 seconds', '95%', 'GET /api/dashboard/org-insights'],
    ['PERF-DASH-004', 'Celebrations', '200 concurrent requests', '<1 second', '99%', 'GET /api/dashboard/celebrations'],
    ['PERF-DASH-005', 'Attendance Report Excel', '10 concurrent exports', '<10 seconds', '95%', 'GET /api/reports/attendance/excel'],
    ['PERF-DASH-006', 'Payroll Summary', '20 concurrent requests', '<3 seconds', '99%', 'GET /api/payroll/salary-months'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], dash_perf)

doc.add_heading('12. Test Scenarios - File Operations', level=1)
doc.add_paragraph(
    'File operations include profile photos, payslip PDFs, chat attachments, and Excel exports. '
    'Performance testing validates file transfer speeds and storage I/O.'
)
file_perf = [
    ['PERF-FILE-001', 'Profile Photo Upload', '50 concurrent uploads', '<3 seconds', '95%', 'POST /api/files/upload'],
    ['PERF-FILE-002', 'Payslip PDF Download', '100 concurrent downloads', '<3 seconds', '99%', 'GET /api/payroll/payslip/{id}/pdf'],
    ['PERF-FILE-003', 'Chat File Upload', '50 concurrent uploads', '<5 seconds', '95%', 'POST /api/communities/{id}/attachments'],
    ['PERF-FILE-004', 'Excel Report Export', '20 concurrent exports', '<10 seconds', '95%', 'GET /api/reports/attendance/excel'],
    ['PERF-FILE-005', 'Asset QR Code', '100 concurrent requests', '<1 second', '99%', 'GET /api/assets/{id}/qr'],
    ['PERF-FILE-006', 'Bulk Employee Import', '1 import with 100 rows', '<30 seconds', '95%', 'POST /api/auth/employees/bulk'],
]
add_styled_table(doc, ['Scenario ID', 'Scenario', 'Load', 'Target', 'Pass Criteria', 'Endpoint'], file_perf)

doc.add_page_break()

# ═══ 13. PERFORMANCE METRICS & KPIs ═══
doc.add_heading('13. Performance Metrics & KPIs', level=1)
doc.add_paragraph(
    'The following Key Performance Indicators (KPIs) define the acceptable performance thresholds '
    'for the PIXOUS HR Portal. All metrics must be met during performance testing for the release '
    'to be approved for production deployment.'
)

doc.add_heading('13.1 Response Time Thresholds', level=2)
response_times = [
    ['Authentication (Login)', '<2 seconds', 'Critical user-facing operation'],
    ['Token Refresh', '<1 second', 'Transparent to user, must be fast'],
    ['Simple GET Requests', '<500ms', 'Dashboard widgets, balances, status checks'],
    ['CRUD Operations', '<2 seconds', 'Create, update, delete operations'],
    ['Search Operations', '<3 seconds', 'Search with filters and pagination'],
    ['File Upload (Small)', '<3 seconds', 'Profile photos, small attachments (<5MB)'],
    ['File Upload (Large)', '<10 seconds', 'Large files, Excel imports (>5MB)'],
    ['PDF Generation', '<5 seconds', 'Payslip PDF generation per employee'],
    ['Batch Operations', '<30 seconds', 'Payroll runs, bulk imports'],
    ['Report Generation', '<10 seconds', 'Excel report exports with large datasets'],
    ['WebSocket Message Delivery', '<500ms', 'Real-time chat and notification delivery'],
    ['Face Recognition', '<3 seconds', 'Face verification for attendance'],
]
add_styled_table(doc, ['Operation', 'Target Response Time', 'Notes'], response_times)

doc.add_heading('13.2 System Resource Thresholds', level=2)
resource_thresholds = [
    ['CPU Utilization', '<80%', 'Average across all cores during peak load'],
    ['Memory Utilization', '<85%', 'JVM heap and system memory combined'],
    ['Database Connections', '<80% of pool', 'HikariCP max pool size is 50'],
    ['Disk I/O', '<70% capacity', 'Read/write operations per second'],
    ['Network Bandwidth', '<80% capacity', 'Inbound and outbound traffic'],
    ['Redis Memory', '<80% of allocated', 'Cache eviction threshold'],
    ['Kafka Consumer Lag', '<1000 messages', 'Chat message processing backlog'],
    ['Thread Pool Utilization', '<80%', 'Tomcat max threads (200 default)'],
    ['JVM GC Pause Time', '<200ms', 'Maximum garbage collection pause'],
    ['Open File Descriptors', '<80% of limit', 'Operating system file handle limit'],
]
add_styled_table(doc, ['Metric', 'Threshold', 'Notes'], resource_thresholds)

doc.add_heading('13.3 Throughput Requirements', level=2)
throughput = [
    ['Login Requests', '200 req/sec', 'POST /api/auth/login'],
    ['API Requests (Aggregate)', '500 req/sec', 'All API endpoints combined'],
    ['WebSocket Connections', '500 simultaneous', 'STOMP/SockJS connections'],
    ['WebSocket Messages', '1000 msg/sec', 'Chat message broadcast rate'],
    ['File Uploads', '50 req/sec', 'Profile photos and attachments'],
    ['File Downloads', '100 req/sec', 'Payslip PDFs and reports'],
    ['Database Queries', '1000 queries/sec', 'MySQL query throughput'],
    ['PDF Generations', '20 PDFs/sec', 'Payslip PDF creation rate'],
    ['Excel Exports', '10 exports/sec', 'Report generation rate'],
]
add_styled_table(doc, ['Operation', 'Target Throughput', 'Notes'], throughput)

doc.add_page_break()

# ═══ 14. ROLE-WISE PERFORMANCE IMPACT ═══
doc.add_heading('14. Role-wise Performance Impact Analysis', level=1)
doc.add_paragraph(
    'Different user roles have different usage patterns and performance impact on the system. '
    'This section analyzes the performance characteristics of each role to ensure optimal resource '
    'allocation and user experience.'
)

doc.add_heading('14.1 Employee Role (Highest Volume)', level=2)
doc.add_paragraph(
    'Employees represent the largest user base (typically 80% of total users). Their usage is '
    'characterized by frequent self-service operations: attendance punches, leave applications, '
    'payslip downloads, helpdesk tickets, and chat messages.'
)
emp_perf = [
    ['Morning Login Burst', '8:30-9:30 AM', '200+ simultaneous logins', 'Login API must handle burst'],
    ['Punch In/Out', '9:00 AM / 6:00 PM', '200+ GPS punches within 30 min', 'Geofence validation must be fast'],
    ['Leave Application', 'Throughout day', '50 concurrent applications', 'Balance calculation must be atomic'],
    ['Payslip Download', 'Month-end', '100+ concurrent PDF downloads', 'PDF caching recommended'],
    ['Chat Activity', 'Throughout day', '500+ concurrent WebSocket connections', 'Message broadcast scalability'],
    ['Helpdesk Tickets', 'Throughout day', '100+ concurrent ticket operations', 'File upload performance'],
]
add_styled_table(doc, ['Activity', 'Peak Time', 'Expected Load', 'Performance Consideration'], emp_perf)

doc.add_heading('14.2 HR Manager Role (Batch Operations)', level=2)
doc.add_paragraph(
    'HR managers perform batch operations that can have significant system impact: payroll processing, '
    'bulk employee imports, leave approvals, and report generation. These operations must be optimized '
    'for batch efficiency.'
)
hr_perf = [
    ['Payroll Run', 'Month-end', '100 employees in single batch', 'Batch processing with progress tracking'],
    ['Bulk Employee Import', 'Hiring cycles', '100+ employees from Excel', 'Background processing with notification'],
    ['Leave Bulk Approval', 'After holidays', '50+ leave requests', 'Transaction batching for DB efficiency'],
    ['Report Generation', 'Weekly/Monthly', 'Excel exports with 500+ rows', 'Streaming response for large files'],
    ['Employee Directory', 'Throughout day', 'Paginated list with filters', 'Database index optimization'],
]
add_styled_table(doc, ['Activity', 'Peak Time', 'Expected Load', 'Performance Consideration'], hr_perf)

doc.add_heading('14.3 Team Lead Role (Approval Queues)', level=2)
doc.add_paragraph(
    'Team leads primarily interact with approval queues and team views. Their performance impact '
    'is moderate but time-sensitive as approvals affect employee workflows.'
)
tl_perf = [
    ['Leave Approvals', 'After holidays', '20+ pending approvals', 'Queue loading must be fast'],
    ['Team Attendance View', 'Morning', '50+ team members displayed', 'Efficient team member queries'],
    ['Task Assignment', 'Sprint planning', '20+ task creations', 'Batch task creation optimization'],
    ['Work Report Review', 'End of week', '10+ report reviews', 'Report aggregation queries'],
]
add_styled_table(doc, ['Activity', 'Peak Time', 'Expected Load', 'Performance Consideration'], tl_perf)

doc.add_heading('14.4 System Admin Role (System-Wide Operations)', level=2)
doc.add_paragraph(
    'System admins perform operations that affect the entire system: module management, cache '
    'operations, audit log reviews, and data resets. These operations must be carefully managed '
    'to avoid system-wide performance impact.'
)
admin_perf = [
    ['Cache Clear', 'As needed', 'Full Redis cache flush', 'Must not cause request spikes'],
    ['Module Toggle', 'Configuration', 'System-wide module changes', 'Permission cache refresh impact'],
    ['Audit Log Query', 'Investigation', '100K+ audit entries', 'Indexed queries with pagination'],
    ['Data Reset', 'Emergency', 'Full data restoration', 'Database-intensive, schedule off-hours'],
    ['Bulk Operations', 'Maintenance', 'System-wide changes', 'Background processing required'],
]
add_styled_table(doc, ['Activity', 'Peak Time', 'Expected Load', 'Performance Consideration'], admin_perf)

doc.add_heading('14.5 CEO/Executive Role (Read-Heavy Dashboards)', level=2)
doc.add_paragraph(
    'Executives primarily consume aggregated data through dashboards and reports. Their performance '
    'impact is read-heavy with complex queries that aggregate data across multiple modules.'
)
cto_perf = [
    ['Executive Dashboard', 'Morning', 'KPI aggregation across all modules', 'Pre-computed aggregates recommended'],
    ['Org Insights', 'Weekly', 'Organization-wide statistics', 'Cached materialized views'],
    ['Report Downloads', 'Monthly', 'Comprehensive Excel reports', 'Background generation with notification'],
    ['Attendance Overview', 'Daily', 'All-employee attendance summary', 'Optimized aggregate queries'],
]
add_styled_table(doc, ['Activity', 'Peak Time', 'Expected Load', 'Performance Consideration'], cto_perf)

doc.add_page_break()

# ═══ 15. PERFORMANCE TEST RESULTS TEMPLATE ═══
doc.add_heading('15. Performance Test Results Template', level=1)
doc.add_paragraph(
    'Use this template to document performance test results for each test execution. Results should '
    'be compared against the defined KPIs and thresholds to determine pass/fail status.'
)
results_template = [
    ['Test Execution ID', 'PERF-EXEC-XXX'],
    ['Test Date', 'YYYY-MM-DD'],
    ['Tester', 'QA Engineer name'],
    ['Environment', 'Staging / Pre-production'],
    ['Test Type', 'Load / Stress / Spike / Endurance / Scalability'],
    ['Total Concurrent Users', 'Number of simulated users'],
    ['Test Duration', 'Total test execution time'],
    ['Average Response Time', 'Mean response time across all requests'],
    ['95th Percentile Response Time', '95% of requests completed within this time'],
    ['99th Percentile Response Time', '99% of requests completed within this time'],
    ['Throughput (req/sec)', 'Total requests processed per second'],
    ['Error Rate', 'Percentage of failed requests'],
    ['Peak CPU Utilization', 'Maximum CPU usage during test'],
    ['Peak Memory Utilization', 'Maximum memory usage during test'],
    ['Database Connection Peak', 'Maximum DB connections used'],
    ['Pass/Fail Status', 'Based on KPI thresholds'],
    ['Bottlenecks Identified', 'List of performance bottlenecks found'],
    ['Recommendations', 'Optimization recommendations'],
]
add_styled_table(doc, ['Field', 'Description'], results_template)

# ═══ 16. BOTTLENECK ANALYSIS FRAMEWORK ═══
doc.add_heading('16. Bottleneck Analysis Framework', level=1)
doc.add_paragraph(
    'When performance issues are identified during testing, use this framework to systematically '
    'analyze and resolve bottlenecks. The analysis follows a top-down approach from application '
    'layer to infrastructure layer.'
)
bottleneck = [
    ['Application Code', 'Inefficient algorithms, N+1 queries, missing indexes', 'Code profiling, query optimization, index creation'],
    ['JVM Configuration', 'Heap size, GC tuning, thread pool sizing', 'JVM flags tuning, GC algorithm selection'],
    ['Database', 'Slow queries, missing indexes, lock contention', 'EXPLAIN analysis, index optimization, query rewriting'],
    ['Connection Pool', 'Pool exhaustion, idle connections, leak detection', 'HikariCP tuning, connection leak detection'],
    ['Caching', 'Cache miss rate, eviction, stale data', 'Redis configuration, cache warming strategies'],
    ['Network', 'Latency, bandwidth, DNS resolution', 'CDN implementation, connection pooling'],
    ['File I/O', 'Disk speed, file handle limits, temp file cleanup', 'SSD upgrade, file handle limits, cleanup cron'],
    ['External Services', 'SMS provider latency, face recognition processing', 'Async processing, circuit breakers, fallbacks'],
]
add_styled_table(doc, ['Layer', 'Common Issues', 'Resolution Approach'], bottleneck)

# ═══ 17. CAPACITY PLANNING ═══
doc.add_heading('17. Capacity Planning Guidelines', level=1)
doc.add_paragraph(
    'Based on performance testing results, use these guidelines to plan production infrastructure '
    'capacity. Capacity planning should account for peak load, growth projections, and safety margins.'
)
capacity = [
    ['Small (100 users)', '2 vCPU, 4GB RAM', '2 vCPU, 8GB RAM', '100 Mbps', 'MySQL + Redis on same server'],
    ['Medium (500 users)', '4 vCPU, 8GB RAM', '4 vCPU, 16GB RAM', '500 Mbps', 'Separate DB and cache servers'],
    ['Large (2000 users)', '8 vCPU, 16GB RAM', '8 vCPU, 32GB RAM', '1 Gbps', 'Load balancer + multiple app servers'],
    ['Enterprise (5000+)', '16 vCPU, 32GB RAM', '16 vCPU, 64GB RAM', '10 Gbps', 'Multi-region, auto-scaling, read replicas'],
]
add_styled_table(doc, ['Scale', 'Minimum Server', 'Recommended Server', 'Network', 'Architecture'], capacity)

doc.add_heading('Growth Projection Model', level=2)
doc.add_paragraph(
    'When planning capacity, consider the following growth factors: user count growth (typically '
    '20-30% annually), data volume growth (attendance and chat data accumulate), feature additions '
    '(new modules increase system complexity), and seasonal peaks (month-end payroll, holiday seasons).'
)
doc.add_paragraph(
    'It is recommended to plan capacity for 2x the current expected peak load to accommodate '
    'unexpected growth and provide a buffer for system stability. Regular performance testing '
    'should be conducted quarterly to validate capacity assumptions.'
)

# ═══ 18. RECOMMENDATIONS ═══
doc.add_heading('18. Recommendations', level=1)
doc.add_paragraph(
    'Based on the performance testing strategy and expected system characteristics, the following '
    'recommendations are provided to ensure optimal performance of the PIXOUS HR Portal.'
)

doc.add_heading('18.1 Immediate Actions', level=2)
immediate = [
    ['Database Indexing', 'Ensure all frequently queried columns have proper indexes', 'High'],
    ['Query Optimization', 'Review and optimize N+1 queries and complex joins', 'High'],
    ['Connection Pool Tuning', 'Configure HikariCP pool size based on expected concurrency', 'High'],
    ['JWT Cache', 'Cache user permissions in Redis to reduce DB hits per request', 'Medium'],
    ['WebSocket Optimization', 'Implement message batching for high-traffic chat groups', 'Medium'],
    ['PDF Caching', 'Cache generated payslip PDFs to avoid regeneration', 'Medium'],
    ['Rate Limiting', 'Implement per-user rate limiting to prevent abuse', 'High'],
]
add_styled_table(doc, ['Action', 'Description', 'Priority'], immediate)

doc.add_heading('18.2 Long-term Improvements', level=2)
longterm = [
    ['CDN Implementation', 'Serve static assets via CDN for faster page loads', 'Medium'],
    ['Database Read Replicas', 'Add read replicas for dashboard and report queries', 'Medium'],
    ['Message Queue Optimization', 'Tune Kafka partition count and consumer group settings', 'Low'],
    ['Monitoring Dashboard', 'Set up Grafana dashboards for real-time performance monitoring', 'High'],
    ['Automated Performance Tests', 'Integrate JMeter tests into CI/CD pipeline', 'Medium'],
    ['Cache Warming', 'Implement cache warming for frequently accessed data', 'Low'],
    ['Background Processing', 'Move batch operations to background jobs with progress tracking', 'Medium'],
]
add_styled_table(doc, ['Improvement', 'Description', 'Priority'], longterm)

doc.add_page_break()

# ═══ 19. APPENDIX: JMETER CONFIGURATION ═══
doc.add_heading('19. Appendix: Sample JMeter Test Plan Configuration', level=1)
doc.add_paragraph(
    'This appendix provides sample Apache JMeter configuration for executing performance tests '
    'against the PIXOUS HR Portal API endpoints. Use this as a starting point for test plan creation.'
)

doc.add_heading('Thread Group Configuration', level=2)
jmeter_config = [
    ['Thread Group Name', 'HR Portal Load Test'],
    ['Number of Threads (Users)', '200'],
    ['Ramp-Up Period', '600 seconds (10 minutes)'],
    ['Loop Count', '1 (or infinite for endurance)'],
    ['Scheduler Duration', '1800 seconds (30 minutes)'],
    ['Delay creation of threads until needed', 'True'],
]
add_styled_table(doc, ['Parameter', 'Value'], jmeter_config)

doc.add_heading('HTTP Request Defaults', level=2)
http_defaults = [
    ['Server Name', 'api.pixoustech.com'],
    ['Protocol', 'https'],
    ['Port', '443'],
    ['Content-Type', 'application/json'],
    ['Connection Timeout', '30000 ms'],
    ['Response Timeout', '30000 ms'],
]
add_styled_table(doc, ['Parameter', 'Value'], http_defaults)

doc.add_heading('Sample HTTP Request - Login', level=2)
login_config = [
    ['Method', 'POST'],
    ['Path', '/api/auth/login'],
    ['Body Data', '{"username": "${username}", "password": "${password}"}'],
    ['Extract Token', 'JSON Path: $.accessToken'],
    ['Header', 'Content-Type: application/json'],
]
add_styled_table(doc, ['Parameter', 'Value'], login_config)

doc.add_heading('Listeners & Reports', level=2)
listeners = [
    ['Summary Report', 'Aggregate statistics for all requests'],
    ['Aggregate Report', 'Detailed statistics with percentiles'],
    ['View Results Tree', 'Debug mode only (disable for load tests)'],
    ['Response Time Graph', 'Visual response time distribution'],
    ['HTML Dashboard Report', 'Comprehensive HTML report generation'],
    ['Backend Listener', 'Send metrics to InfluxDB for Grafana dashboards'],
]
add_styled_table(doc, ['Listener', 'Purpose'], listeners)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PIXOUS Technologies Pvt. Ltd. | Confidential')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save('PIXOUS_HR_Portal_Performance_Testing_Document.docx')
print("Document 3 generated successfully: PIXOUS_HR_Portal_Performance_Testing_Document.docx")
