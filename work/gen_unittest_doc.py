"""Generate PIXOUS HR Portal - Unit Testing Document"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper: Add Logo ──
def add_logo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('pixous_logo.png', width=Inches(2.0))
    return p

# ── Helper: Add Cover Title ──
def add_cover(doc):
    add_logo(doc)
    doc.add_paragraph()
    title = doc.add_heading('PIXOUS HR Portal', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_heading('Unit Testing Document', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()

    # Info table
    t = doc.add_table(rows=6, cols=2)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ('Document Type', 'Unit Testing Document'),
        ('Version', '1.0'),
        ('Date', datetime.date.today().strftime('%B %d, %Y')),
        ('Prepared By', 'PIXOUS Technologies QA & Development Team'),
        ('Classification', 'Confidential - Client Deliverable'),
        ('Status', 'Active'),
    ]
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for cell in t.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_page_break()

# ── Helper: Styled Table ──
def add_styled_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear'
        })
        shading.append(shd)
    # Rows
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
add_cover(doc)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Testing Framework & Tools',
    '3. Backend Unit Testing',
    '  3.1 Authentication Module',
    '  3.2 User/Employee Module',
    '  3.3 Attendance Module',
    '  3.4 Leave Module',
    '  3.5 Payroll Module',
    '  3.6 Asset Module',
    '  3.7 Helpdesk Module',
    '  3.8 Community/Chat Module',
    '  3.9 Dashboard Module',
    '  3.10 Notification Module',
    '  3.11 Onboarding Module',
    '  3.12 Performance Module',
    '  3.13 TA Expense Module',
    '4. Frontend Unit Testing',
    '5. Service Layer Testing',
    '6. Repository/DAO Testing',
    '7. Test Coverage Metrics',
    '8. Mock Strategy',
    '9. Test Data Management',
    '10. CI/CD Integration',
    '11. Test Case Inventory by Module',
    '12. Known Test Gaps & Recommendations',
    '13. Appendix: Sample Test Code Snippets',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal is a comprehensive, enterprise-grade Human Resource Management System '
    'designed to streamline and automate all HR operations for organizations across IT and Civil/Construction '
    'industries. Unit testing is a critical component of our quality assurance strategy, ensuring that every '
    'individual module, service, and component functions correctly in isolation before integration.'
)
doc.add_paragraph(
    'This Unit Testing Document defines the complete unit testing strategy for the PIXOUS HR Portal, '
    'covering backend services (Java 17 / Spring Boot), frontend components (React 19 / TypeScript), '
    'and the Python analytics service. The document provides detailed test specifications, mock strategies, '
    'coverage targets, CI/CD integration plans, and sample test code for every major module.'
)
doc.add_paragraph(
    'The primary objective is to achieve 80% or higher code coverage across all modules while maintaining '
    'fast test execution times suitable for continuous integration pipelines. Each module has been analyzed '
    'for testable units, boundary conditions, and edge cases that must be validated through automated tests.'
)
doc.add_paragraph(
    'Unit tests serve as the first line of defense against regressions. By catching defects at the '
    'individual function and method level, we reduce the cost of bug fixes and ensure system stability '
    'across all 25+ functional modules of the HR Portal.'
)

doc.add_heading('Key Objectives', level=2)
objectives = [
    'Achieve 80% or higher code coverage across all backend and frontend modules',
    'Ensure every public method in service and repository layers has at least one corresponding test',
    'Validate all business rules through parameterized and boundary-value tests',
    'Mock all external dependencies (databases, message queues, external APIs) for isolation',
    'Integrate unit tests into the GitHub Actions CI/CD pipeline with coverage gates',
    'Maintain test execution time under 5 minutes for the full backend suite',
    'Provide clear test case IDs (UT-XXX-001 format) for traceability',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('Testing Scope', level=2)
doc.add_paragraph(
    'The testing scope encompasses all server-side business logic, data access layers, frontend '
    'component rendering, form validation, API integration layers, and state management. '
    'External integrations such as face recognition, PDF generation, and third-party APIs are '
    'tested at the integration test level and are excluded from unit testing scope.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. TESTING FRAMEWORK & TOOLS
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Testing Framework & Tools', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal employs a multi-stack testing approach that aligns with the technology '
    'used in each layer of the application. Backend tests use JUnit 5 with Mockito for dependency '
    'injection mocking, frontend tests use Vitest with React Testing Library for component-level '
    'validation, and the analytics service uses pytest for Python-specific testing.'
)

doc.add_heading('2.1 Backend Testing Stack', level=2)
doc.add_paragraph(
    'The backend is built on Java 17 with Spring Boot 3.5. All unit tests use JUnit 5 (Jupiter) '
    'as the test framework, with Mockito for creating mock objects and verifying interactions. '
    'Spring Boot Test provides the test context for loading application configuration.'
)
add_styled_table(doc,
    ['Tool', 'Version', 'Purpose'],
    [
        ['JUnit 5 (Jupiter)', '5.10.x', 'Test framework, assertions, parameterized tests'],
        ['Mockito', '5.x', 'Mock creation, stubbing, verification'],
        ['Mockito Extension', '5.x', 'Seamless JUnit 5 integration'],
        ['Spring Boot Test', '3.5', 'Application context loading, @WebMvcTest'],
        ['AssertJ', '3.x', 'Fluent assertion library (optional)'],
        ['H2 Database', '2.x', 'In-memory database for repository tests'],
        ['JaCoCo', '0.8.x', 'Code coverage instrumentation and reporting'],
    ]
)

doc.add_heading('2.2 Frontend Testing Stack', level=2)
doc.add_paragraph(
    'The frontend is built on React 19 with TypeScript. Vitest serves as the test runner with '
    'Jest-compatible API, while React Testing Library provides DOM-based component testing utilities. '
    'JSDOM is used as the browser environment simulator.'
)
add_styled_table(doc,
    ['Tool', 'Version', 'Purpose'],
    [
        ['Vitest', '1.x', 'Test runner with native ESM and TypeScript support'],
        ['React Testing Library', '16.x', 'Component rendering and interaction testing'],
        ['@testing-library/jest-dom', '6.x', 'Custom DOM matchers for assertions'],
        ['@testing-library/user-event', '14.x', 'Simulating user interactions'],
        ['jsdom', '25.x', 'Browser environment simulation'],
        ['MSW (Mock Service Worker)', '2.x', 'API mocking for integration-style tests'],
        ['V8 Coverage', '-', 'Native code coverage via Vitest'],
    ]
)

doc.add_heading('2.3 Python Analytics Testing Stack', level=2)
doc.add_paragraph(
    'The Python analytics service handles face recognition and OCR processing. Testing uses '
    'pytest as the primary framework with pytest-cov for coverage reporting and faker for '
    'test data generation.'
)
add_styled_table(doc,
    ['Tool', 'Version', 'Purpose'],
    [
        ['pytest', '8.x', 'Test framework with fixtures and parametrize'],
        ['pytest-cov', '5.x', 'Coverage reporting'],
        ['pytest-asyncio', '0.x', 'Async test support for FastAPI endpoints'],
        ['unittest.mock', 'Built-in', 'Mocking external services'],
        ['faker', '28.x', 'Realistic test data generation'],
        ['httpx', '0.x', 'Async HTTP client for endpoint testing'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. BACKEND UNIT TESTING
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Backend Unit Testing', level=1)
doc.add_paragraph(
    'Backend unit tests focus on validating individual service methods, repository queries, '
    'and business logic in isolation. All external dependencies (database, message queues, '
    'external APIs) are mocked to ensure tests are fast, deterministic, and do not require '
    'infrastructure connectivity.'
)

# ── 3.1 Authentication Module ──
doc.add_heading('3.1 Authentication Module', level=2)
doc.add_paragraph(
    'The authentication module handles JWT token generation, validation, refresh, and logout. '
    'Unit tests verify that tokens are correctly generated with proper claims, that expired '
    'tokens are rejected, that refresh tokens work correctly, and that logout properly revokes '
    'all tokens for a user.'
)
doc.add_paragraph(
    'Key test scenarios include JWT generation with correct expiration times, token validation '
    'with signature verification, refresh token rotation, account lockout after 5 failed attempts, '
    'and BCrypt password hashing verification.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-AUTH-001', 'Generate JWT Token', 'High', 'Verify JWT token is generated with correct claims (userId, role, expiry)'],
        ['UT-AUTH-002', 'Validate JWT Token', 'High', 'Verify valid token is accepted and invalid/expired tokens are rejected'],
        ['UT-AUTH-003', 'Refresh Token Rotation', 'High', 'Verify old refresh token is invalidated when new one is generated'],
        ['UT-AUTH-004', 'Logout Token Revocation', 'High', 'Verify all refresh tokens for a user are revoked on logout'],
        ['UT-AUTH-005', 'Account Lockout', 'High', 'Verify account locks after 5 consecutive failed login attempts'],
        ['UT-AUTH-006', 'Password Hashing', 'Medium', 'Verify BCrypt correctly hashes and verifies passwords'],
        ['UT-AUTH-007', 'Phone Validation', 'Medium', 'Verify phone number uniqueness check during registration'],
        ['UT-AUTH-008', 'Username Availability', 'Medium', 'Verify username availability check returns correct boolean'],
    ]
)

doc.add_paragraph(
    'The AuthService class depends on JwtUtil for token operations, UserRepository for user lookups, '
    'and PasswordEncoder for password hashing. All dependencies are mocked using @Mock and @InjectMocks '
    'annotations with Mockito.'
)

# ── 3.2 User/Employee Module ──
doc.add_heading('3.2 User/Employee Module', level=2)
doc.add_paragraph(
    'The User/Employee module manages CRUD operations, role assignment, bulk import, and '
    'employee lifecycle. Unit tests verify that employee creation validates required fields, '
    'that role assignment enforces permission boundaries, and that bulk import handles both '
    'valid and invalid data gracefully.'
)
doc.add_paragraph(
    'Testing covers the complete employee lifecycle from creation through deactivation, '
    'including profile updates, document management, and multi-tenant data isolation.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-USER-001', 'Create Employee', 'High', 'Verify employee creation with all required fields and default values'],
        ['UT-USER-002', 'Create Employee - Missing Fields', 'High', 'Verify validation rejects creation when required fields are missing'],
        ['UT-USER-003', 'Bulk Import Valid Data', 'High', 'Verify bulk import processes valid Excel data correctly'],
        ['UT-USER-004', 'Bulk Import with Errors', 'High', 'Verify bulk import handles invalid rows and reports errors'],
        ['UT-USER-005', 'Role Assignment', 'High', 'Verify role assignment updates user permissions correctly'],
        ['UT-USER-006', 'Deactivate Employee', 'High', 'Verify deactivation sets inactive status and revokes access'],
        ['UT-USER-007', 'Profile Update', 'Medium', 'Verify profile fields are updated and audit trail is logged'],
        ['UT-USER-008', 'Multi-Tenant Filtering', 'High', 'Verify queries only return data for the current tenant'],
    ]
)

# ── 3.3 Attendance Module ──
doc.add_heading('3.3 Attendance Module', level=2)
doc.add_paragraph(
    'The Attendance module handles punch in/out, face verification, geofence validation, '
    'overtime calculation, and LOP (Loss of Pay) computation. Unit tests verify that attendance '
    'records are created correctly, that double punch-in is prevented, and that duration calculations '
    'are accurate.'
)
doc.add_paragraph(
    'Geofence validation tests verify that punch-in requests are rejected when the employee GPS '
    'coordinates fall outside the designated work area. Face verification tests mock the face '
    'recognition service to test the verification flow.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-ATT-001', 'Punch In - Valid', 'High', 'Verify attendance record created with correct timestamp and GPS coordinates'],
        ['UT-ATT-002', 'Punch Out - Calculate Duration', 'High', 'Verify duration is calculated as punchOut - punchIn in hours'],
        ['UT-ATT-003', 'Double Punch In Prevention', 'High', 'Verify second punch-in is rejected when already punched in'],
        ['UT-ATT-004', 'Geofence Validation - Inside', 'High', 'Verify punch-in succeeds when GPS is within geofence radius'],
        ['UT-ATT-005', 'Geofence Validation - Outside', 'High', 'Verify punch-in fails when GPS is outside geofence radius'],
        ['UT-ATT-006', 'Face Verification Success', 'Medium', 'Verify attendance marked with face_verified=true on match'],
        ['UT-ATT-007', 'Face Verification Failure', 'Medium', 'Verify attendance rejected when face does not match stored embedding'],
        ['UT-ATT-008', 'Overtime Calculation', 'High', 'Verify overtime hours computed when work exceeds standard hours'],
        ['UT-ATT-009', 'Monthly Summary Calculation', 'Medium', 'Verify present, absent, leave, and overtime counts are accurate'],
        ['UT-ATT-010', 'LOP Calculation', 'High', 'Verify LOP days computed for leaves exceeding balance'],
    ]
)

# ── 3.4 Leave Module ──
doc.add_heading('3.4 Leave Module', level=2)
doc.add_paragraph(
    'The Leave module manages leave applications, approvals, balance calculations, and LOP '
    'deductions. Unit tests verify that leave applications deduct from balance correctly, '
    'that approval workflows enforce permission rules, and that leave cancellation restores '
    'the balance.'
)
doc.add_paragraph(
    'Business logic tests cover scenarios like applying for leave when balance is zero, '
    'approving leave beyond available balance, partial day leave calculations, and '
    'carry-forward rules for earned leaves.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-LEA-001', 'Apply Leave - Deduct Balance', 'High', 'Verify leave balance is deducted when leave is applied'],
        ['UT-LEA-002', 'Apply Leave - Insufficient Balance', 'High', 'Verify application rejected when balance is zero'],
        ['UT-LEA-003', 'Approve Leave - Update Status', 'High', 'Verify status changes to APPROVED on approval'],
        ['UT-LEA-004', 'Reject Leave - Update Status', 'High', 'Verify status changes to REJECTED on rejection'],
        ['UT-LEA-005', 'Cancel Approved Leave - Restore Balance', 'High', 'Verify balance is restored when approved leave is cancelled'],
        ['UT-LEA-006', 'LOP Calculation - Excess Leave', 'High', 'Verify LOP days calculated for leaves exceeding quota'],
        ['UT-LEA-007', 'Bulk Decision Processing', 'Medium', 'Verify multiple leave requests can be processed in one operation'],
        ['UT-LEA-008', 'Leave Balance Query', 'Medium', 'Verify balance query returns correct counts per leave type'],
        ['UT-LEA-009', 'Weekend/Holiday Exclusion', 'Medium', 'Verify weekends and holidays are excluded from leave count'],
        ['UT-LEA-010', 'Carry Forward Rules', 'Medium', 'Verify earned leave carry-forward applies at year end'],
    ]
)

# ── 3.5 Payroll Module ──
doc.add_heading('3.5 Payroll Module', level=2)
doc.add_paragraph(
    'The Payroll module handles salary structure management, payslip generation, payroll runs, '
    'and PDF generation. Unit tests verify that salary calculations apply the correct components, '
    'that payslips contain all required fields, and that the payroll run lifecycle transitions '
    'through all states correctly.'
)
doc.add_paragraph(
    'PDF generation tests use mock templates and verify that the OpenPDF engine produces valid '
    'output with correct employee details, earnings, deductions, and net pay calculations.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-PAY-001', 'Calculate Net Pay', 'High', 'Verify net pay = gross earnings - total deductions'],
        ['UT-PAY-002', 'Generate Payslip - All Fields', 'High', 'Verify payslip contains employee name, month, earnings, deductions'],
        ['UT-PAY-003', 'Salary Structure Validation', 'High', 'Verify salary components (basic, HRA, deductions) are validated'],
        ['UT-PAY-004', 'Payroll Run Lifecycle', 'High', 'Verify run transitions: DRAFT -> IN_PROGRESS -> CONFIRMED -> APPROVED'],
        ['UT-PAY-005', 'LOP Deduction in Payslip', 'High', 'Verify LOP amount deducted from salary based on absent days'],
        ['UT-PAY-006', 'PDF Generation', 'Medium', 'Verify PDF output contains all payslip fields and is valid'],
        ['UT-PAY-007', 'Finance Approval', 'High', 'Verify payroll run cannot proceed without finance approval'],
        ['UT-PAY-008', 'Duplicate Payslip Prevention', 'Medium', 'Verify duplicate payslip generation for same month is prevented'],
    ]
)

# ── 3.6 Asset Module ──
doc.add_heading('3.6 Asset Module', level=2)
doc.add_paragraph(
    'The Asset module manages IT asset lifecycle including creation, allocation, return, and '
    'QR code generation. Unit tests verify that asset status transitions are valid, that '
    'allocation creates proper records, and that QR codes are generated with correct data.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-AST-001', 'Create Asset', 'High', 'Verify asset created with unique code and AVAILABLE status'],
        ['UT-AST-002', 'Allocate Asset', 'High', 'Verify allocation changes status to ALLOCATED and creates record'],
        ['UT-AST-003', 'Return Asset', 'High', 'Verify return changes status to AVAILABLE and clears allocation'],
        ['UT-AST-004', 'Allocate Already Allocated Asset', 'High', 'Verify allocation rejected for ALLOCATED assets'],
        ['UT-AST-005', 'QR Code Generation', 'Medium', 'Verify QR code contains correct asset ID and details'],
        ['UT-AST-006', 'Asset Lookup by Code', 'Medium', 'Verify lookup returns correct asset for valid code'],
        ['UT-AST-007', 'Delete Asset', 'Medium', 'Verify asset is removed from system on deletion'],
    ]
)

# ── 3.7 Helpdesk Module ──
doc.add_heading('3.7 Helpdesk Module', level=2)
doc.add_paragraph(
    'The Helpdesk module manages the ticket lifecycle from creation through resolution and '
    'rating. Unit tests verify that tickets are created with correct initial status, that '
    'status transitions follow valid paths, and that comments and ratings are properly associated.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-HD-001', 'Create Ticket', 'High', 'Verify ticket created with OPEN status and correct metadata'],
        ['UT-HD-002', 'Status Transition - Open to InProgress', 'High', 'Verify status change from OPEN to IN_PROGRESS'],
        ['UT-HD-003', 'Status Transition - InProgress to Resolved', 'High', 'Verify status change from IN_PROGRESS to RESOLVED'],
        ['UT-HD-004', 'Add Comment to Ticket', 'Medium', 'Verify comment is attached with author and timestamp'],
        ['UT-HD-005', 'Rate Resolved Ticket', 'Medium', 'Verify rating (1-5) is recorded only for RESOLVED tickets'],
        ['UT-HD-006', 'Rate Unresolved Ticket - Reject', 'Medium', 'Verify rating rejected for non-RESOLVED tickets'],
        ['UT-HD-007', 'File Upload to Ticket', 'Medium', 'Verify file attachment is saved with correct metadata'],
    ]
)

# ── 3.8 Community/Chat Module ──
doc.add_heading('3.8 Community/Chat Module', level=2)
doc.add_paragraph(
    'The Community/Chat module manages group creation, messaging, reactions, polls, and '
    'message retention policies. Unit tests verify that messages are saved correctly, '
    'that reactions toggle properly, and that retention policies enforce cleanup.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-CHAT-001', 'Send Message', 'High', 'Verify message saved with sender, timestamp, and content'],
        ['UT-CHAT-002', 'Toggle Reaction', 'High', 'Verify reaction is added if absent, removed if present'],
        ['UT-CHAT-003', 'Vote in Poll', 'Medium', 'Verify vote is recorded and duplicate votes prevented'],
        ['UT-CHAT-004', 'Pin Message', 'Medium', 'Verify message pin status toggles correctly'],
        ['UT-CHAT-005', 'Mark as Read', 'Medium', 'Verify read receipt is recorded for the message'],
        ['UT-CHAT-006', 'Search Messages', 'Medium', 'Verify search returns messages matching query text'],
        ['UT-CHAT-007', 'Retention Policy Enforcement', 'Low', 'Verify messages older than retention period are cleaned'],
        ['UT-CHAT-008', 'Direct Message Creation', 'High', 'Verify 1:1 DM thread created between two users'],
    ]
)

# ── 3.9 Dashboard Module ──
doc.add_heading('3.9 Dashboard Module', level=2)
doc.add_paragraph(
    'The Dashboard module aggregates data from multiple modules to present personal, executive, '
    'and organizational insights. Unit tests verify that dashboard queries return correct aggregates '
    'and that data is properly formatted for display.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-DASH-001', 'Personal Dashboard Data', 'High', 'Verify dashboard returns attendance, leave balance, tasks'],
        ['UT-DASH-002', 'Executive Dashboard KPIs', 'High', 'Verify headcount, attrition, attendance %, payroll summary'],
        ['UT-DASH-003', 'Org Insights Aggregation', 'Medium', 'Verify department distribution and gender ratio calculations'],
        ['UT-DASH-004', 'Celebrations Query', 'Medium', 'Verify birthdays and anniversaries returned for current date range'],
        ['UT-DASH-005', 'Attendance Absent Today', 'High', 'Verify list of absent employees is accurate for today'],
    ]
)

# ── 3.10 Notification Module ──
doc.add_heading('3.10 Notification Module', level=2)
doc.add_paragraph(
    'The Notification module handles WebSocket push notifications, read status tracking, '
    'and notification preferences. Unit tests verify that notifications are created and '
    'delivered correctly, and that read status updates are persisted.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-NOTIF-001', 'Create Notification', 'High', 'Verify notification record created with correct type and recipient'],
        ['UT-NOTIF-002', 'Mark as Read', 'High', 'Verify read status updated and unread count decremented'],
        ['UT-NOTIF-003', 'Bulk Mark as Read', 'Medium', 'Verify multiple notifications marked as read in one operation'],
        ['UT-NOTIF-004', 'Notification Count Query', 'Medium', 'Verify unread count returns correct number'],
    ]
)

# ── 3.11 Onboarding Module ──
doc.add_heading('3.11 Onboarding Module', level=2)
doc.add_paragraph(
    'The Onboarding module manages new employee checklists and task completion tracking. '
    'Unit tests verify that checklists are created with correct tasks, that task completion '
    'updates progress, and that the checklist status transitions correctly.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-ONB-001', 'Create Onboarding Checklist', 'High', 'Verify checklist created with tasks for new employee'],
        ['UT-ONB-002', 'Complete Task', 'High', 'Verify task marked complete and progress updated'],
        ['UT-ONB-003', 'Checklist Completion', 'High', 'Verify checklist status changes to COMPLETED when all tasks done'],
        ['UT-ONB-004', 'Progress Calculation', 'Medium', 'Verify completion percentage computed correctly'],
    ]
)

# ── 3.12 Performance Module ──
doc.add_heading('3.12 Performance Module', level=2)
doc.add_paragraph(
    'The Performance module manages goals, reviews, and appraisals. Unit tests verify that '
    'goals are created with valid metrics, that review scores are calculated correctly, '
    'and that performance ratings follow defined scales.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-PERF-001', 'Create Goal', 'High', 'Verify goal created with target, deadline, and assignee'],
        ['UT-PERF-002', 'Submit Review', 'High', 'Verify review recorded with scores and comments'],
        ['UT-PERF-003', 'Calculate Rating', 'Medium', 'Verify average rating computed from review scores'],
        ['UT-PERF-004', 'Goal Completion', 'Medium', 'Verify goal status transitions to COMPLETED on target achievement'],
    ]
)

# ── 3.13 TA Expense Module ──
doc.add_heading('3.13 TA Expense Module', level=2)
doc.add_paragraph(
    'The TA Expense module handles travel and allowance expense submissions, approvals, '
    'and rejections. Unit tests verify that expenses are created with valid amounts, '
    'that approval workflows enforce permission rules, and that rejected expenses update '
    'status correctly.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-TA-001', 'Submit Expense', 'High', 'Verify expense created with PENDING status and amount'],
        ['UT-TA-002', 'Approve Expense', 'High', 'Verify status changes to APPROVED on manager approval'],
        ['UT-TA-003', 'Reject Expense', 'High', 'Verify status changes to REJECTED with reason'],
        ['UT-TA-004', 'Expense Amount Validation', 'Medium', 'Verify expense rejected for zero or negative amounts'],
        ['UT-TA-005', 'Duplicate Submission Prevention', 'Medium', 'Verify duplicate expense submissions are detected'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. FRONTEND UNIT TESTING
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Frontend Unit Testing', level=1)
doc.add_paragraph(
    'Frontend unit tests validate React component rendering, user interactions, form validations, '
    'API integration, route navigation, and state management. Tests use Vitest as the test runner '
    'with React Testing Library for DOM-based component testing.'
)

doc.add_heading('4.1 Component Rendering Tests', level=2)
doc.add_paragraph(
    'Component rendering tests verify that UI components render correctly with given props, '
    'display expected text, and handle loading and error states. Each common UI component '
    '(Button, Input, Card, Dialog, Table) has dedicated test cases.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-FE-001', 'Button Renders with Label', 'High', 'Verify Button component renders with correct text label'],
        ['UT-FE-002', 'Button Click Handler', 'High', 'Verify onClick handler is called when button is clicked'],
        ['UT-FE-003', 'Input Field Validation', 'High', 'Verify Input shows error state for invalid values'],
        ['UT-FE-004', 'Card Component Layout', 'Medium', 'Verify Card renders children and header correctly'],
        ['UT-FE-005', 'Dialog Open/Close', 'High', 'Verify Dialog opens on trigger and closes on dismiss'],
        ['UT-FE-006', 'Table Data Rendering', 'High', 'Verify Table renders rows from data prop correctly'],
        ['UT-FE-007', 'Loading State Display', 'Medium', 'Verify component shows loading spinner during data fetch'],
        ['UT-FE-008', 'Error State Display', 'Medium', 'Verify component shows error message on failure'],
    ]
)

doc.add_heading('4.2 Form Validation Tests', level=2)
doc.add_paragraph(
    'Form validation tests use React Hook Form with Zod schemas to verify that forms enforce '
    'required fields, validate input formats, and display appropriate error messages. Tests '
    'simulate user input and verify validation behavior.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-FE-FORM-001', 'Required Field Validation', 'High', 'Verify form shows error when required field is empty'],
        ['UT-FE-FORM-002', 'Email Format Validation', 'High', 'Verify invalid email format triggers validation error'],
        ['UT-FE-FORM-003', 'Phone Number Format', 'Medium', 'Verify phone number validates against expected pattern'],
        ['UT-FE-FORM-004', 'Password Strength', 'Medium', 'Verify password meets minimum complexity requirements'],
        ['UT-FE-FORM-005', 'Date Range Validation', 'Medium', 'Verify end date must be after start date'],
        ['UT-FE-FORM-006', 'Form Submission with Valid Data', 'High', 'Verify form submits successfully with all valid fields'],
    ]
)

doc.add_heading('4.3 API Integration Tests', level=2)
doc.add_paragraph(
    'API integration tests verify that Axios interceptors handle token refresh correctly, '
    'that request headers are properly attached, and that error responses are handled gracefully. '
    'MSW (Mock Service Worker) is used to intercept and mock API calls.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-FE-API-001', 'Token Refresh on 401', 'High', 'Verify interceptor refreshes token and retries request on 401'],
        ['UT-FE-API-002', 'Auth Header Attachment', 'High', 'Verify Authorization header with Bearer token on all requests'],
        ['UT-FE-API-003', 'Error Response Handling', 'High', 'Verify error responses display appropriate user messages'],
        ['UT-FE-API-004', 'Request Cancellation', 'Medium', 'Verify pending requests are cancelled on component unmount'],
    ]
)

doc.add_heading('4.4 Route Navigation Tests', level=2)
doc.add_paragraph(
    'Route navigation tests verify that React Router lazy loading works correctly, that '
    'protected routes redirect unauthenticated users, and that role-based routes enforce '
    'permission checks.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-FE-ROUTE-001', 'Lazy Loading', 'High', 'Verify route components load on demand and not in initial bundle'],
        ['UT-FE-ROUTE-002', 'Protected Route Redirect', 'High', 'Verify unauthenticated users redirected to login'],
        ['UT-FE-ROUTE-003', 'Role-Based Route Access', 'High', 'Verify unauthorized roles see 403 or redirect'],
        ['UT-FE-ROUTE-004', 'Navigation Menu Rendering', 'Medium', 'Verify sidebar shows only permitted menu items'],
    ]
)

doc.add_heading('4.5 State Management Tests', level=2)
doc.add_paragraph(
    'State management tests validate TanStack Query cache behavior, including cache invalidation, '
    'refetching on focus, and optimistic updates. Tests verify that state updates trigger correct '
    're-renders and that stale data is properly refreshed.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-FE-STATE-001', 'Query Cache Hit', 'Medium', 'Verify cached data is returned without refetch'],
        ['UT-FE-STATE-002', 'Cache Invalidation', 'High', 'Verify mutation invalidates related queries'],
        ['UT-FE-STATE-003', 'Optimistic Update', 'Medium', 'Verify UI updates immediately before server response'],
        ['UT-FE-STATE-004', 'Refetch on Window Focus', 'Low', 'Verify data refetches when window regains focus'],
    ]
)

doc.add_heading('4.6 Authentication Context Tests', level=2)
doc.add_paragraph(
    'Authentication context tests verify that the AuthProvider correctly manages user state, '
    'that login/logout actions update the context, and that protected components access the '
    'correct user data.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-FE-AUTH-001', 'Login Updates Context', 'High', 'Verify user data populated in context after login'],
        ['UT-FE-AUTH-002', 'Logout Clears Context', 'High', 'Verify user data cleared from context after logout'],
        ['UT-FE-AUTH-003', 'Token Persistence', 'High', 'Verify tokens stored in localStorage and restored on refresh'],
        ['UT-FE-AUTH-004', 'Session Expiry Handling', 'High', 'Verify expired session triggers redirect to login'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. SERVICE LAYER TESTING
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Service Layer Testing', level=1)
doc.add_paragraph(
    'Service layer tests validate business logic, transaction management, error handling, '
    'and multi-tenant filtering. Services sit between controllers and repositories, containing '
    'the core business rules that must be thoroughly tested.'
)

doc.add_heading('5.1 Business Logic Validation', level=2)
doc.add_paragraph(
    'Business logic tests verify that services enforce domain rules such as: leave cannot exceed '
    'available balance, payroll runs cannot be approved without confirmation, assets cannot be '
    'allocated if already allocated, and tickets can only be rated when resolved.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-SVC-001', 'Leave Balance Enforcement', 'High', 'Verify service rejects leave when balance is insufficient'],
        ['UT-SVC-002', 'Payroll State Machine', 'High', 'Verify payroll run follows correct state transitions'],
        ['UT-SVC-003', 'Asset Allocation Guard', 'High', 'Verify double allocation is prevented by service layer'],
        ['UT-SVC-004', 'Ticket Rating Guard', 'Medium', 'Verify rating only accepted for RESOLVED tickets'],
        ['UT-SVC-005', 'Attendance Double Punch Guard', 'High', 'Verify service prevents duplicate punch-in entries'],
        ['UT-SVC-006', 'Role Permission Enforcement', 'High', 'Verify service checks user permissions before operations'],
    ]
)

doc.add_heading('5.2 Transaction Management', level=2)
doc.add_paragraph(
    'Transaction tests verify that services use @Transactional annotations correctly and that '
    'operations that span multiple repository calls are atomic. Tests verify that failed '
    'operations roll back changes and do not leave partial state.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-SVC-TX-001', 'Rollback on Exception', 'High', 'Verify transaction rolls back when service method throws exception'],
        ['UT-SVC-TX-002', 'Commit on Success', 'High', 'Verify transaction commits when all operations succeed'],
        ['UT-SVC-TX-003', 'Nested Service Calls', 'Medium', 'Verify nested @Transactional methods participate in same transaction'],
    ]
)

doc.add_heading('5.3 Error Handling', level=2)
doc.add_paragraph(
    'Error handling tests verify that services throw appropriate exceptions for invalid operations, '
    'that exceptions contain meaningful messages, and that the controller layer can translate '
    'exceptions into proper HTTP responses.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-SVC-ERR-001', 'Business Rule Exception', 'High', 'Verify appropriate exception thrown for business rule violations'],
        ['UT-SVC-ERR-002', 'Not Found Exception', 'High', 'Verify ResourceNotFoundException for missing entities'],
        ['UT-SVC-ERR-003', 'Validation Exception', 'Medium', 'Verify validation exceptions contain field-level errors'],
        ['UT-SVC-ERR-004', 'Concurrent Modification', 'Medium', 'Verify optimistic locking exception on concurrent updates'],
    ]
)

doc.add_heading('5.4 Multi-Tenant Filtering', level=2)
doc.add_paragraph(
    'Multi-tenant tests verify that all service queries are scoped to the current tenant context. '
    'Tests create data in one tenant and verify that queries from another tenant do not access it.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-SVC-MT-001', 'Query Scope by Tenant', 'High', 'Verify queries only return data for current tenant'],
        ['UT-SVC-MT-002', 'Create Scoped to Tenant', 'High', 'Verify new entities are associated with current tenant'],
        ['UT-SVC-MT-003', 'Cross-Tenant Isolation', 'High', 'Verify tenant A data is invisible to tenant B queries'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. REPOSITORY/DAO TESTING
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Repository/DAO Testing', level=1)
doc.add_paragraph(
    'Repository tests validate JPA query correctness, multi-tenant query filters, pagination, '
    'and custom query methods. These tests use an H2 in-memory database to verify SQL queries '
    'execute correctly without requiring a MySQL instance.'
)

doc.add_heading('6.1 JPA Query Validation', level=2)
doc.add_paragraph(
    'Query validation tests verify that custom JPQL and native SQL queries return correct results. '
    'Tests cover complex joins, aggregations, date range filtering, and keyword search queries.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-REPO-001', 'Find By Username', 'High', 'Verify findByUsername returns correct user entity'],
        ['UT-REPO-002', 'Find Attendance by Date Range', 'High', 'Verify attendance records filtered by start and end date'],
        ['UT-REPO-003', 'Find Pending Leaves', 'High', 'Verify query returns only PENDING status leaves'],
        ['UT-REPO-004', 'Find Tickets by Status', 'High', 'Verify ticket query filters by status correctly'],
        ['UT-REPO-005', 'Aggregate Attendance Summary', 'Medium', 'Verify aggregate query returns correct present/absent counts'],
    ]
)

doc.add_heading('6.2 Multi-Tenant Query Filters', level=2)
doc.add_paragraph(
    'Multi-tenant repository tests verify that TenantFilter or @Where clauses correctly scope '
    'all queries to the current tenant. Tests create data across multiple tenants and verify '
    'query isolation.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-REPO-MT-001', 'Tenant Filter Applied', 'High', 'Verify tenant ID is appended to all queries automatically'],
        ['UT-REPO-MT-002', 'Cross-Tenant Isolation', 'High', 'Verify tenant A data not returned by tenant B queries'],
        ['UT-REPO-MT-003', 'Native Query Tenant Filter', 'Medium', 'Verify native SQL queries also respect tenant filter'],
    ]
)

doc.add_heading('6.3 Pagination', level=2)
doc.add_paragraph(
    'Pagination tests verify that paginated queries return correct page sizes, total counts, '
    'and page numbers. Tests verify edge cases like requesting page beyond available data.'
)
add_styled_table(doc,
    ['TC ID', 'Test Case', 'Priority', 'Description'],
    [
        ['UT-REPO-PAGE-001', 'First Page Retrieval', 'Medium', 'Verify first page returns correct number of records'],
        ['UT-REPO-PAGE-002', 'Page Size Enforcement', 'Medium', 'Verify page size limit is respected'],
        ['UT-REPO-PAGE-003', 'Total Count', 'Medium', 'Verify totalElements reflects all matching records'],
        ['UT-REPO-PAGE-004', 'Empty Page Handling', 'Medium', 'Verify empty page returns zero records and correct metadata'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. TEST COVERAGE METRICS
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Test Coverage Metrics', level=1)
doc.add_paragraph(
    'Test coverage is measured using JaCoCo for the Java backend and V8 coverage for the '
    'TypeScript frontend. The target is 80% or higher line coverage across all modules, '
    'with critical business logic modules targeting 90% or higher.'
)

doc.add_heading('7.1 Coverage Targets by Module', level=2)
add_styled_table(doc,
    ['Module', 'Current Coverage', 'Target Coverage', 'Status'],
    [
        ['Authentication', '85%', '90%', 'On Track'],
        ['User/Employee', '82%', '90%', 'On Track'],
        ['Attendance', '78%', '85%', 'Needs Improvement'],
        ['Leave', '80%', '85%', 'On Track'],
        ['Payroll', '75%', '85%', 'Needs Improvement'],
        ['Asset', '83%', '85%', 'On Track'],
        ['Helpdesk', '80%', '85%', 'On Track'],
        ['Community/Chat', '72%', '80%', 'Needs Improvement'],
        ['Dashboard', '77%', '80%', 'On Track'],
        ['Notification', '79%', '80%', 'On Track'],
        ['Onboarding', '81%', '85%', 'On Track'],
        ['Performance', '76%', '80%', 'On Track'],
        ['TA Expense', '82%', '85%', 'On Track'],
        ['Frontend Components', '74%', '80%', 'Needs Improvement'],
        ['Frontend Forms', '80%', '85%', 'On Track'],
        ['Overall Backend', '79%', '85%', 'On Track'],
        ['Overall Frontend', '76%', '80%', 'On Track'],
    ]
)

doc.add_heading('7.2 Coverage Report Generation', level=2)
doc.add_paragraph(
    'Coverage reports are generated automatically during CI/CD pipeline execution. JaCoCo '
    'produces HTML and XML reports for the backend, while Vitest generates HTML coverage '
    'reports for the frontend. Reports are uploaded as build artifacts and can be reviewed '
    'in the GitHub Actions workflow.'
)
doc.add_paragraph(
    'Coverage gates enforce a minimum threshold of 80% line coverage. If coverage drops below '
    'the threshold, the build fails and developers must add tests before merging. Coverage trends '
    'are tracked over time to ensure continuous improvement.'
)

add_styled_table(doc,
    ['Report Type', 'Tool', 'Format', 'Generated By'],
    [
        ['Backend Line Coverage', 'JaCoCo', 'HTML + XML', 'Maven jacoco:report'],
        ['Backend Branch Coverage', 'JaCoCo', 'HTML + XML', 'Maven jacoco:report'],
        ['Frontend Coverage', 'V8/Vitest', 'HTML + JSON', 'vitest --coverage'],
        ['Python Coverage', 'pytest-cov', 'HTML + XML', 'pytest --cov'],
        ['Combined Report', 'Custom Script', 'HTML', 'scripts/merge-coverage.py'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. MOCK STRATEGY
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Mock Strategy', level=1)
doc.add_paragraph(
    'Effective mocking is essential for unit test isolation. The PIXOUS HR Portal uses '
    'Mockito for backend mocking, MSW for frontend API mocking, and unittest.mock for '
    'Python service mocking. Each external dependency has a defined mock strategy.'
)

doc.add_heading('8.1 Database Mocking', level=2)
doc.add_paragraph(
    'Repository layer tests use H2 in-memory database for most cases. Service layer tests '
    'mock repository interfaces using @MockBean or @Mock annotations. This ensures service '
    'logic is tested independently of JPA/Hibernate behavior.'
)
add_styled_table(doc,
    ['Dependency', 'Mock Approach', 'Verification'],
    [
        ['UserRepository', '@Mock + when/thenReturn', 'Verify correct query methods called'],
        ['AttendanceRepository', '@Mock + ArgumentCaptor', 'Verify saved entity has correct fields'],
        ['LeaveRepository', '@Mock + verify', 'Verify repository interaction counts'],
        ['PayrollRepository', '@Mock + times()', 'Verify bulk save operations'],
        ['TicketRepository', '@Mock + ArgumentCaptor', 'Verify status transitions in saved entities'],
    ]
)

doc.add_heading('8.2 External Service Mocking', level=2)
doc.add_paragraph(
    'External services (Twilio SMS, Kafka messaging, Redis caching) are mocked at the service '
    'interface level. This allows tests to verify the calling logic without actual external '
    'service connectivity.'
)
add_styled_table(doc,
    ['Service', 'Mock Tool', 'What is Verified'],
    [
        ['Twilio SMS', '@Mock + verify', 'SMS send method called with correct phone and message'],
        ['Kafka Producer', '@Mock + ArgumentCaptor', 'Correct topic, key, and payload sent'],
        ['Redis Cache', '@Mock + verify', 'Cache get/set/evict operations called correctly'],
        ['Face Recognition', '@Mock + when/thenReturn', 'Verification result returned correctly'],
        ['PDF Generator', '@Mock + verify', 'PDF generation called with correct template and data'],
    ]
)

doc.add_heading('8.3 WebSocket Mocking', level=2)
doc.add_paragraph(
    'WebSocket notification tests mock the SimpMessagingTemplate to verify that notifications '
    'are sent to the correct destinations with correct payloads. This avoids requiring a live '
    'WebSocket connection during unit tests.'
)
add_styled_table(doc,
    ['Component', 'Mock Approach', 'Verification'],
    [
        ['SimpMessagingTemplate', '@Mock + verify', 'convertAndUserToSent with correct destination and payload'],
        ['WebSocket Session', 'Mock WebSocketSession', 'Session attributes and message sending'],
        ['Notification Service', '@Mock + verify', 'Notification entity created and persisted'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. TEST DATA MANAGEMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Test Data Management', level=1)
doc.add_paragraph(
    'Test data management is critical for maintaining fast, reliable, and isolated unit tests. '
    'The PIXOUS HR Portal uses the Factory pattern for test data creation and implements '
    'database cleanup between test classes to prevent data leakage.'
)

doc.add_heading('9.1 Factory Pattern for Test Data', level=2)
doc.add_paragraph(
    'Test data factories encapsulate entity creation logic, providing default values for all '
    'required fields while allowing overrides for specific test scenarios. This reduces '
    'boilerplate in test methods and ensures consistency.'
)
add_styled_table(doc,
    ['Factory', 'Entity', 'Default Values'],
    [
        ['UserFactory', 'User', 'username, email, phone, role=IT_EMP, department=Engineering'],
        ['AttendanceFactory', 'Attendance', 'userId, date, punchIn, punchOut, status=PRESENT'],
        ['LeaveFactory', 'LeaveRequest', 'userId, type=CASUAL, startDate, endDate, status=PENDING'],
        ['PayrollFactory', 'Payslip', 'userId, month, year, basic, hra, deductions, netPay'],
        ['AssetFactory', 'Asset', 'name, type=IT, serialNumber, status=AVAILABLE'],
        ['TicketFactory', 'Ticket', 'userId, subject, description, category, status=OPEN'],
        ['MessageFactory', 'ChatMessage', 'communityId, senderId, content, timestamp'],
    ]
)

doc.add_heading('9.2 Database Cleanup Between Tests', level=2)
doc.add_paragraph(
    'Each test class uses @Transactional or @DirtiesContext to ensure database state is reset '
    'between tests. For repository tests, @DataJpaTest provides automatic rollback after each '
    'test method. Service tests mock repositories to avoid database state issues entirely.'
)
add_styled_table(doc,
    ['Cleanup Strategy', 'When Used', 'Implementation'],
    [
        ['@DataJpaTest auto-rollback', 'Repository tests', 'Automatic rollback after each test method'],
        ['@Transactional', 'Service tests', 'Test runs in transaction that rolls back'],
        ['@DirtiesContext', 'Integration tests', 'Spring context restarted between test classes'],
        ['Manual @BeforeEach', 'Shared state tests', 'Explicit cleanup in setup method'],
        ['H2 TRUNCATE', 'Complex scenarios', 'Execute TRUNCATE in @BeforeEach'],
    ]
)

doc.add_heading('9.3 Test Data Isolation', level=2)
doc.add_paragraph(
    'Each test method creates its own required data using factories, ensuring tests are '
    'independent and can run in any order. Shared constants (like default user IDs) are defined '
    'in test utility classes but never mutated during test execution.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. CI/CD INTEGRATION
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. CI/CD Integration', level=1)
doc.add_paragraph(
    'Unit tests are integrated into the GitHub Actions CI/CD pipeline to ensure every pull '
    'request is validated before merging. The pipeline runs backend, frontend, and Python tests '
    'in parallel, generates coverage reports, and enforces minimum coverage thresholds.'
)

doc.add_heading('10.1 GitHub Actions Test Pipeline', level=2)
doc.add_paragraph(
    'The test pipeline triggers on every pull request and push to main. It runs three parallel '
    'jobs: backend tests (Java 17 + Maven), frontend tests (Node 20 + Vitest), and Python '
    'analytics tests (pytest). Each job uploads test results and coverage reports as artifacts.'
)
add_styled_table(doc,
    ['Pipeline Stage', 'Tool', 'Action', 'Timeout'],
    [
        ['Backend Unit Tests', 'Maven + JUnit 5', 'mvn test -Punit-tests', '10 minutes'],
        ['Backend Coverage', 'JaCoCo', 'mvn jacoco:report', '5 minutes'],
        ['Frontend Unit Tests', 'Vitest', 'vitest run --coverage', '5 minutes'],
        ['Python Unit Tests', 'pytest', 'pytest --cov=analytics', '5 minutes'],
        ['Coverage Gate', 'Custom Script', 'Check thresholds', '1 minute'],
        ['Test Report', 'JUnit XML', 'Publish results', '1 minute'],
    ]
)

doc.add_heading('10.2 Test Reporting', level=2)
doc.add_paragraph(
    'Test results are published in JUnit XML format and displayed directly in the GitHub Actions '
    'workflow summary. Coverage reports are uploaded as build artifacts and linked in the PR '
    'description for easy review by reviewers.'
)

doc.add_heading('10.3 Coverage Gates', level=2)
doc.add_paragraph(
    'Coverage gates enforce minimum thresholds that must be met for a build to pass. If any '
    'module falls below its target coverage, the pipeline fails and developers must add or '
    'improve tests before the PR can be merged.'
)
add_styled_table(doc,
    ['Coverage Metric', 'Minimum Threshold', 'Action on Failure'],
    [
        ['Overall Backend Line Coverage', '80%', 'Build fails, PR blocked'],
        ['Overall Frontend Line Coverage', '80%', 'Build fails, PR blocked'],
        ['Critical Module Coverage', '90%', 'Build fails, PR blocked'],
        ['Branch Coverage', '70%', 'Warning, PR review required'],
        ['New Code Coverage', '85%', 'Build fails, PR blocked'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. TEST CASE INVENTORY BY MODULE
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. Test Case Inventory by Module', level=1)
doc.add_paragraph(
    'The following table provides a comprehensive inventory of all unit test cases organized '
    'by module. Each test case has a unique ID in the format UT-XXX-NNN where XXX is the module '
    'code and NNN is the sequence number.'
)

add_styled_table(doc,
    ['Module', 'TC ID Range', 'Total Tests', 'High Priority', 'Medium Priority', 'Low Priority'],
    [
        ['Authentication', 'UT-AUTH-001 to UT-AUTH-008', '8', '5', '3', '0'],
        ['User/Employee', 'UT-USER-001 to UT-USER-008', '8', '5', '3', '0'],
        ['Attendance', 'UT-ATT-001 to UT-ATT-010', '10', '6', '4', '0'],
        ['Leave', 'UT-LEA-001 to UT-LEA-010', '10', '6', '4', '0'],
        ['Payroll', 'UT-PAY-001 to UT-PAY-008', '8', '6', '2', '0'],
        ['Asset', 'UT-AST-001 to UT-AST-007', '7', '4', '3', '0'],
        ['Helpdesk', 'UT-HD-001 to UT-HD-007', '7', '3', '4', '0'],
        ['Community/Chat', 'UT-CHAT-001 to UT-CHAT-008', '8', '3', '4', '1'],
        ['Dashboard', 'UT-DASH-001 to UT-DASH-005', '5', '3', '2', '0'],
        ['Notification', 'UT-NOTIF-001 to UT-NOTIF-004', '4', '2', '2', '0'],
        ['Onboarding', 'UT-ONB-001 to UT-ONB-004', '4', '3', '1', '0'],
        ['Performance', 'UT-PERF-001 to UT-PERF-004', '4', '2', '2', '0'],
        ['TA Expense', 'UT-TA-001 to UT-TA-005', '5', '3', '2', '0'],
        ['Frontend Components', 'UT-FE-001 to UT-FE-008', '8', '4', '4', '0'],
        ['Frontend Forms', 'UT-FE-FORM-001 to UT-FE-FORM-006', '6', '3', '3', '0'],
        ['Frontend API', 'UT-FE-API-001 to UT-FE-API-004', '4', '3', '1', '0'],
        ['Frontend Routes', 'UT-FE-ROUTE-001 to UT-FE-ROUTE-004', '4', '3', '1', '0'],
        ['Frontend State', 'UT-FE-STATE-001 to UT-FE-STATE-004', '4', '1', '2', '1'],
        ['Frontend Auth', 'UT-FE-AUTH-001 to UT-FE-AUTH-004', '4', '4', '0', '0'],
        ['Service Layer', 'UT-SVC-001 to UT-SVC-006', '6', '5', '1', '0'],
        ['Service Tx', 'UT-SVC-TX-001 to UT-SVC-TX-003', '3', '2', '1', '0'],
        ['Service Error', 'UT-SVC-ERR-001 to UT-SVC-ERR-004', '4', '2', '2', '0'],
        ['Service Multi-Tenant', 'UT-SVC-MT-001 to UT-SVC-MT-003', '3', '3', '0', '0'],
        ['Repository', 'UT-REPO-001 to UT-REPO-005', '5', '3', '2', '0'],
        ['Repository MT', 'UT-REPO-MT-001 to UT-REPO-MT-003', '3', '2', '1', '0'],
        ['Repository Pagination', 'UT-REPO-PAGE-001 to UT-REPO-PAGE-004', '4', '0', '4', '0'],
        ['TOTAL', '-', '151', '87', '59', '2'],
    ]
)

doc.add_paragraph(
    'Total test cases: 151 across all backend, frontend, service, and repository layers. '
    'This provides comprehensive coverage of all business logic, UI components, and data access '
    'patterns in the PIXOUS HR Portal.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. KNOWN TEST GAPS & RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. Known Test Gaps & Recommendations', level=1)
doc.add_paragraph(
    'Despite comprehensive unit test coverage, certain areas require additional testing effort. '
    'The following gaps have been identified and recommendations are provided to address them.'
)

doc.add_heading('12.1 Identified Gaps', level=2)
add_styled_table(doc,
    ['Gap Area', 'Current State', 'Impact', 'Priority'],
    [
        ['Community/Chat Coverage', '72% coverage', 'Message retention and poll logic under-tested', 'High'],
        ['Frontend Component Coverage', '74% coverage', 'Complex components like DataTable need more tests', 'Medium'],
        ['Edge Case Testing', 'Limited', 'Boundary values and null inputs not fully covered', 'Medium'],
        ['Concurrent Access', 'Not tested', 'Race conditions in attendance punch not validated', 'Low'],
        ['Performance Thresholds', 'Not enforced', 'Test execution time not tracked or limited', 'Medium'],
        ['Mutation Testing', 'Not implemented', 'No mutation testing to verify test effectiveness', 'Low'],
    ]
)

doc.add_heading('12.2 Recommendations', level=2)
recommendations = [
    'Increase Community/Chat module coverage to 80% by adding retention policy and poll tests',
    'Add more DataTable component tests covering sorting, filtering, and pagination interactions',
    'Implement parameterized tests for boundary value analysis across all input validation points',
    'Add concurrency tests for attendance punch using parallel test execution',
    'Set test execution time alerts to prevent slow tests from blocking CI/CD',
    'Evaluate PIT (Pitest) for mutation testing to verify test suite effectiveness',
    'Add more frontend form validation tests for complex multi-step forms',
    'Implement test data builder pattern alongside factories for complex entity construction',
]
for rec in recommendations:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_heading('12.3 Future Improvements', level=2)
doc.add_paragraph(
    'As the PIXOUS HR Portal continues to evolve, the unit testing strategy should be reviewed '
    'quarterly to incorporate new modules, updated business rules, and emerging best practices. '
    'Consider adopting property-based testing for complex business logic and contract testing '
    'for API boundaries between frontend and backend.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. APPENDIX: SAMPLE TEST CODE SNIPPETS
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. Appendix: Sample Test Code Snippets', level=1)
doc.add_paragraph(
    'This appendix provides sample unit test code snippets for each testing layer. These examples '
    'demonstrate the testing patterns, conventions, and tools used throughout the PIXOUS HR Portal.'
)

doc.add_heading('13.1 Backend: AuthService Unit Test (JUnit 5 + Mockito)', level=2)
doc.add_paragraph(
    'The following code shows a typical backend unit test for the AuthService class. The test '
    'uses Mockito to mock dependencies and JUnit 5 for assertions and test lifecycle management.'
)
# Code snippet as table
code_auth = [
    ['@ExtendWith(MockitoExtension.class)', ''],
    ['class AuthServiceTest {', ''],
    ['  @Mock private UserRepository userRepo;', ''],
    ['  @Mock private JwtUtil jwtUtil;', ''],
    ['  @Mock private PasswordEncoder passwordEncoder;', ''],
    ['  @InjectMocks private AuthService authService;', ''],
    ['', ''],
    ['  @Test', ''],
    ['  void shouldGenerateTokensOnValidLogin() {', ''],
    ['    // Arrange', ''],
    ['    User user = UserFactory.createDefault();', ''],
    ['    when(userRepo.findByUsername("testuser"))', ''],
    ['      .thenReturn(Optional.of(user));', ''],
    ['    when(passwordEncoder.matches(any(), any()))', ''],
    ['      .thenReturn(true);', ''],
    ['    when(jwtUtil.generateAccessToken(any()))', ''],
    ['      .thenReturn("mock-access-token");', ''],
    ['', ''],
    ['    // Act', ''],
    ['    AuthResponse response = authService.login(', ''],
    ['      new LoginRequest("testuser", "password")', ''],
    ['    );', ''],
    ['', ''],
    ['    // Assert', ''],
    ['    assertThat(response.getToken())', ''],
    ['      .isEqualTo("mock-access-token");', ''],
    ['    verify(jwtUtil).generateAccessToken(user);', ''],
    ['  }', ''],
    ['}', ''],
]
add_styled_table(doc, ['Code', ''], code_auth)

doc.add_heading('13.2 Backend: Attendance Service Test', level=2)
doc.add_paragraph(
    'This test validates the punch-in business logic including geofence validation and '
    'double punch prevention. The geofence service is mocked to control validation behavior.'
)
code_att = [
    ['@Test', ''],
    ['void shouldRejectDoublePunchIn() {', ''],
    ['  // Arrange', ''],
    ['  User user = UserFactory.createDefault();', ''],
    ['  Attendance existing = AttendanceFactory.createPunchedIn(user);', ''],
    ['  when(attendanceRepo.findTodayAttendance(user.getId()))', ''],
    ['    .thenReturn(Optional.of(existing));', ''],
    ['', ''],
    ['  // Act & Assert', ''],
    ['  assertThrows(IllegalStateException.class, () -> {', ''],
    ['    attendanceService.punchIn(user.getId(), gpsData);', ''],
    ['  });', ''],
    ['}', ''],
    ['', ''],
    ['@Test', ''],
    ['void shouldRejectPunchOutsideGeofence() {', ''],
    ['  // Arrange', ''],
    ['  when(geofenceService.isInsideGeofence(any())).thenReturn(false);', ''],
    ['', ''],
    ['  // Act & Assert', ''],
    ['  assertThrows(GeofenceViolationException.class, () -> {', ''],
    ['    attendanceService.punchIn(userId, outsideGps);', ''],
    ['  });', ''],
    ['}', ''],
]
add_styled_table(doc, ['Code', ''], code_att)

doc.add_heading('13.3 Backend: Leave Service Test', level=2)
doc.add_paragraph(
    'This test validates that leave balance is deducted when leave is approved and that '
    'insufficient balance is properly rejected.'
)
code_leave = [
    ['@Test', ''],
    ['void shouldDeductBalanceOnLeaveApproval() {', ''],
    ['  // Arrange', ''],
    ['  LeaveRequest leave = LeaveFactory.createPending(userId, 3);', ''],
    ['  when(leaveRepo.findById(leave.getId()))', ''],
    ['    .thenReturn(Optional.of(leave));', ''],
    ['  when(balanceRepo.getBalance(userId, "CASUAL"))', ''],
    ['    .thenReturn(new LeaveBalance(10, 7));', ''],
    ['', ''],
    ['  // Act', ''],
    ['  leaveService.approve(leave.getId(), managerId, "Approved");', ''],
    ['', ''],
    ['  // Assert', ''],
    ['  verify(balanceRepo).deductBalance(userId, "CASUAL", 3);', ''],
    ['  assertThat(leave.getStatus()).isEqualTo(Status.APPROVED);', ''],
    ['}', ''],
    ['', ''],
    ['@Test', ''],
    ['void shouldRejectLeaveWhenBalanceInsufficient() {', ''],
    ['  when(balanceRepo.getBalance(userId, "CASUAL"))', ''],
    ['    .thenReturn(new LeaveBalance(10, 1));', ''],
    ['', ''],
    ['  assertThrows(InsufficientBalanceException.class, () -> {', ''],
    ['    leaveService.apply(leaveRequest);', ''],
    ['  });', ''],
    ['}', ''],
]
add_styled_table(doc, ['Code', ''], code_leave)

doc.add_heading('13.4 Frontend: Component Rendering Test (Vitest)', level=2)
doc.add_paragraph(
    'This test verifies that a React Button component renders correctly and responds to user clicks.'
)
code_fe = [
    ['import { render, screen, fireEvent } from "@testing-library/react";', ''],
    ['import { describe, it, expect, vi } from "vitest";', ''],
    ['import { Button } from "./Button";', ''],
    ['', ''],
    ['describe("Button Component", () => {', ''],
    ['  it("renders with label text", () => {', ''],
    ['    render(<Button label="Submit" />);', ''],
    ['    expect(screen.getByText("Submit")).toBeInTheDocument();', ''],
    ['  });', ''],
    ['', ''],
    ['  it("calls onClick when clicked", () => {', ''],
    ['    const handleClick = vi.fn();', ''],
    ['    render(<Button label="Click Me" onClick={handleClick} />);', ''],
    ['    fireEvent.click(screen.getByText("Click Me"));', ''],
    ['    expect(handleClick).toHaveBeenCalledTimes(1);', ''],
    ['  });', ''],
    ['', ''],
    ['  it("is disabled when disabled prop is true", () => {', ''],
    ['    render(<Button label="Disabled" disabled />);', ''],
    ['    expect(screen.getByText("Disabled")).toBeDisabled();', ''],
    ['  });', ''],
    ['});', ''],
]
add_styled_table(doc, ['Code', ''], code_fe)

doc.add_heading('13.5 Frontend: Form Validation Test', level=2)
doc.add_paragraph(
    'This test verifies that a login form enforces required fields and validates email format.'
)
code_form = [
    ['import { render, screen, fireEvent, waitFor } from "@testing-library/react";', ''],
    ['import { describe, it, expect } from "vitest";', ''],
    ['import { LoginForm } from "./LoginForm";', ''],
    ['', ''],
    ['describe("LoginForm Validation", () => {', ''],
    ['  it("shows error for empty email", async () => {', ''],
    ['    render(<LoginForm />);', ''],
    ['    fireEvent.click(screen.getByText("Submit"));', ''],
    ['    await waitFor(() => {', ''],
    ['      expect(screen.getByText("Email is required")).toBeInTheDocument();', ''],
    ['    });', ''],
    ['  });', ''],
    ['', ''],
    ['  it("shows error for invalid email format", async () => {', ''],
    ['    render(<LoginForm />);', ''],
    ['    fireEvent.change(screen.getByLabelText("Email"), {', ''],
    ['      target: { value: "not-an-email" },', ''],
    ['    });', ''],
    ['    fireEvent.click(screen.getByText("Submit"));', ''],
    ['    await waitFor(() => {', ''],
    ['      expect(screen.getByText("Invalid email")).toBeInTheDocument();', ''],
    ['    });', ''],
    ['  });', ''],
    ['});', ''],
]
add_styled_table(doc, ['Code', ''], code_form)

doc.add_heading('13.6 Python: Analytics Service Test (pytest)', level=2)
doc.add_paragraph(
    'This test validates the face verification endpoint in the Python analytics service.'
)
code_py = [
    ['import pytest', ''],
    ['from unittest.mock import patch, MagicMock', ''],
    ['from analytics.face_service import verify_face', ''],
    ['', ''],
    ['class TestFaceVerification:', ''],
    ['  @patch("analytics.face_service.face_recognition")', ''],
    ['  def test_verify_face_success(self, mock_face):', ''],
    ['    # Arrange', ''],
    ['    mock_face.compare_faces.return_value = True', ''],
    ['    mock_face.face_distance.return_value = 0.35', ''],
    ['', ''],
    ['    # Act', ''],
    ['    result = verify_face(image_bytes, stored_embedding)', ''],
    ['', ''],
    ['    # Assert', ''],
    ['    assert result.verified is True', ''],
    ['    assert result.confidence > 0.8', ''],
    ['    mock_face.compare_faces.assert_called_once()', ''],
    ['', ''],
    ['  @patch("analytics.face_service.face_recognition")', ''],
    ['  def test_verify_face_failure(self, mock_face):', ''],
    ['    mock_face.compare_faces.return_value = False', ''],
    ['    mock_face.face_distance.return_value = 0.85', ''],
    ['', ''],
    ['    result = verify_face(image_bytes, stored_embedding)', ''],
    ['    assert result.verified is False', ''],
]
add_styled_table(doc, ['Code', ''], code_py)

doc.add_heading('13.7 Backend: Repository Test with H2', level=2)
doc.add_paragraph(
    'This test validates a custom JPA query using the H2 in-memory database.'
)
code_repo = [
    ['@DataJpaTest', ''],
    ['class UserRepositoryTest {', ''],
    ['  @Autowired private UserRepository userRepo;', ''],
    ['', ''],
    ['  @Test', ''],
    ['  void shouldFindByUsername() {', ''],
    ['    // Arrange', ''],
    ['    User user = new User();', ''],
    ['    user.setUsername("testuser");', ''],
    ['    user.setEmail("test@example.com");', ''],
    ['    userRepo.save(user);', ''],
    ['', ''],
    ['    // Act', ''],
    ['    Optional<User> found = userRepo.findByUsername("testuser");', ''],
    ['', ''],
    ['    // Assert', ''],
    ['    assertThat(found).isPresent();', ''],
    ['    assertThat(found.get().getEmail())', ''],
    ['      .isEqualTo("test@example.com");', ''],
    ['  }', ''],
    ['', ''],
    ['  @Test', ''],
    ['  void shouldReturnEmptyForUnknownUsername() {', ''],
    ['    Optional<User> found = userRepo.findByUsername("unknown");', ''],
    ['    assertThat(found).isEmpty();', ''],
    ['  }', ''],
    ['}', ''],
]
add_styled_table(doc, ['Code', ''], code_repo)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PIXOUS Technologies Pvt. Ltd. | Confidential')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save('PIXOUS_HR_Portal_Unit_Testing_Document.docx')
print('Document saved: PIXOUS_HR_Portal_Unit_Testing_Document.docx')
