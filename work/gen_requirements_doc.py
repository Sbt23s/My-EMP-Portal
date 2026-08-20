"""Generate PIXOUS HR Portal - Requirements Document"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

DARK_BLUE = RGBColor(0, 51, 102)
MED_BLUE = RGBColor(0, 102, 153)
GREY = RGBColor(128, 128, 128)

def add_logo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('pixous_logo.png', width=Inches(2.0))

def add_styled_table(doc, headers, rows_data):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): '003366', qn('w:val'): 'clear'})
        shading.append(shd)
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE if level <= 1 else MED_BLUE
    return h

def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(10)
    return p


# ═══ COVER PAGE ═══
add_logo(doc)
doc.add_paragraph()
title = doc.add_heading('PIXOUS HR Portal', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = DARK_BLUE

subtitle = doc.add_heading('Comprehensive Requirements Document', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = MED_BLUE

doc.add_paragraph()
t = doc.add_table(rows=6, cols=2)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ('Document Type', 'Software Requirements Specification (SRS)'),
    ('Version', '1.0'),
    ('Date', datetime.date.today().strftime('%B %d, %Y')),
    ('Prepared By', 'PIXOUS Technologies - Product & Engineering Team'),
    ('Reviewed By', 'PIXOUS Technologies - Quality Assurance Team'),
    ('Classification', 'Confidential - Client Deliverable'),
]
for i, (k, v) in enumerate(info):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
    for p in t.rows[i].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)
doc.add_page_break()

# ═══ REVISION HISTORY ═══
add_heading(doc, 'Revision History', level=1)
rev = [
    ['1.0', datetime.date.today().strftime('%Y-%m-%d'), 'PIXOUS Technologies', 'Initial draft of the SRS document'],
]
add_styled_table(doc, ['Version', 'Date', 'Author', 'Description'], rev)
doc.add_page_break()

# ═══ TABLE OF CONTENTS ═══
add_heading(doc, 'Table of Contents', level=1)
toc = [
    '1. Executive Summary',
    '2. Project Scope',
    '3. Business Requirements',
    '4. Functional Requirements',
    '5. Non-Functional Requirements',
    '6. Technical Requirements',
    '7. User Role Requirements',
    '8. Module Requirements',
    '9. Integration Requirements',
    '10. Deployment Requirements',
    '11. Acceptance Criteria',
    '12. Traceability Matrix',
    '13. Glossary',
    '14. Appendix',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)
doc.add_page_break()


# ═══ 1. EXECUTIVE SUMMARY ═══
add_heading(doc, '1. Executive Summary', level=1)

add_para(doc,
    'This document defines the comprehensive software requirements for the PIXOUS HR Portal, '
    'a full-featured Human Resource Management System (HRMS) designed to serve organizations across '
    'multiple industries including Information Technology (IT) and Civil/Construction sectors.')

add_para(doc,
    'The PIXOUS HR Portal is a multi-tenant SaaS platform that provides end-to-end employee lifecycle '
    'management, from onboarding through offboarding. The system supports role-based access control with '
    'over 14 distinct permission codes, ensuring that each user role has precisely calibrated access to '
    'system features and organizational data.')

add_para(doc,
    'The platform is built on a modern technology stack comprising Spring Boot 3.2 for the backend REST API, '
    'React with TypeScript for the web frontend, React Native for Android mobile applications, and Flutter '
    'for iOS mobile applications. The database layer utilizes MySQL 8.4 with Flyway for schema migration '
    'management (101 migrations to date), while Redis provides optional caching with in-memory fallback.')

add_para(doc,
    'Key capabilities include GPS and face-recognition attendance tracking, leave management with multi-level '
    'approval workflows, payroll processing with automated PDF payslip generation, asset management with QR code '
    'tracking, helpdesk ticketing with SLA monitoring, real-time internal communication through chat and community '
    'groups, task management, performance goal tracking, and executive dashboards with KPI visualization.')

add_para(doc,
    'The system delivers real-time notifications via WebSocket (STOMP/SockJS), supports Apache Kafka for '
    'chat event streaming, and is designed for horizontal scalability through Docker containerization with '
    'a target availability of 99.9% uptime and response times under 2 seconds for 200 concurrent users.')

add_para(doc,
    'This requirements document serves as the authoritative reference for all stakeholders including '
    'product managers, software engineers, QA engineers, DevOps personnel, and client representatives. '
    'It is intended to be reviewed alongside the associated Test Cases Document and Workflow Document.')


# ═══ 2. PROJECT SCOPE ═══
add_heading(doc, '2. Project Scope', level=1)

add_para(doc,
    'The PIXOUS HR Portal project encompasses the design, development, testing, and deployment of a '
    'comprehensive Human Resource Management System. The scope covers all modules required for complete '
    'employee lifecycle management within mid-to-large organizations.')

add_heading(doc, '2.1 In Scope', level=2)
in_scope = [
    'Authentication and authorization with JWT tokens and role-based access control (RBAC)',
    'Employee directory management with bulk import, edit, and deactivation workflows',
    'Attendance tracking via GPS punch-in/out and face recognition with geofence validation',
    'Leave management with configurable leave types, balances, approval workflows, and calendar integration',
    'Payroll processing with salary structure definition, payslip generation, batch runs, and PDF export',
    'Asset management with lifecycle tracking, QR code generation, allocation, and return workflows',
    'Helpdesk ticketing with SLA monitoring, comments, status tracking, and rating system',
    'Internal communication through community groups, direct messages, voice messages, file sharing, and polls',
    'Task management with assignment, workload visualization, and task-level chat threads',
    'Performance management with goals, reviews, and appraisal workflows',
    'Executive dashboards with KPI widgets, organizational insights, and celebration notifications',
    'Onboarding checklists with task assignments and progress tracking',
    'Travel and expense (TA) submission and approval workflows',
    'Report generation with Excel export for attendance, payroll, and organizational data',
    'System administration with organization setup, module management, branding, and audit logging',
    'Multi-platform support: Web (React), Mobile Android (React Native), Mobile iOS (Flutter)',
    'Real-time push notifications via WebSocket',
    'SaaS multi-tenant architecture supporting multiple organizations',
]
for item in in_scope:
    add_bullet(doc, item)

add_heading(doc, '2.2 Out of Scope', level=2)
out_scope = [
    'Custom model training or machine learning pipeline development beyond face recognition',
    'Direct integration with third-party payroll processors or banking APIs',
    'Recruitment and applicant tracking system (ATS)',
    'Learning management system (LMS) integration',
    'Time clock hardware integration (biometric devices)',
    'On-premises deployment (system is designed for cloud/Docker deployment)',
]
for item in out_scope:
    add_bullet(doc, item)

add_heading(doc, '2.3 Assumptions', level=2)
assumptions = [
    'The client organization will provide test data and sample employee records during the UAT phase',
    'Internet connectivity will be available for GPS and face recognition features to function',
    'The client will provision appropriate cloud infrastructure for deployment (AWS, GCP, or Azure)',
    'SSL certificates and domain configuration will be managed by the client infrastructure team',
    'Third-party email and SMS providers will be configured for notification delivery',
]
for item in assumptions:
    add_bullet(doc, item)


# ═══ 3. BUSINESS REQUIREMENTS ═══
add_heading(doc, '3. Business Requirements', level=1)

add_para(doc,
    'This section enumerates the business requirements that the PIXOUS HR Portal must satisfy. '
    'Each requirement is uniquely identified with a BR-XXX identifier for traceability purposes.')

br_rows = [
    ['BR-001', 'Multi-Industry Support', 'HIGH', 'The system must support organizations in IT and Civil/Construction industries with industry-specific configurations for attendance policies, leave types, and reporting requirements.'],
    ['BR-002', 'Role-Based Access Control', 'HIGH', 'The system must implement RBAC with 14+ permission codes covering all functional areas. Each role (Employee, Team Lead, HR Manager, Finance, CEO, Admin, etc.) must have precisely defined access permissions.'],
    ['BR-003', 'Employee Lifecycle Management', 'HIGH', 'The system must manage the complete employee lifecycle from onboarding (checklists, task assignments) through active employment (attendance, leave, payroll, performance) to offboarding (asset return, access revocation).'],
    ['BR-004', 'Attendance with GPS and Face Recognition', 'HIGH', 'The system must support GPS-based punch-in/out with geofence validation, face recognition for identity verification, and real-time attendance status updates via WebSocket.'],
    ['BR-005', 'Leave Management with Approval Workflows', 'HIGH', 'The system must provide configurable leave types (Casual, Sick, Earned, Maternity, Paternity, LOP), automated balance tracking, multi-level approval workflows, and calendar integration.'],
    ['BR-006', 'Payroll Processing with PDF Payslips', 'HIGH', 'The system must automate payroll processing including salary structure definition, component-based calculations, batch payroll runs, and PDF payslip generation with company branding.'],
    ['BR-007', 'Asset Management with QR Codes', 'MEDIUM', 'The system must track organizational assets through their lifecycle with QR code generation, employee allocation, acknowledgment workflows, and return tracking.'],
    ['BR-008', 'Helpdesk Ticketing with SLA', 'MEDIUM', 'The system must provide a ticketing system with categories, priorities, SLA tracking, status workflows (Open, In Progress, Resolved, Closed), and satisfaction ratings.'],
    ['BR-009', 'Internal Communication', 'HIGH', 'The system must support community groups, direct messaging, voice messages, file sharing, message reactions, polls, and real-time chat via WebSocket and Kafka.'],
    ['BR-010', 'Task Management', 'MEDIUM', 'The system must support task creation, assignment, status tracking, workload visualization, and task-level chat threads for team collaboration.'],
    ['BR-011', 'Performance Goals and Reviews', 'MEDIUM', 'The system must enable goal setting, periodic performance reviews, appraisal workflows, and performance history tracking for employees.'],
    ['BR-012', 'Executive Dashboards and KPIs', 'HIGH', 'The system must provide role-specific dashboards with KPI widgets, organizational insights, attendance summaries, payroll summaries, and celebration notifications.'],
    ['BR-013', 'Multi-Platform Support', 'HIGH', 'The system must be accessible via Web browsers (React), Android mobile (React Native), and iOS mobile (Flutter) with consistent feature parity and responsive design.'],
    ['BR-014', 'SaaS Multi-Tenant Architecture', 'HIGH', 'The system must support multiple organizations (tenants) on a shared infrastructure with strict data isolation, configurable branding, and per-tenant module management.'],
    ['BR-015', 'Real-Time WebSocket Notifications', 'HIGH', 'The system must deliver real-time push notifications for attendance updates, leave approvals, chat messages, task assignments, and system announcements via WebSocket (STOMP/SockJS).'],
]
add_styled_table(doc, ['ID', 'Requirement', 'Priority', 'Description'], br_rows)
doc.add_page_break()


# ═══ 4. FUNCTIONAL REQUIREMENTS ═══
add_heading(doc, '4. Functional Requirements', level=1)

add_para(doc,
    'This section details the functional requirements organized by module. Each requirement specifies '
    'the expected behavior, input/output conditions, and acceptance criteria for the corresponding feature.')

# FR-AUTH
add_heading(doc, '4.1 FR-AUTH: Authentication and Authorization', level=2)
add_para(doc,
    'The authentication and authorization module provides secure user access management using industry-standard '
    'security mechanisms. All API endpoints must be protected by JWT-based authentication.')
auth_rows = [
    ['FR-AUTH-001', 'User Login', 'The system must authenticate users via username/email and password. Failed attempts must be tracked.'],
    ['FR-AUTH-002', 'JWT Token Issuance', 'Upon successful login, the system must issue an access token (4-hour expiry) and a refresh token (7-day expiry) using HS256 signing.'],
    ['FR-AUTH-003', 'Token Refresh', 'The system must support automatic token refresh via refresh token rotation. The client must transparently refresh expired access tokens.'],
    ['FR-AUTH-004', 'Account Lockout', 'The system must lock user accounts after 5 consecutive failed login attempts. Locked accounts must display a message directing users to contact the administrator.'],
    ['FR-AUTH-005', 'Password Hashing', 'All passwords must be hashed using BCrypt with a minimum cost factor of 10. Plaintext passwords must never be stored or transmitted.'],
    ['FR-AUTH-006', 'Role-Based Access Control', 'The system must enforce RBAC with 14+ permission codes. Each API endpoint must validate the requesting user role and permissions.'],
    ['FR-AUTH-007', 'Session Management', 'The system must invalidate all tokens on logout. Concurrent sessions must be tracked and manageable by administrators.'],
    ['FR-AUTH-008', 'Password Change', 'Users must be able to change their password by providing the current password and a new password meeting complexity requirements.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], auth_rows)
doc.add_paragraph()

# FR-ATT
add_heading(doc, '4.2 FR-ATT: Attendance Management', level=2)
add_para(doc,
    'The attendance module provides comprehensive time tracking with GPS verification, face recognition, '
    'geofence validation, and real-time team visibility.')
att_rows = [
    ['FR-ATT-001', 'GPS Punch In', 'Employees can punch in using GPS coordinates. The system must capture latitude, longitude, timestamp, and device information.'],
    ['FR-ATT-002', 'GPS Punch Out', 'Employees can punch out using GPS coordinates. The system must calculate total work hours and update attendance records.'],
    ['FR-ATT-003', 'Face Recognition Punch', 'The system must verify employee identity via face recognition before recording attendance. Face embeddings must be stored securely.'],
    ['FR-ATT-004', 'Geofence Validation', 'The system must validate that punch-in/out occurs within designated geofence boundaries. Out-of-range punches must be flagged for review.'],
    ['FR-ATT-005', 'Team Attendance View', 'Managers and team leads must see real-time attendance status for their direct reports. Updates must be delivered via WebSocket.'],
    ['FR-ATT-006', 'Monthly Calendar', 'Employees must see a color-coded monthly attendance calendar showing present, absent, leave, and holiday days.'],
    ['FR-ATT-007', 'Attendance Insights', 'The system must provide attendance pattern analysis including late arrivals, early departures, and overtime hours.'],
    ['FR-ATT-008', 'Double Punch Prevention', 'The system must prevent duplicate punch-in without a prior punch-out. Users must receive clear error messaging.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], att_rows)
doc.add_paragraph()

# FR-LEA
add_heading(doc, '4.3 FR-LEA: Leave Management', level=2)
add_para(doc,
    'The leave management module handles the complete leave lifecycle including balance tracking, '
    'application submission, approval workflows, and calendar integration.')
lea_rows = [
    ['FR-LEA-001', 'Leave Types', 'The system must support configurable leave types: Casual, Sick, Earned, Maternity, Paternity, and LOP (Loss of Pay).'],
    ['FR-LEA-002', 'Leave Balances', 'Employees must view current leave balances for each leave type with remaining and total entitlements.'],
    ['FR-LEA-003', 'Apply for Leave', 'Employees must apply for leave by selecting type, dates, reason, and submitting the request. The system must validate balance availability.'],
    ['FR-LEA-004', 'Approval Workflow', 'Leave requests must route to the appropriate approver (Team Lead or HR Manager) based on organizational hierarchy.'],
    ['FR-LEA-005', 'Leave Calendar', 'Managers must view team leave calendars showing approved leaves with date ranges and leave types.'],
    ['FR-LEA-006', 'LOP Management', 'When employees exceed their leave entitlement, LOP days must be automatically calculated and reflected in payroll processing.'],
    ['FR-LEA-007', 'Cancel Leave', 'Employees must be able to cancel pending or approved leave requests. Approved cancellations must restore leave balances.'],
    ['FR-LEA-008', 'Bulk Operations', 'HR managers must be able to approve or reject multiple leave requests in a single batch operation.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], lea_rows)
doc.add_paragraph()

# FR-PAY
add_heading(doc, '4.4 FR-PAY: Payroll Processing', level=2)
add_para(doc,
    'The payroll module automates salary calculations, payslip generation, batch processing, '
    'and PDF export with company branding.')
pay_rows = [
    ['FR-PAY-001', 'Salary Structure', 'HR must define salary structures with components: Basic, HRA, Conveyance, Special Allowance, PF, ESI, and Professional Tax.'],
    ['FR-PAY-002', 'Payslip Generation', 'The system must generate payslips with earnings, deductions, LOP adjustments, and net pay calculations.'],
    ['FR-PAY-003', 'PDF Payslip Export', 'Payslips must be exportable as PDF documents with PIXOUS company branding, employee details, and all financial breakdowns.'],
    ['FR-PAY-004', 'Batch Payroll Runs', 'HR must be able to initiate monthly payroll runs for all employees. The system must process calculations in batch and track run status.'],
    ['FR-PAY-005', 'Payroll Approval', 'Payroll runs must go through a two-stage approval: HR confirmation followed by Finance approval before finalization.'],
    ['FR-PAY-006', 'Employee Salary View', 'Employees must be able to view their own salary structure and payslip history. Salary information of other employees must be restricted.'],
    ['FR-PAY-007', 'LOP Integration', 'The payroll system must integrate with the leave module to automatically deduct LOP days from salary calculations.'],
    ['FR-PAY-008', 'Payslip History', 'Employees and HR must be able to access historical payslips organized by month and year with PDF download capability.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], pay_rows)
doc.add_paragraph()

# FR-AST
add_heading(doc, '4.5 FR-AST: Asset Management', level=2)
add_para(doc,
    'The asset management module tracks organizational assets through their complete lifecycle '
    'from procurement through allocation to return or disposal.')
ast_rows = [
    ['FR-AST-001', 'Asset CRUD', 'Administrators must be able to create, read, update, and delete asset records with fields: name, type, serial number, and status.'],
    ['FR-AST-002', 'Asset Allocation', 'Assets must be allocatable to employees with allocation date, expected return date, and acknowledgment requirement.'],
    ['FR-AST-003', 'QR Code Generation', 'Each asset must have a QR code generated and downloadable as PNG for physical tagging and tracking.'],
    ['FR-AST-004', 'Asset Acknowledgment', 'Employees must acknowledge receipt of allocated assets. The system must record acknowledgment date and digital signature.'],
    ['FR-AST-005', 'Asset Return', 'Allocated assets must be returnable with condition assessment, return date, and updated status tracking.'],
    ['FR-AST-006', 'Asset Status Tracking', 'The system must maintain asset status: Available, Allocated, Under Maintenance, Returned, or Disposed.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], ast_rows)
doc.add_paragraph()

# FR-HD
add_heading(doc, '4.6 FR-HD: Helpdesk and Ticketing', level=2)
add_para(doc,
    'The helpdesk module provides a ticketing system with SLA monitoring, status workflows, '
    'comments, and satisfaction ratings.')
hd_rows = [
    ['FR-HD-001', 'Create Ticket', 'Employees must create tickets with subject, description, category, priority, and optional file attachments.'],
    ['FR-HD-002', 'Ticket Status Workflow', 'Tickets must progress through statuses: Open, In Progress, Resolved, and Closed with appropriate transitions.'],
    ['FR-HD-003', 'Comments and Notes', 'All parties must be able to add comments and internal notes to tickets with timestamps and author information.'],
    ['FR-HD-004', 'SLA Monitoring', 'The system must track SLA compliance based on ticket priority with escalation alerts for approaching deadlines.'],
    ['FR-HD-005', 'Ticket Rating', 'Employees must rate resolved tickets on a 1-5 star scale. Ratings must be visible to management.'],
    ['FR-HD-006', 'File Attachments', 'Tickets must support file attachments for screenshots, documents, and other supporting evidence.'],
    ['FR-HD-007', 'Ticket Assignment', 'Admin and agents must be able to assign tickets to specific team members or departments.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], hd_rows)
doc.add_paragraph()

# FR-CHAT
add_heading(doc, '4.7 FR-CHAT: Community and Chat', level=2)
add_para(doc,
    'The communication module provides internal messaging through community groups, direct messages, '
    'voice messages, file sharing, reactions, and polls.')
chat_rows = [
    ['FR-CHAT-001', 'Community Groups', 'Users must be able to create, join, and manage community groups with configurable membership and moderation.'],
    ['FR-CHAT-002', 'Direct Messages', 'Users must be able to send private 1:1 direct messages to other organization members.'],
    ['FR-CHAT-003', 'Text Messaging', 'All chat interfaces must support rich text messaging with timestamps, sender information, and read receipts.'],
    ['FR-CHAT-004', 'Voice Messages', 'Users must be able to record, send, and play voice messages with duration indicators.'],
    ['FR-CHAT-005', 'File Sharing', 'Chat interfaces must support file uploads and downloads for documents, images, and other media.'],
    ['FR-CHAT-006', 'Message Reactions', 'Users must be able to react to messages with emojis. Reaction counts must be visible.'],
    ['FR-CHAT-007', 'Message Pinning', 'Group administrators must be able to pin important messages for visibility.'],
    ['FR-CHAT-008', 'Polls', 'Users must be able to create and vote in polls within community groups. Results must be displayed in real-time.'],
    ['FR-CHAT-009', 'Kafka Event Streaming', 'Chat events must be streamed via Apache Kafka for reliable message delivery and persistence.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], chat_rows)
doc.add_paragraph()

# FR-TASK
add_heading(doc, '4.8 FR-TASK: Task Management', level=2)
add_para(doc,
    'The task management module enables creation, assignment, tracking, and collaboration on tasks '
    'within teams and across departments.')
task_rows = [
    ['FR-TASK-001', 'Task CRUD', 'Users with appropriate roles must be able to create, read, update, and delete tasks with title, description, assignee, and due date.'],
    ['FR-TASK-002', 'Task Assignment', 'Tasks must be assignable to individual employees or team members with notification delivery.'],
    ['FR-TASK-003', 'Task Status', 'Tasks must support statuses: Pending, In Progress, Completed, and Cancelled with transition rules.'],
    ['FR-TASK-004', 'Workload View', 'Managers must see team workload visualizations showing task counts per team member by status.'],
    ['FR-TASK-005', 'Task Chat', 'Each task must have a dedicated chat thread for discussion and collaboration on the task.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], task_rows)
doc.add_paragraph()

# FR-DASH
add_heading(doc, '4.9 FR-DASH: Dashboard and Reporting', level=2)
add_para(doc,
    'The dashboard module provides role-specific views with KPI widgets, organizational insights, '
    'and data visualization for decision-making.')
dash_rows = [
    ['FR-DASH-001', 'Personal Dashboard', 'Employees must see personal widgets: attendance summary, leave balance, pending tasks, notifications, and celebrations.'],
    ['FR-DASH-002', 'Executive Dashboard', 'Executives must see organizational KPIs: headcount, attrition rate, attendance percentage, payroll summary, and ticket resolution rates.'],
    ['FR-DASH-003', 'Org Insights', 'HR and executives must see organizational charts: department distribution, gender ratio, location spread, and age demographics.'],
    ['FR-DASH-004', 'Celebrations', 'The system must display birthdays and work anniversaries with employee photos and dates.'],
    ['FR-DASH-005', 'Report Export', 'Authorized users must be able to export attendance and payroll reports as Excel files.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], dash_rows)
doc.add_paragraph()

# FR-NOTIF
add_heading(doc, '4.10 FR-NOTIF: Notifications', level=2)
add_para(doc,
    'The notification module delivers real-time push notifications for system events across all platforms.')
notif_rows = [
    ['FR-NOTIF-001', 'WebSocket Push', 'The system must deliver real-time notifications via WebSocket (STOMP/SockJS) for attendance, leave, chat, and task events.'],
    ['FR-NOTIF-002', 'Notification Bell', 'The UI must display a notification bell icon with unread count badge. Users must be able to view and mark notifications as read.'],
    ['FR-NOTIF-003', 'Event Types', 'The system must support notifications for: leave requests, leave approvals/rejections, task assignments, chat messages, ticket updates, and announcements.'],
    ['FR-NOTIF-004', 'Desktop Notifications', 'The web application must support browser notification APIs for desktop push alerts when the application is in the background.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], notif_rows)
doc.add_paragraph()

# FR-ONB
add_heading(doc, '4.11 FR-ONB: Onboarding', level=2)
add_para(doc,
    'The onboarding module streamlines new employee integration with checklists, task assignments, '
    'and progress tracking.')
onb_rows = [
    ['FR-ONB-001', 'Onboarding Checklists', 'HR must create configurable onboarding checklists with tasks, assignees, and deadlines.'],
    ['FR-ONB-002', 'Task Tracking', 'Each onboarding task must have status tracking (Pending, In Progress, Completed) with progress indicators.'],
    ['FR-ONB-003', 'Progress Dashboard', 'HR must see onboarding progress for all new hires with completion percentages and overdue indicators.'],
    ['FR-ONB-004', 'Automated Notifications', 'The system must notify relevant parties when onboarding tasks are due, completed, or overdue.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], onb_rows)
doc.add_paragraph()

# FR-PERF
add_heading(doc, '4.12 FR-PERF: Performance Management', level=2)
add_para(doc,
    'The performance module enables goal setting, periodic reviews, and appraisal workflows '
    'for employee development and organizational growth.')
perf_rows = [
    ['FR-PERF-001', 'Goal Setting', 'Managers and employees must be able to define performance goals with metrics, targets, and review periods.'],
    ['FR-PERF-002', 'Performance Reviews', 'The system must support periodic performance reviews with self-assessment, manager assessment, and rating scales.'],
    ['FR-PERF-003', 'Appraisal Workflows', 'The system must manage appraisal cycles with configurable review periods, rating criteria, and approval hierarchies.'],
    ['FR-PERF-004', 'Performance History', 'Employees and managers must be able to view historical performance data and trend analysis.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], perf_rows)
doc.add_paragraph()

# FR-EXP
add_heading(doc, '4.13 FR-EXP: Travel and Expense', level=2)
add_para(doc,
    'The expense module manages employee expense submissions, approvals, and category-based tracking.')
exp_rows = [
    ['FR-EXP-001', 'Expense Submission', 'Employees must submit expenses with amount, category, date, description, and receipt attachments.'],
    ['FR-EXP-002', 'Expense Approval', 'Expenses must go through manager approval workflows with status tracking (Submitted, Approved, Rejected, Reimbursed).'],
    ['FR-EXP-003', 'Expense Categories', 'The system must support configurable expense categories: Travel, Food, Accommodation, Equipment, and Miscellaneous.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], exp_rows)
doc.add_paragraph()

# FR-RPT
add_heading(doc, '4.14 FR-RPT: Reports', level=2)
add_para(doc,
    'The reports module generates operational and analytical reports with export capabilities.')
rpt_rows = [
    ['FR-RPT-001', 'Attendance Reports', 'HR must generate attendance reports with date range filters, department filters, and individual employee summaries.'],
    ['FR-RPT-002', 'Payroll Reports', 'Finance and HR must generate payroll reports with salary breakdowns, tax summaries, and department-wise totals.'],
    ['FR-RPT-003', 'Excel Export', 'All reports must be exportable as Excel (.xlsx) files with proper formatting, headers, and data integrity.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], rpt_rows)
doc.add_paragraph()

# FR-ADMIN
add_heading(doc, '4.15 FR-ADMIN: System Administration', level=2)
add_para(doc,
    'The administration module provides system-wide configuration, branding, module management, '
    'and audit logging capabilities.')
admin_rows = [
    ['FR-ADMIN-001', 'Organization Setup', 'Admins must configure organization details including name, logo, departments, designations, and office locations.'],
    ['FR-ADMIN-002', 'Module Management', 'Admins must be able to enable or disable system modules. Disabled modules must be hidden from all users.'],
    ['FR-ADMIN-003', 'Branding', 'Admins must configure application branding including logo, primary colors, and theme settings.'],
    ['FR-ADMIN-004', 'Holiday Management', 'Admins must define company-wide holidays for attendance exclusion and calendar display.'],
    ['FR-ADMIN-005', 'Audit Logging', 'The system must log all significant user actions with timestamp, user, action type, IP address, and details.'],
    ['FR-ADMIN-006', 'Cache Management', 'Admins must be able to clear Redis cache to force data refresh across the application.'],
    ['FR-ADMIN-007', 'Login History', 'Admins must be able to review login history with timestamps, IP addresses, browsers, and success/failure status.'],
    ['FR-ADMIN-008', 'Global Announcements', 'Admins must be able to create and publish global announcements with entrance effects visible to all users on login.'],
]
add_styled_table(doc, ['ID', 'Feature', 'Description'], admin_rows)
doc.add_page_break()


# ═══ 5. NON-FUNCTIONAL REQUIREMENTS ═══
add_heading(doc, '5. Non-Functional Requirements', level=1)

add_para(doc,
    'Non-functional requirements define the quality attributes and constraints that the PIXOUS HR Portal '
    'must satisfy. These requirements complement the functional requirements and define the system\'s '
    'operational characteristics.')

nfr_rows = [
    ['NFR-001', 'Performance', 'HIGH', 'The system must respond to user requests within 2 seconds under normal load. The system must support at least 200 concurrent users without performance degradation. Database queries must execute within 200ms for standard operations.'],
    ['NFR-002', 'Security', 'CRITICAL', 'All API communication must use HTTPS with TLS 1.2+. Passwords must be hashed with BCrypt (cost factor >= 10). The system must implement OWASP Top 10 protections including SQL injection prevention, XSS mitigation, CSRF tokens, and secure session management.'],
    ['NFR-003', 'Scalability', 'HIGH', 'The system must support horizontal scaling through Docker containerization. The architecture must enable adding application instances behind a load balancer without code changes. Database connection pooling must be configured for scaling.'],
    ['NFR-004', 'Availability', 'HIGH', 'The system must target 99.9% uptime (approximately 8.76 hours of downtime per year). Scheduled maintenance windows must be communicated at least 48 hours in advance. The system must implement health checks and automatic restart capabilities.'],
    ['NFR-005', 'Usability', 'MEDIUM', 'The web interface must be responsive and functional across desktop, tablet, and mobile viewports. The system must support Progressive Web App (PWA) capabilities for offline indicators and home screen installation. Navigation must be intuitive with consistent UI patterns.'],
    ['NFR-006', 'Maintainability', 'MEDIUM', 'The codebase must follow modular architecture principles with clear separation of concerns. Database schema changes must be managed through Flyway migrations (currently 101 migrations). Code must include adequate documentation and follow established coding standards.'],
    ['NFR-007', 'Compatibility', 'MEDIUM', 'The web application must function correctly in Chrome 120+, Firefox 120+, Safari 17+, and Edge 120+. Mobile applications must support Android 12+ (React Native) and iOS 16+ (Flutter). Testing must cover all supported platforms.'],
    ['NFR-008', 'Data Integrity', 'HIGH', 'The database must enforce ACID compliance through MySQL 8.4 InnoDB engine. Optimistic locking must be implemented for concurrent data modifications. Data validation must occur at both application and database levels. Foreign key constraints must maintain referential integrity.'],
]
add_styled_table(doc, ['ID', 'Category', 'Priority', 'Description'], nfr_rows)
doc.add_page_break()


# ═══ 6. TECHNICAL REQUIREMENTS ═══
add_heading(doc, '6. Technical Requirements', level=1)

add_heading(doc, '6.1 Technology Stack', level=2)
tech_rows = [
    ['Backend Framework', 'Spring Boot 3.2 with Java 17+', 'REST API development, dependency injection, security'],
    ['Web Frontend', 'React 18 with TypeScript', 'Single-page application with type safety and component architecture'],
    ['Android Mobile', 'React Native', 'Cross-platform mobile application for Android devices'],
    ['iOS Mobile', 'Flutter', 'Cross-platform mobile application for iOS devices'],
    ['Database', 'MySQL 8.4 (InnoDB)', 'Primary data storage with ACID compliance'],
    ['Schema Migration', 'Flyway', 'Version-controlled database schema management (101 migrations)'],
    ['Cache Layer', 'Redis 7 (optional)', 'Session caching, permission caching, with in-memory fallback'],
    ['Message Streaming', 'Apache Kafka 3.9', 'Chat event streaming and message persistence'],
    ['Real-Time', 'WebSocket (STOMP/SockJS)', 'Bidirectional real-time communication for notifications and chat'],
    ['Build Tool', 'Maven / npm / Flutter CLI', 'Dependency management and build automation for each platform'],
    ['Containerization', 'Docker / Docker Compose', 'Application packaging and orchestration for deployment'],
    ['Version Control', 'Git', 'Source code management with branching strategy'],
    ['API Documentation', 'OpenAPI / Swagger', 'REST API specification and documentation generation'],
]
add_styled_table(doc, ['Component', 'Technology', 'Purpose'], tech_rows)

add_heading(doc, '6.2 Database Design', level=2)
add_para(doc,
    'The database schema uses MySQL 8.4 with the InnoDB storage engine to ensure ACID compliance '
    'and referential integrity. Schema changes are managed through Flyway migrations, with 101 migration '
    'scripts currently tracked in the version control system. The schema includes tables for users, '
    'employees, departments, designations, attendance records, leave requests, leave balances, payslips, '
    'salary structures, assets, asset allocations, helpdesk tickets, ticket comments, chat messages, '
    'community groups, tasks, goals, performance reviews, expenses, audit logs, and system configuration.')

add_heading(doc, '6.3 API Design', level=2)
add_para(doc,
    'The REST API follows RESTful conventions with consistent resource naming, HTTP method semantics, '
    'and status code usage. All API responses follow a standardized envelope format with data, message, '
    'and pagination fields. The API is documented using OpenAPI/Swagger specifications for client generation '
    'and testing. API versioning is managed through URL path prefixing (/api/v1/).')

add_heading(doc, '6.4 Authentication Architecture', level=2)
add_para(doc,
    'Authentication uses JSON Web Tokens (JWT) with HS256 signing. The system issues short-lived access '
    'tokens (4-hour expiry) and long-lived refresh tokens (7-day expiry). Token refresh implements rotation '
    'where each refresh operation issues a new refresh token and invalidates the previous one. '
    'The backend validates tokens on every protected API endpoint using Spring Security filters.')

add_heading(doc, '6.5 Caching Strategy', level=2)
add_para(doc,
    'The system implements an optional Redis caching layer with automatic fallback to in-memory caching '
    'when Redis is unavailable. Cached data includes user permissions, organization settings, and frequently '
    'accessed reference data. Cache invalidation follows a write-through strategy with TTL-based expiration. '
    'Administrators can manually clear the cache through the admin interface.')

add_heading(doc, '6.6 Messaging Architecture', level=2)
add_para(doc,
    'Apache Kafka serves as the event streaming platform for chat messages and real-time system events. '
    'Kafka topics are organized by event type (chat.messages, attendance.updates, leave.notifications). '
    'The system uses consumer groups for horizontal scaling of event processors. Message retention is '
    'configured for 7 days with configurable archival policies.')

add_heading(doc, '6.7 Real-Time Communication', level=2)
add_para(doc,
    'The WebSocket layer uses STOMP protocol over SockJS for broad browser compatibility. The WebSocket '
    'server handles subscription management, message routing, and connection lifecycle. Topics follow the '
    'naming convention /topic/{event-type}/{organization-id}. The system supports both authenticated '
    'WebSocket connections (via JWT in handshake headers) and automatic reconnection with session recovery.')

add_heading(doc, '6.8 Security Implementation', level=2)
add_para(doc,
    'Security is implemented at multiple layers: transport encryption (HTTPS/TLS 1.2+), application-level '
    'authentication (JWT), authorization (RBAC with 14+ permission codes), data protection (BCrypt password '
    'hashing, AES-256 encryption for sensitive fields), and input validation (parameterized queries, output '
    'encoding). The system implements OWASP Top 10 protections including SQL injection prevention, cross-site '
    'scripting (XSS) mitigation, cross-site request forgery (CSRF) tokens, and secure HTTP headers.')
doc.add_page_break()


# ═══ 7. USER ROLE REQUIREMENTS ═══
add_heading(doc, '7. User Role Requirements', level=1)

add_para(doc,
    'The PIXOUS HR Portal implements a comprehensive role-based access control system with 14+ distinct '
    'permission codes. Each role is designed to provide precisely the access needed for that function '
    'without over-provisioning privileges. The following table details each role and its capabilities.')

role_rows = [
    ['Super Admin', 'SYSTEM_ADMIN', 'Full system access including organization setup, user management, module configuration, branding, audit logs, and cache management. Can access all tenants in multi-tenant mode.'],
    ['HR Manager', 'HR_MGR', 'Employee lifecycle management, leave approval, payroll processing, payslip generation, attendance reports, onboarding management, and organization configuration within assigned industry.'],
    ['Team Lead / Manager', 'TEAM_LEAD', 'Team attendance viewing, leave approval for direct reports, task creation and assignment, team workload monitoring, and performance reviews for team members.'],
    ['Finance Officer', 'FINANCE', 'Payroll approval workflow, salary structure review, financial reports, expense approval, and payroll reconciliation. Cannot modify employee data or attendance records.'],
    ['CEO / Executive', 'EXECUTIVE', 'Executive dashboard access, organizational KPIs, headcount analytics, attrition monitoring, and celebration notifications. Read-only access to operational data.'],
    ['Asset Manager', 'ASSET_MGR', 'Full asset lifecycle management including creation, allocation, return, QR code generation, and maintenance tracking. Cannot modify employee records or payroll data.'],
    ['Helpdesk Agent', 'HD_AGENT', 'Ticket management including assignment, status updates, comments, resolution, and SLA monitoring. Cannot access employee records or payroll information.'],
    ['Employee', 'EMPLOYEE', 'Self-service access to attendance (GPS/face punch), leave application, payslip viewing, helpdesk ticket creation, chat participation, task management, and personal dashboard.'],
    ['IT Administrator', 'IT_ADMIN', 'Technical system administration including module management, API configuration, cache management, and audit log review. Limited access to HR operational features.'],
    ['Project Manager', 'PM', 'Task creation and assignment, project workload visualization, team performance monitoring, and cross-functional coordination. Cannot approve leaves or process payroll.'],
    ['Recruiter', 'RECRUITER', 'Access to onboarding checklists, new hire tracking, and recruitment-related announcements. Limited system access focused on hiring workflows.'],
    ['Compliance Officer', 'COMPLIANCE', 'Audit log review, compliance reporting, data access monitoring, and regulatory report generation. Read-only access across most modules.'],
    ['Department Head', 'DEPT_HEAD', 'Department-level employee management, leave approval for department members, department attendance reports, and departmental goal setting.'],
    ['Intern', 'INTERN', 'Limited self-service access including attendance punch, basic task participation, and community chat. Cannot apply for leave without manager pre-approval.'],
]
add_styled_table(doc, ['Role Name', 'Role Code', 'Capabilities and Access Scope'], role_rows)
doc.add_page_break()


# ═══ 8. MODULE REQUIREMENTS ═══
add_heading(doc, '8. Module Requirements', level=1)

add_para(doc,
    'The PIXOUS HR Portal is composed of 25+ functional modules that can be independently enabled or '
    'disabled by system administrators. This modular architecture allows organizations to activate only '
    'the capabilities they need, reducing complexity and training overhead.')

module_rows = [
    ['MOD-01', 'Authentication Module', 'JWT-based login, token refresh, password management, account lockout', 'Core'],
    ['MOD-02', 'Employee Directory', 'Employee CRUD, bulk import, profile management, deactivation', 'Core'],
    ['MOD-03', 'Organization Setup', 'Departments, designations, locations, holidays configuration', 'Core'],
    ['MOD-04', 'Attendance Module', 'GPS punch, face recognition, geofence, calendar, insights', 'Core'],
    ['MOD-05', 'Leave Module', 'Leave types, balances, apply, approve, calendar, LOP', 'Core'],
    ['MOD-06', 'Payroll Module', 'Salary structure, payslips, batch runs, PDF generation', 'Core'],
    ['MOD-07', 'Asset Management', 'Asset CRUD, allocation, QR codes, acknowledgment, return', 'Optional'],
    ['MOD-08', 'Helpdesk Module', 'Tickets, comments, status, ratings, SLA monitoring', 'Optional'],
    ['MOD-09', 'Community & Chat', 'Groups, DMs, voice, files, reactions, polls, Kafka streaming', 'Optional'],
    ['MOD-10', 'Task Module', 'Task CRUD, assignment, workload view, task chat', 'Optional'],
    ['MOD-11', 'Dashboard Module', 'Personal, executive, org insights, celebrations', 'Core'],
    ['MOD-12', 'Notification Module', 'WebSocket push, notification bell, desktop alerts', 'Core'],
    ['MOD-13', 'Onboarding Module', 'Checklists, tasks, progress tracking', 'Optional'],
    ['MOD-14', 'Performance Module', 'Goals, reviews, appraisals, performance history', 'Optional'],
    ['MOD-15', 'Expense Module', 'TA submission, approval, categories', 'Optional'],
    ['MOD-16', 'Report Module', 'Attendance reports, payroll reports, Excel export', 'Core'],
    ['MOD-17', 'Admin Module', 'Organization setup, module management, branding, audit', 'Core'],
    ['MOD-18', 'Calendar Module', 'Leave calendar, holiday calendar, team availability', 'Core'],
    ['MOD-19', 'Search Module', 'Employee search, global search, filtered search', 'Core'],
    ['MOD-20', 'Document Module', 'Document upload, storage, and retrieval', 'Optional'],
    ['MOD-21', 'Announcement Module', 'Global announcements with entrance effects', 'Optional'],
    ['MOD-22', 'PWA Module', 'Progressive Web App installation and offline indicators', 'Optional'],
    ['MOD-23', 'Geofence Module', 'GPS geofence configuration and validation', 'Core'],
    ['MOD-24', 'Face Recognition Module', 'Face embedding storage and verification', 'Optional'],
    ['MOD-25', 'Audit Module', 'Action logging, login history, data access tracking', 'Core'],
]
add_styled_table(doc, ['Module ID', 'Module Name', 'Key Features', 'Category'], module_rows)
doc.add_page_break()


# ═══ 9. INTEGRATION REQUIREMENTS ═══
add_heading(doc, '9. Integration Requirements', level=1)

add_para(doc,
    'The PIXOUS HR Portal must integrate with various external systems and services to deliver '
    'complete functionality. The following integration points are defined.')

int_rows = [
    ['INT-001', 'Email Service', 'SMTP/Email API integration for sending notifications, password resets, and system alerts', 'Required'],
    ['INT-002', 'SMS Service', 'SMS gateway integration for OTP verification and critical notifications', 'Optional'],
    ['INT-003', 'Cloud Storage', 'S3-compatible storage for file attachments, documents, and asset images', 'Required'],
    ['INT-004', 'Face Recognition API', 'Integration with face recognition service for attendance verification', 'Optional'],
    ['INT-005', 'Map/Geofence API', 'Google Maps or similar API for GPS validation and geofence configuration', 'Required'],
    ['INT-006', 'PDF Generation', 'Server-side PDF generation service for payslips and reports', 'Required'],
    ['INT-007', 'QR Code Generation', 'QR code generation service for asset tagging and identification', 'Optional'],
    ['INT-008', 'Monitoring/Prometheus', 'Application metrics collection for performance monitoring and alerting', 'Optional'],
    ['INT-009', 'Log Aggregation', 'Centralized logging integration (ELK stack or similar) for troubleshooting', 'Optional'],
    ['INT-010', 'SSO Integration', 'OAuth 2.0 / SAML integration for enterprise single sign-on', 'Future'],
]
add_styled_table(doc, ['Integration ID', 'System', 'Description', 'Priority'], int_rows)
doc.add_page_break()


# ═══ 10. DEPLOYMENT REQUIREMENTS ═══
add_heading(doc, '10. Deployment Requirements', level=1)

add_para(doc,
    'The PIXOUS HR Portal must be deployable across multiple environments with consistent behavior '
    'and reliable operations. The following deployment requirements define the infrastructure and '
    'operational needs.')

dep_rows = [
    ['DEP-001', 'Docker Containerization', 'All application components must be containerized using Docker with multi-stage builds for optimized image sizes.'],
    ['DEP-002', 'Docker Compose', 'A docker-compose.yml must be provided for local development and testing with all required services.'],
    ['DEP-003', 'Environment Configuration', 'Application configuration must be externalized through environment variables and configuration files per environment.'],
    ['DEP-004', 'Database Migration', 'Flyway migrations must run automatically on application startup to ensure schema consistency across environments.'],
    ['DEP-005', 'Health Checks', 'The application must expose health check endpoints (/actuator/health) for container orchestration and load balancer integration.'],
    ['DEP-006', 'Logging', 'Application logs must be structured (JSON format) and output to stdout/stderr for container log aggregation.'],
    ['DEP-007', 'Graceful Shutdown', 'The application must handle SIGTERM signals gracefully, completing in-flight requests before terminating.'],
    ['DEP-008', 'Resource Limits', 'Docker containers must define CPU and memory limits to prevent resource exhaustion.'],
    ['DEP-009', 'Backup Strategy', 'Database backups must be automated with point-in-time recovery capability and off-site storage.'],
    ['DEP-010', 'SSL/TLS', 'All production deployments must use HTTPS with valid SSL certificates. HTTP must redirect to HTTPS.'],
]
add_styled_table(doc, ['Requirement ID', 'Description', 'Details'], dep_rows)
doc.add_page_break()


# ═══ 11. ACCEPTANCE CRITERIA ═══
add_heading(doc, '11. Acceptance Criteria', level=1)

add_para(doc,
    'The following acceptance criteria define the minimum requirements that must be satisfied for the '
    'PIXOUS HR Portal to be considered ready for production deployment. Each criterion must be verified '
    'through manual testing and/or automated test execution.')

ac_rows = [
    ['AC-001', 'Authentication', 'All users can log in with valid credentials. JWT tokens are issued and refreshed correctly. Account lockout activates after 5 failed attempts.'],
    ['AC-002', 'Role-Based Access', 'Each of the 14+ roles can access only their permitted features. Unauthorized access attempts return 403 Forbidden.'],
    ['AC-003', 'Attendance', 'GPS punch-in/out records location and timestamp. Face recognition verifies identity. Geofence validation rejects out-of-range punches.'],
    ['AC-004', 'Leave Management', 'Leave application, approval, rejection, and cancellation workflows function correctly. Leave balances update accurately.'],
    ['AC-005', 'Payroll', 'Salary calculations are accurate. Payslips generate correctly with all components. PDF exports contain correct data and branding.'],
    ['AC-006', 'Helpdesk', 'Tickets can be created, assigned, updated, resolved, and rated. SLA tracking functions correctly.'],
    ['AC-007', 'Chat and Communication', 'Text, voice, and file messages are delivered in real-time. Community groups and direct messages function correctly.'],
    ['AC-008', 'Task Management', 'Tasks can be created, assigned, tracked, and completed. Workload views display accurate data.'],
    ['AC-009', 'Dashboard', 'Personal and executive dashboards display accurate, real-time data. KPI calculations are correct.'],
    ['AC-010', 'Notifications', 'WebSocket notifications are delivered in real-time for all configured event types. Notification history is maintained.'],
    ['AC-011', 'Performance', 'Goal setting and review workflows function correctly. Performance history is maintained and displayed.'],
    ['AC-012', 'Reports', 'Attendance and payroll reports generate correctly. Excel exports contain accurate data with proper formatting.'],
    ['AC-013', 'Cross-Browser', 'The application functions correctly in Chrome, Firefox, Safari, and Edge. Responsive design adapts to all viewport sizes.'],
    ['AC-014', 'Performance', 'API response times are under 2 seconds for 200 concurrent users. No memory leaks observed during extended testing.'],
    ['AC-015', 'Security', 'No critical or high-severity vulnerabilities found in security scanning. OWASP Top 10 protections are verified.'],
]
add_styled_table(doc, ['Criterion ID', 'Area', 'Acceptance Requirement'], ac_rows)
doc.add_page_break()


# ═══ 12. TRACEABILITY MATRIX ═══
add_heading(doc, '12. Traceability Matrix', level=1)

add_para(doc,
    'The traceability matrix maps each requirement to its corresponding test cases and verification '
    'method. This matrix ensures complete coverage and provides a clear audit trail from requirement '
    'to test execution.')

trace_rows = [
    ['BR-001', 'Multi-Industry Support', 'TC-BRW-001 to TC-BRW-006', 'Manual Testing', 'Verified'],
    ['BR-002', 'Role-Based Access Control', 'TC-AUTH-EMP-001 to TC-AUTH-EMP-010', 'Automated + Manual', 'Verified'],
    ['BR-003', 'Employee Lifecycle', 'TC-HR-EMP-001 to TC-HR-EMP-006', 'Manual Testing', 'Verified'],
    ['BR-004', 'Attendance with GPS/Face', 'TC-EMP-ATT-001 to TC-EMP-ATT-007', 'Manual Testing', 'In Progress'],
    ['BR-005', 'Leave Management', 'TC-EMP-LEA-001 to TC-EMP-LEA-005', 'Manual Testing', 'Verified'],
    ['BR-006', 'Payroll Processing', 'TC-HR-PAY-001 to TC-HR-PAY-005', 'Manual Testing', 'Verified'],
    ['BR-007', 'Asset Management', 'TC-AST-ADM-001 to TC-AST-EMP-002', 'Manual Testing', 'Verified'],
    ['BR-008', 'Helpdesk Ticketing', 'TC-HD-EMP-001 to TC-HD-EMP-002', 'Manual Testing', 'Verified'],
    ['BR-009', 'Internal Communication', 'TC-CHAT-001 to TC-CHAT-006', 'Manual Testing', 'In Progress'],
    ['BR-010', 'Task Management', 'TC-TASK-001 to TC-TASK-004', 'Manual Testing', 'Verified'],
    ['BR-011', 'Performance Goals', 'TC-ONB-001 to TC-ONB-002', 'Manual Testing', 'Planned'],
    ['BR-012', 'Executive Dashboards', 'TC-CTO-DASH-001 to TC-CTO-DASH-004', 'Manual Testing', 'Verified'],
    ['BR-013', 'Multi-Platform Support', 'TC-BRW-001 to TC-BRW-006', 'Manual Testing', 'In Progress'],
    ['BR-014', 'SaaS Multi-Tenant', 'TC-ADM-SET-001 to TC-ADM-SET-004', 'Manual Testing', 'Verified'],
    ['BR-015', 'WebSocket Notifications', 'TC-ATT-XROLE-003', 'Manual Testing', 'Verified'],
]
add_styled_table(doc, ['Requirement ID', 'Requirement', 'Test Cases', 'Method', 'Status'], trace_rows)
doc.add_page_break()


# ═══ 13. GLOSSARY ═══
add_heading(doc, '13. Glossary', level=1)

add_para(doc,
    'The following glossary defines technical terms, acronyms, and domain-specific language used '
    'throughout this requirements document. Understanding these terms is essential for accurate '
    'interpretation of the requirements.')

glossary_rows = [
    ['ACID', 'Atomicity, Consistency, Isolation, Durability - database transaction properties ensuring reliability'],
    ['BCrypt', 'A password hashing function designed for security, used for storing user passwords'],
    ['CRUD', 'Create, Read, Update, Delete - the four basic operations of persistent storage'],
    ['Flyway', 'A database migration tool that manages schema changes through versioned SQL scripts'],
    ['Geofence', 'A virtual boundary defined by GPS coordinates for location-based access control'],
    ['HRMS', 'Human Resource Management System - software for managing HR processes'],
    ['JWT', 'JSON Web Token - a compact, URL-safe means of representing claims between parties'],
    ['Kafka', 'Apache Kafka - a distributed event streaming platform for high-throughput data feeds'],
    ['LOP', 'Loss of Pay - leave days deducted from salary when leave entitlement is exceeded'],
    ['OAuth 2.0', 'An open standard for access delegation, commonly used for third-party login'],
    ['PWA', 'Progressive Web App - a type of web application that provides app-like experience'],
    ['RBAC', 'Role-Based Access Control - restricting system access based on user roles'],
    ['Redis', 'An in-memory data structure store used for caching and session management'],
    ['SaaS', 'Software as a Service - a software licensing and delivery model'],
    ['SLA', 'Service Level Agreement - a commitment between a service provider and client'],
    ['SockJS', 'A JavaScript library providing cross-domain communication channels'],
    ['STOMP', 'Simple Text Oriented Messaging Protocol - a messaging protocol for WebSocket'],
    ['WebSocket', 'A protocol providing full-duplex communication channels over a single TCP connection'],
]
add_styled_table(doc, ['Term', 'Definition'], glossary_rows)
doc.add_page_break()


# ═══ 14. APPENDIX ═══
add_heading(doc, '14. Appendix', level=1)

add_heading(doc, '14.1 Document References', level=2)
doc_refs = [
    ['DOC-001', 'PIXOUS HR Portal Test Cases Document', 'Comprehensive test cases for all modules and roles'],
    ['DOC-002', 'PIXOUS HR Portal Workflow Document', 'Detailed workflow diagrams and process flows'],
    ['DOC-003', 'OpenAPI Specification', 'REST API documentation generated from Swagger annotations'],
    ['DOC-004', 'Database Schema Documentation', 'ER diagrams and table descriptions for MySQL schema'],
    ['DOC-005', 'Deployment Guide', 'Step-by-step deployment instructions for production environments'],
]
add_styled_table(doc, ['Reference ID', 'Document', 'Description'], doc_refs)

add_heading(doc, '14.2 Change Request Process', level=2)
add_para(doc,
    'Any changes to the requirements defined in this document must follow the formal change request process. '
    'Change requests must include the proposed modification, justification, impact analysis on existing '
    'features, estimated effort, and approval from the project sponsor and technical lead.')

add_para(doc,
    'Minor changes (typo fixes, clarifications) can be processed through the document revision process '
    'without formal change requests. Major changes (new requirements, scope modifications) require a '
    'formal change request reviewed and approved by the Change Control Board (CCB).')

add_heading(doc, '14.3 Risk Assessment', level=2)
risk_rows = [
    ['RSK-001', 'Face Recognition Accuracy', 'Medium', 'Mitigation: Implement confidence threshold and manual override option'],
    ['RSK-002', 'GPS Accuracy in Indoor Locations', 'Medium', 'Mitigation: Support WiFi-based location fallback and manual punch option'],
    ['RSK-003', 'Kafka Message Loss', 'Low', 'Mitigation: Configure persistent storage and replication factor >= 3'],
    ['RSK-004', 'Database Performance at Scale', 'Medium', 'Mitigation: Implement connection pooling, query optimization, and read replicas'],
    ['RSK-005', 'Mobile App Store Approval', 'Low', 'Mitigation: Follow platform guidelines strictly and submit early for review'],
]
add_styled_table(doc, ['Risk ID', 'Risk', 'Severity', 'Mitigation Strategy'], risk_rows)

add_heading(doc, '14.4 Future Enhancements', level=2)
add_para(doc,
    'The following features are planned for future releases but are not part of the current scope:')
future = [
    'Recruitment and Applicant Tracking System (ATS) integration',
    'Learning Management System (LMS) for employee training',
    'Advanced analytics with machine learning-powered insights',
    'Chatbot integration for HR query resolution',
    'Biometric device integration for attendance',
    'Multi-language support (internationalization)',
    'Advanced workflow designer for custom approval processes',
    'API marketplace for third-party integrations',
]
for item in future:
    add_bullet(doc, item)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('--- End of Document ---')
run.font.size = Pt(10)
run.font.color.rgb = GREY
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PIXOUS Technologies Pvt. Ltd. | Confidential')
run.font.size = Pt(9)
run.font.color.rgb = GREY

doc.save('PIXOUS_HR_Portal_Requirements_Document.docx')
print('Document generated: PIXOUS_HR_Portal_Requirements_Document.docx')
