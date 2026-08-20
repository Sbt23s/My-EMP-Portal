"""Generate PIXOUS HR Portal - Security Testing Document"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper: Add Logo ──
def add_logo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture('pixous_logo.png', width=Inches(2.0))
    return p

# ── Helper: Add Cover Title ──
def add_cover(doc):
    add_logo(doc)
    doc.add_paragraph()
    title = doc.add_heading('PIXOUS HR Portal', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading('Security Testing Document', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(0, 102, 153)
    
    doc.add_paragraph()
    
    # Info table
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ('Document Type', 'Security Testing Document'),
        ('Version', '1.0'),
        ('Date', datetime.date.today().strftime('%B %d, %Y')),
        ('Prepared By', 'PIXOUS Technologies Security Team'),
        ('Classification', 'Confidential - Client Deliverable'),
    ]
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for cell in t.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_page_break()

# ── Helper: Styled Table ──
def add_styled_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear'
        })
        shading.append(shd)
    # Rows
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
add_cover(doc)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Security Testing Scope',
    '3. Authentication Security',
    '4. Authorization & RBAC',
    '5. API Security',
    '6. Data Security',
    '7. Infrastructure Security',
    '8. WebSocket Security',
    '9. Mobile App Security',
    '10. Third-Party Integration Security',
    '11. OWASP Top 10 Compliance Checklist',
    '12. Security Test Cases',
    '13. Vulnerability Assessment Results Template',
    '14. Risk Matrix',
    '15. Security Recommendations',
    '16. Appendix: Security Scanning Tools',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal is a comprehensive enterprise-grade Human Resource Management System that '
    'handles sensitive employee data, financial information, and business-critical operations. This '
    'security testing document outlines the systematic approach to identifying, evaluating, and '
    'mitigating security vulnerabilities across all layers of the application.'
)
doc.add_paragraph(
    'Security testing covers authentication mechanisms, authorization controls, API security, data '
    'protection, infrastructure hardening, WebSocket security, mobile application security, and '
    'third-party integration safeguards. The goal is to ensure the PIXOUS HR Portal meets industry '
    'security standards and protects against common attack vectors.'
)
doc.add_paragraph(
    'This document follows the OWASP Testing Guide v4 methodology and aligns with OWASP Top 10 2021 '
    'risk categories. All testing activities are designed to validate both preventive and detective '
    'security controls within the system.'
)
doc.add_paragraph(
    'The testing scope includes the Spring Boot backend API, React frontend application, React Native '
    'and Flutter mobile applications, WebSocket communication layer, Docker-based deployment, and all '
    'third-party integrations including Twilio, Fast2SMS, Kafka, and Redis.'
)
doc.add_paragraph(
    'Key areas of focus include JWT token lifecycle management, role-based access control enforcement '
    'with 14+ permission codes, input validation and output encoding, data encryption at rest and in '
    'transit, and container security within the Docker deployment architecture.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. SECURITY TESTING SCOPE
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Security Testing Scope', level=1)
doc.add_paragraph(
    'The security testing scope encompasses all components of the PIXOUS HR Portal that process, store, '
    'or transmit sensitive data. This includes the backend API services, frontend web application, mobile '
    'applications, database layer, caching infrastructure, message queue, and deployment containers.'
)

doc.add_heading('2.1 In-Scope Components', level=2)
scope_items = [
    ['Spring Boot Backend API', 'REST API endpoints, Spring Security configuration, JWT handling'],
    ['React Frontend', 'Client-side authentication, token storage, input validation'],
    ['React Native / Flutter Mobile', 'Secure storage, certificate pinning, runtime protection'],
    ['MySQL Database', 'Data encryption, access controls, query parameterization'],
    ['Redis Cache', 'Session management, permission caching, data exposure risks'],
    ['Apache Kafka', 'Message security, topic access controls, data in transit'],
    ['WebSocket (STOMP)', 'Authentication, topic authorization, message integrity'],
    ['Docker Containers', 'Image security, network isolation, environment variables'],
    ['Third-Party APIs', 'Twilio, Fast2SMS, face recognition service integration'],
]
add_styled_table(doc, ['Component', 'Security Focus Areas'], scope_items)

doc.add_heading('2.2 Out-of-Scope Items', level=2)
out_scope = [
    'Physical security of PIXOUS Technologies office infrastructure',
    'Social engineering attacks targeting PIXOUS employees',
    'Denial of Service (DoS) testing against production systems',
    'Third-party infrastructure beyond PIXOUS control (AWS, GCP)',
]
for item in out_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. AUTHENTICATION SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Authentication Security', level=1)
doc.add_paragraph(
    'Authentication is the first line of defense in securing the PIXOUS HR Portal. The system implements '
    'JWT-based stateless authentication with multiple security layers including token rotation, account '
    'lockout, and BCrypt password hashing. This section details the testing approach for each authentication '
    'component.'
)

doc.add_heading('3.1 JWT Token Security', level=2)
doc.add_paragraph(
    'The PIXOUS HR Portal uses JSON Web Tokens (JWT) for stateless authentication. Access tokens are '
    'generated with a 4-hour time-to-live (TTL) and refresh tokens are used to obtain new access tokens '
    'without requiring user re-authentication.'
)
jwt_items = [
    ['Token Expiry', 'Validate access tokens expire after 4 hours', 'HIGH'],
    ['Token Rotation', 'Verify refresh tokens are rotated on each use', 'HIGH'],
    ['Token Revocation', 'Test logout revokes all refresh tokens for user', 'HIGH'],
    ['Token Tampering', 'Attempt to modify JWT payload and signature', 'CRITICAL'],
    ['Algorithm Confusion', 'Test for none algorithm or weak signing keys', 'CRITICAL'],
    ['Token Storage', 'Verify tokens stored securely in client-side storage', 'MEDIUM'],
    ['Token Leakage', 'Check tokens in URL parameters, logs, error messages', 'HIGH'],
    ['Cross-tenant Token Use', 'Attempt to use token from one tenant in another', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], jwt_items)

doc.add_heading('3.2 Password Policy', level=2)
doc.add_paragraph(
    'Passwords are hashed using BCrypt with a work factor of 12, providing strong protection against '
    'rainbow table and brute force attacks. The password policy enforces complexity requirements to '
    'ensure users create strong passwords.'
)
password_items = [
    ['Minimum Length', 'Verify minimum 8 characters enforced', 'MEDIUM'],
    ['Complexity Requirements', 'Test uppercase, lowercase, number, special char requirements', 'MEDIUM'],
    ['Password History', 'Verify last 5 passwords cannot be reused', 'MEDIUM'],
    ['BCrypt Work Factor', 'Confirm BCrypt work factor >= 10', 'HIGH'],
    ['Password in Transit', 'Verify HTTPS for all password transmissions', 'CRITICAL'],
    ['Password Storage', 'Confirm passwords never stored in plaintext', 'CRITICAL'],
    ['Reset Flow', 'Test password reset link expiry and single-use enforcement', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], password_items)

doc.add_heading('3.3 Account Lockout', level=2)
doc.add_paragraph(
    'The system implements account lockout after 5 consecutive failed login attempts. This prevents '
    'brute force attacks while allowing legitimate users who may have mistyped their password.'
)
lockout_items = [
    ['Lockout Threshold', 'Verify account locks after 5 failed attempts', 'HIGH'],
    ['Lockout Duration', 'Test automatic unlock after lockout period expires', 'MEDIUM'],
    ['Failed Attempt Counter', 'Verify counter resets after successful login', 'MEDIUM'],
    ['Lockout Notification', 'Test user notification on account lockout', 'LOW'],
    ['Admin Unlock', 'Verify admin can manually unlock accounts', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], lockout_items)

doc.add_heading('3.4 Session Management', level=2)
doc.add_paragraph(
    'Session management in the PIXOUS HR Portal relies on JWT tokens rather than server-side sessions. '
    'This section validates the implementation of token-based session controls.'
)
session_items = [
    ['Session Timeout', 'Verify tokens expire and require refresh', 'MEDIUM'],
    ['Concurrent Sessions', 'Test behavior with multiple active sessions', 'MEDIUM'],
    ['Session Invalidation', 'Verify logout invalidates all associated tokens', 'HIGH'],
    ['Session Fixation', 'Test that tokens are regenerated on privilege change', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], session_items)

doc.add_heading('3.5 Token Storage (Client-Side)', level=2)
doc.add_paragraph(
    'Client-side token storage is a critical security consideration. The React frontend stores JWT '
    'access tokens in memory and refresh tokens in HTTP-only cookies or secure storage mechanisms.'
)
storage_items = [
    ['localStorage Exposure', 'Verify tokens are not stored in localStorage', 'HIGH'],
    ['HTTP-only Cookies', 'Test refresh token cookie has HTTP-only flag', 'HIGH'],
    ['Secure Cookie Flag', 'Verify Secure flag is set on authentication cookies', 'HIGH'],
    ['SameSite Attribute', 'Test SameSite=Strict or Lax on auth cookies', 'MEDIUM'],
    ['Token in URL', 'Ensure tokens never appear in URL parameters', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], storage_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. AUTHORIZATION & RBAC
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Authorization & RBAC', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal implements Role-Based Access Control (RBAC) with 14+ distinct permission codes '
    'across multiple user roles. Authorization testing validates that users can only access resources and '
    'perform actions permitted by their assigned roles.'
)

doc.add_heading('4.1 Role-Based Access Control Testing', level=2)
doc.add_paragraph(
    'Each of the 14+ permission codes must be tested to verify proper enforcement. Permission codes include '
    'USER_MANAGE, LEAVE_APPLY, LEAVE_APPROVE, ATTENDANCE_SELF, ATTENDANCE_TEAM, PAYROLL_VIEW, PAYROLL_RUN, '
    'PAYROLL_APPROVE, ASSET_MANAGE, HELPDESK_RAISE, HELPDESK_AGENT, REPORT_VIEW, DASHBOARD_EXEC, '
    'ORG_MANAGE, and COMMUNITY_MANAGE.'
)
rbac_items = [
    ['USER_MANAGE', 'HR, Admin', 'Verify non-HR users cannot create/edit employees', 'HIGH'],
    ['LEAVE_APPLY', 'All Employees', 'Verify all authenticated users can apply for leave', 'MEDIUM'],
    ['LEAVE_APPROVE', 'Manager, HR, Admin', 'Verify only authorized users can approve leave', 'HIGH'],
    ['ATTENDANCE_SELF', 'All Employees', 'Verify users can only punch own attendance', 'MEDIUM'],
    ['ATTENDANCE_TEAM', 'Manager, HR, Admin', 'Verify team leads see only their team data', 'HIGH'],
    ['PAYROLL_VIEW', 'Employee (own), HR, Admin', 'Verify employees see only own payslips', 'CRITICAL'],
    ['PAYROLL_RUN', 'HR, Admin', 'Verify only HR/Admin can run payroll', 'CRITICAL'],
    ['PAYROLL_APPROVE', 'Finance, Admin', 'Verify only Finance can approve payroll runs', 'CRITICAL'],
    ['ASSET_MANAGE', 'Asset Manager, Admin', 'Verify only authorized roles manage assets', 'HIGH'],
    ['HELPDESK_RAISE', 'All Employees', 'Verify all users can raise support tickets', 'MEDIUM'],
    ['HELPDESK_AGENT', 'Asset Manager, Admin', 'Verify only agents can resolve tickets', 'HIGH'],
    ['REPORT_VIEW', 'Manager, HR, Admin', 'Verify report access is role-restricted', 'HIGH'],
    ['DASHBOARD_EXEC', 'CEO, Admin', 'Verify executive dashboard is role-restricted', 'HIGH'],
    ['ORG_MANAGE', 'HR, Admin', 'Verify org config is admin-only', 'HIGH'],
    ['COMMUNITY_MANAGE', 'HR, Admin', 'Verify chat group management is restricted', 'MEDIUM'],
]
add_styled_table(doc, ['Permission Code', 'Expected Roles', 'Test Description', 'Severity'], rbac_items)

doc.add_heading('4.2 Vertical Privilege Escalation', level=2)
doc.add_paragraph(
    'Vertical privilege escalation testing verifies that low-privilege users cannot elevate their access '
    'level to perform administrative functions. This includes testing API endpoint access with different '
    'user roles and attempting to modify role assignments.'
)
escalation_items = [
    ['Employee to Admin', 'Attempt admin API calls with employee JWT token', 'CRITICAL'],
    ['Employee to HR', 'Attempt HR operations (create user) with employee token', 'CRITICAL'],
    ['Role Modification', 'Attempt to modify own role via API', 'CRITICAL'],
    ['Permission Injection', 'Attempt to add unauthorized permissions to JWT', 'CRITICAL'],
    ['Hidden Endpoints', 'Enumerate and test undocumented admin endpoints', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], escalation_items)

doc.add_heading('4.3 Horizontal Privilege Escalation', level=2)
doc.add_paragraph(
    'Horizontal privilege escalation testing verifies that users cannot access data or perform actions '
    'on behalf of other users at the same privilege level. This is particularly important for operations '
    'like viewing payslips, attendance records, and personal information.'
)
horizontal_items = [
    ['User A accessing User B payslips', 'Modify userId in API request to access another user data', 'CRITICAL'],
    ['User A updating User B profile', 'Attempt profile update with different userId', 'CRITICAL'],
    ['Cross-team data access', 'Team lead accessing another team lead\'s team data', 'HIGH'],
    ['Employee viewing other employee attendance', 'Modify employeeId in attendance API', 'HIGH'],
    ['Ticket access across users', 'Attempt to view tickets raised by other employees', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], horizontal_items)

doc.add_heading('4.4 Permission Boundary Testing', level=2)
doc.add_paragraph(
    'Boundary testing validates edge cases where permission boundaries might be unclear or incorrectly '
    'enforced. This includes testing with expired tokens, tokens from deleted users, and tokens with '
    'modified permission claims.'
)
boundary_items = [
    ['Expired Token Access', 'Attempt API calls with expired JWT token', 'HIGH'],
    ['Deleted User Token', 'Use token from a deactivated user account', 'HIGH'],
    ['Modified Permissions', 'Alter permission claims in JWT payload', 'CRITICAL'],
    ['Missing Permission', 'Test API call without required permission in token', 'HIGH'],
    ['Wildcard Permissions', 'Test for wildcard or overly broad permission patterns', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], boundary_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. API SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. API Security', level=1)
doc.add_paragraph(
    'API security testing covers all REST API endpoints exposed by the Spring Boot backend. This includes '
    'input validation, injection attacks, cross-site scripting, CSRF protection, rate limiting, CORS '
    'configuration, and endpoint enumeration protection.'
)

doc.add_heading('5.1 Input Validation and Sanitization', level=2)
doc.add_paragraph(
    'All user-supplied input must be validated and sanitized before processing. The PIXOUS HR Portal uses '
    'Spring Boot validation annotations and custom validators to enforce input constraints.'
)
input_items = [
    ['Boundary Values', 'Test minimum and maximum length inputs for all fields', 'MEDIUM'],
    ['Special Characters', 'Input special characters in text fields', 'MEDIUM'],
    ['Unicode Handling', 'Test Unicode characters in names, addresses', 'LOW'],
    ['Null/Empty Input', 'Submit null or empty values where fields are required', 'MEDIUM'],
    ['Type Mismatch', 'Submit string where integer is expected', 'MEDIUM'],
    ['File Upload Validation', 'Test file type, size, and content validation', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], input_items)

doc.add_heading('5.2 SQL Injection Testing', level=2)
doc.add_paragraph(
    'SQL injection testing validates that all database queries use parameterized statements or prepared '
    'statements to prevent SQL injection attacks. The PIXOUS HR Portal uses Spring Data JPA which provides '
    'built-in SQL injection protection when used correctly.'
)
sqli_items = [
    ['Login Bypass', 'Attempt SQL injection in login form fields', 'CRITICAL'],
    ['Search Fields', 'Inject SQL in search and filter parameters', 'HIGH'],
    ['Sort Parameters', 'Test SQL injection via sort/order parameters', 'HIGH'],
    ['WHERE Clause Injection', 'Attempt to modify WHERE clauses via input', 'CRITICAL'],
    [' UNION SELECT', 'Test UNION-based SQL injection on all endpoints', 'CRITICAL'],
    ['Blind SQL Injection', 'Test time-based and boolean-based blind injection', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], sqli_items)

doc.add_heading('5.3 XSS (Cross-Site Scripting) Prevention', level=2)
doc.add_paragraph(
    'Cross-Site Scripting testing verifies that user input is properly encoded when rendered in HTML, '
    'JavaScript, or URL contexts. The React frontend provides automatic XSS protection through JSX '
    'encoding, but server-side output encoding must also be validated.'
)
xss_items = [
    ['Reflected XSS', 'Inject script in URL parameters reflected in page', 'HIGH'],
    ['Stored XSS', 'Submit script in fields displayed to other users', 'CRITICAL'],
    ['DOM-based XSS', 'Test client-side JavaScript for DOM manipulation vulnerabilities', 'HIGH'],
    ['Event Handler Injection', 'Attempt to inject event handlers in input fields', 'MEDIUM'],
    ['SVG/Image XSS', 'Test SVG upload or image metadata for script injection', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], xss_items)

doc.add_heading('5.4 CSRF (Cross-Site Request Forgery) Protection', level=2)
doc.add_paragraph(
    'CSRF protection ensures that requests to the application originate from legitimate users and not '
    'from malicious third-party sites. The PIXOUS HR Portal uses state-changing operations that must '
    'be protected against CSRF attacks.'
)
csrf_items = [
    ['Token Validation', 'Verify CSRF tokens on state-changing requests', 'HIGH'],
    ['SameSite Cookies', 'Verify SameSite attribute on session cookies', 'MEDIUM'],
    ['Origin Validation', 'Test Origin and Referer header validation', 'HIGH'],
    ['Preflight Requests', 'Verify CORS preflight for complex requests', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], csrf_items)

doc.add_heading('5.5 Rate Limiting Verification', level=2)
doc.add_paragraph(
    'Rate limiting prevents brute force attacks and API abuse by limiting the number of requests a client '
    'can make within a specified time window. The PIXOUS HR Portal implements rate limiting at the API '
    'gateway level.'
)
rate_items = [
    ['Login Rate Limit', 'Verify rate limiting on authentication endpoints', 'HIGH'],
    ['Password Reset Rate Limit', 'Test rate limiting on password reset requests', 'HIGH'],
    ['API Rate Limit', 'Validate general API rate limiting configuration', 'MEDIUM'],
    ['Burst Request Handling', 'Test behavior under burst traffic conditions', 'MEDIUM'],
    ['Rate Limit Bypass', 'Attempt to bypass rate limiting via header manipulation', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], rate_items)

doc.add_heading('5.6 CORS Policy Testing', level=2)
doc.add_paragraph(
    'Cross-Origin Resource Sharing (CORS) configuration must be restrictive to prevent unauthorized '
    'cross-origin access to API endpoints while allowing legitimate frontend requests.'
)
cors_items = [
    ['Allowed Origins', 'Verify only trusted origins are in CORS allowlist', 'HIGH'],
    ['Credentials Handling', 'Test Access-Control-Allow-Credentials configuration', 'HIGH'],
    ['Methods Restriction', 'Verify only necessary HTTP methods are allowed', 'MEDIUM'],
    ['Headers Restriction', 'Test allowed headers are appropriately restricted', 'MEDIUM'],
    ['Wildcard Prevention', 'Ensure Access-Control-Allow-Origin is not wildcard (*)', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], cors_items)

doc.add_heading('5.7 API Endpoint Enumeration Protection', level=2)
doc.add_paragraph(
    'API endpoint enumeration protection ensures that the application does not expose unnecessary '
    'information about available endpoints, their parameters, or internal implementation details.'
)
enum_items = [
    ['Error Messages', 'Verify error messages do not reveal internal paths', 'MEDIUM'],
    ['Swagger/OpenAPI', 'Ensure API docs are not publicly accessible in production', 'HIGH'],
    ['Directory Traversal', 'Test for exposed directory listings', 'HIGH'],
    ['Debug Endpoints', 'Verify debug/actuator endpoints are not exposed', 'HIGH'],
    ['HTTP Methods', 'Test 405 responses for unsupported HTTP methods', 'LOW'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], enum_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. DATA SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Data Security', level=1)
doc.add_paragraph(
    'Data security testing validates that sensitive information is properly protected at rest, in transit, '
    'and during processing. The PIXOUS HR Portal handles PII, salary data, passwords, and file uploads '
    'that require specific security controls.'
)

doc.add_heading('6.1 Data Encryption at Rest (MySQL)', level=2)
doc.add_paragraph(
    'The MySQL database stores employee records, salary data, attendance logs, and other sensitive '
    'information. Encryption at rest ensures data is protected even if physical storage is compromised.'
)
encryption_rest = [
    ['TDE Configuration', 'Verify Transparent Data Encryption is enabled on MySQL', 'HIGH'],
    ['Column-level Encryption', 'Test encryption of sensitive columns (SSN, bank details)', 'HIGH'],
    ['Backup Encryption', 'Verify database backups are encrypted', 'HIGH'],
    ['Key Management', 'Validate encryption key rotation and storage practices', 'HIGH'],
    ['Data Classification', 'Confirm all sensitive fields are classified and encrypted', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], encryption_rest)

doc.add_heading('6.2 Data Encryption in Transit (HTTPS/TLS)', level=2)
doc.add_paragraph(
    'All data transmitted between client and server must be encrypted using TLS 1.2 or higher. This '
    'includes API requests, WebSocket connections, and mobile app communications.'
)
encryption_transit = [
    ['TLS Version', 'Verify TLS 1.2 or higher is enforced', 'HIGH'],
    ['Certificate Validity', 'Test SSL certificate expiration and chain validation', 'HIGH'],
    ['HSTS Header', 'Verify Strict-Transport-Security header is present', 'MEDIUM'],
    ['Mixed Content', 'Test for mixed HTTP/HTTPS content loading', 'MEDIUM'],
    ['Cipher Suites', 'Verify strong cipher suites are configured', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], encryption_transit)

doc.add_heading('6.3 PII (Personally Identifiable Information) Protection', level=2)
doc.add_paragraph(
    'PII protection ensures that personal information such as employee names, addresses, phone numbers, '
    'email addresses, and government IDs are handled according to privacy regulations and best practices.'
)
pii_items = [
    ['Data Minimization', 'Verify only necessary PII is collected and stored', 'MEDIUM'],
    ['Access Controls', 'Test PII access is limited to authorized roles', 'HIGH'],
    ['Data Masking', 'Verify PII is masked in logs and error messages', 'HIGH'],
    ['Retention Policies', 'Test data retention and deletion policies', 'MEDIUM'],
    ['Consent Management', 'Verify user consent is recorded where required', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], pii_items)

doc.add_heading('6.4 Salary Data Confidentiality', level=2)
doc.add_paragraph(
    'Salary data is highly sensitive and must be accessible only to authorized personnel such as HR '
    'managers, finance officers, and the employees themselves. Access controls must be strictly enforced.'
)
salary_items = [
    ['Role-based Access', 'Verify only HR, Finance, Admin can view salary data', 'CRITICAL'],
    ['Employee Self-Access', 'Verify employees can only view their own salary', 'CRITICAL'],
    ['API Authorization', 'Test salary endpoints reject unauthorized requests', 'CRITICAL'],
    ['Data Exposure', 'Check salary data is not exposed in logs or analytics', 'HIGH'],
    ['Export Controls', 'Verify salary export requires appropriate permissions', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], salary_items)

doc.add_heading('6.5 Password Vault Security', level=2)
doc.add_paragraph(
    'The PIXOUS HR Portal stores passwords using BCrypt hashing. The password vault security testing '
    'validates that passwords are properly protected throughout their lifecycle.'
)
vault_items = [
    ['BCrypt Work Factor', 'Verify BCrypt work factor is appropriate (>= 10)', 'HIGH'],
    ['Salt Generation', 'Confirm unique salt is generated for each password', 'HIGH'],
    ['Hash Storage', 'Verify hashes are stored securely in database', 'HIGH'],
    ['Password Transmission', 'Test passwords are only sent over HTTPS', 'CRITICAL'],
    ['Password Logging', 'Ensure passwords are never logged', 'CRITICAL'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], vault_items)

doc.add_heading('6.6 File Upload Security', level=2)
doc.add_paragraph(
    'The PIXOUS HR Portal supports file uploads for profile photos, attachments, and documents. File '
    'upload security testing validates that uploads are properly validated and stored securely.'
)
file_items = [
    ['File Type Validation', 'Verify allowed file types are enforced', 'HIGH'],
    ['File Size Limits', 'Test maximum file size limits are enforced', 'MEDIUM'],
    ['Malicious File Detection', 'Test upload of executable or script files', 'HIGH'],
    ['File Storage Location', 'Verify uploaded files are not web-accessible directly', 'HIGH'],
    ['Filename Sanitization', 'Test filename sanitization to prevent path traversal', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], file_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. INFRASTRUCTURE SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Infrastructure Security', level=1)
doc.add_paragraph(
    'Infrastructure security testing validates the security of the Docker deployment, network configuration, '
    'SSL/TLS certificates, port exposure, and environment variable management.'
)

doc.add_heading('7.1 Docker Container Security', level=2)
doc.add_paragraph(
    'The PIXOUS HR Portal is deployed using Docker containers. Container security testing validates that '
    'containers are properly configured and do not expose unnecessary risks.'
)
docker_items = [
    ['Base Image Security', 'Verify containers use minimal, up-to-date base images', 'HIGH'],
    ['Non-root User', 'Confirm containers run as non-root user', 'HIGH'],
    ['Read-only Filesystem', 'Test container filesystem is read-only where possible', 'MEDIUM'],
    ['Resource Limits', 'Verify CPU and memory limits are set', 'MEDIUM'],
    ['Secret Management', 'Test secrets are not hardcoded in Dockerfiles', 'CRITICAL'],
    ['Image Scanning', 'Validate Docker images are scanned for vulnerabilities', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], docker_items)

doc.add_heading('7.2 Network Configuration', level=2)
doc.add_paragraph(
    'Network configuration testing validates that internal services (Redis, Kafka, MySQL) are not '
    'exposed to the public internet and that network segmentation is properly implemented.'
)
network_items = [
    ['Redis Exposure', 'Verify Redis is not exposed on public ports', 'CRITICAL'],
    ['Kafka Exposure', 'Verify Kafka is not exposed on public ports', 'CRITICAL'],
    ['MySQL Exposure', 'Verify MySQL is not exposed on public ports', 'CRITICAL'],
    ['Internal Communication', 'Test internal services communicate on private network', 'HIGH'],
    ['Firewall Rules', 'Validate firewall rules restrict unnecessary access', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], network_items)

doc.add_heading('7.3 SSL/TLS Certificate Validation', level=2)
doc.add_paragraph(
    'SSL/TLS certificate validation ensures that the application uses valid, properly configured '
    'certificates for encrypted communication.'
)
ssl_items = [
    ['Certificate Expiry', 'Verify certificates are not expired', 'HIGH'],
    ['Certificate Chain', 'Test complete certificate chain is validated', 'HIGH'],
    ['Hostname Verification', 'Verify hostname matches certificate CN/SAN', 'HIGH'],
    ['Weak Certificates', 'Test for weak key sizes or deprecated algorithms', 'HIGH'],
    ['OCSP Stapling', 'Verify OCSP stapling is configured if used', 'LOW'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], ssl_items)

doc.add_heading('7.4 Port Scanning Resistance', level=2)
doc.add_paragraph(
    'Port scanning resistance testing validates that the application does not expose unnecessary ports '
    'and responds appropriately to port scanning attempts.'
)
port_items = [
    ['Open Port Audit', 'Identify all open ports on production servers', 'HIGH'],
    ['Unnecessary Services', 'Verify unnecessary services are disabled', 'MEDIUM'],
    ['Banner Grabbing', 'Test for information disclosure in service banners', 'MEDIUM'],
    ['Port Knocking', 'Evaluate if port knocking could enhance security', 'LOW'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], port_items)

doc.add_heading('7.5 Environment Variable Security', level=2)
doc.add_paragraph(
    'Environment variable security testing validates that sensitive configuration values are properly '
    'managed and not exposed through application endpoints or logs.'
)
env_items = [
    ['.env File Exposure', 'Verify .env files are not web-accessible', 'CRITICAL'],
    ['Secret Values', 'Test secrets are not in application logs', 'HIGH'],
    ['Config Exposure', 'Verify sensitive config not in error responses', 'HIGH'],
    ['Key Rotation', 'Test ability to rotate secrets without downtime', 'MEDIUM'],
    ['Default Credentials', 'Ensure no default passwords in deployment', 'CRITICAL'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], env_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. WEBSOCKET SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. WebSocket Security', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal uses WebSocket (STOMP/SockJS) for real-time notifications and chat functionality. '
    'WebSocket security testing validates authentication, authorization, and message integrity.'
)

doc.add_heading('8.1 STOMP Authentication', level=2)
doc.add_paragraph(
    'STOMP protocol authentication ensures that WebSocket connections are properly authenticated before '
    'allowing subscription to topics or sending messages.'
)
stomp_items = [
    ['Connection Authentication', 'Verify JWT token required for STOMP connection', 'HIGH'],
    ['Token Validation', 'Test invalid/expired tokens are rejected', 'HIGH'],
    ['Reconnection Auth', 'Verify authentication is enforced on reconnection', 'MEDIUM'],
    ['Frame Validation', 'Test STOMP frame parsing for malformed input', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], stomp_items)

doc.add_heading('8.2 Topic Subscription Authorization', level=2)
doc.add_paragraph(
    'Topic subscription authorization validates that users can only subscribe to topics they are '
    'authorized to receive. This prevents unauthorized access to notification channels.'
)
topic_items = [
    ['User-specific Topics', 'Verify users can only subscribe to own notification topics', 'HIGH'],
    ['Team Topics', 'Test team lead can only subscribe to own team topics', 'HIGH'],
    ['Admin Topics', 'Verify admin-only topics reject non-admin subscriptions', 'HIGH'],
    ['Topic Enumeration', 'Test for topic enumeration via subscription attempts', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], topic_items)

doc.add_heading('8.3 Message Integrity', level=2)
doc.add_paragraph(
    'Message integrity testing validates that WebSocket messages are not tampered with during transit '
    'and that message payloads are properly validated.'
)
message_items = [
    ['Message Tampering', 'Test if messages can be modified in transit', 'HIGH'],
    ['Payload Validation', 'Verify server validates incoming message payloads', 'MEDIUM'],
    ['Message Replay', 'Test for message replay attack resistance', 'MEDIUM'],
    ['Rate Limiting', 'Verify WebSocket connections have rate limiting', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], message_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. MOBILE APP SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Mobile App Security', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal mobile applications (React Native and Flutter) handle sensitive data and must '
    'implement additional security controls for the mobile environment.'
)

doc.add_heading('9.1 Secure Storage', level=2)
doc.add_paragraph(
    'Secure storage testing validates that sensitive data such as authentication tokens and user credentials '
    'are stored using platform-provided secure storage mechanisms.'
)
storage_items = [
    ['Expo Secure Store', 'Verify React Native uses expo-secure-store for tokens', 'HIGH'],
    ['Flutter Secure Storage', 'Verify Flutter uses flutter_secure_storage', 'HIGH'],
    ['Keychain/iOS KeyStore', 'Test platform-native keychain integration', 'HIGH'],
    ['SharedPreferences', 'Ensure sensitive data not in SharedPreferences', 'HIGH'],
    ['Data-at-rest Encryption', 'Verify local data is encrypted', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], storage_items)

doc.add_heading('9.2 Certificate Pinning', level=2)
doc.add_paragraph(
    'Certificate pinning prevents man-in-the-middle attacks by validating that the server certificate '
    'matches a known certificate fingerprint.'
)
cert_items = [
    ['Pin Validation', 'Verify certificate pinning is implemented', 'HIGH'],
    ['Pin Bypass', 'Attempt to bypass certificate pinning', 'HIGH'],
    ['Backup Pins', 'Test backup pin configuration for certificate rotation', 'MEDIUM'],
    ['Pin Failure Handling', 'Test behavior when pin validation fails', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], cert_items)

doc.add_heading('9.3 Runtime Application Self-Protection', level=2)
doc.add_paragraph(
    'Runtime application self-protection (RASP) testing validates that the mobile application can detect '
    'and respond to security threats at runtime.'
)
rasp_items = [
    ['Jailbreak Detection', 'Verify app detects jailbroken/rooted devices', 'MEDIUM'],
    ['Debugger Detection', 'Test for debugger detection mechanisms', 'MEDIUM'],
    ['Code Obfuscation', 'Verify binary is obfuscated to prevent reverse engineering', 'MEDIUM'],
    ['Tamper Detection', 'Test app integrity verification', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], rasp_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. THIRD-PARTY INTEGRATION SECURITY
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Third-Party Integration Security', level=1)
doc.add_paragraph(
    'The PIXOUS HR Portal integrates with several third-party services that require specific security '
    'controls to protect API keys, sensitive data, and communication channels.'
)

doc.add_heading('10.1 Twilio/Fast2SMS API Key Security', level=2)
doc.add_paragraph(
    'SMS service API keys must be securely stored and never exposed in client-side code or logs. '
    'This testing validates proper API key management for Twilio and Fast2SMS integrations.'
)
sms_items = [
    ['Key Storage', 'Verify API keys stored in server-side environment variables', 'HIGH'],
    ['Key Rotation', 'Test ability to rotate API keys without downtime', 'MEDIUM'],
    ['Key Exposure', 'Ensure API keys not in client-side code or logs', 'CRITICAL'],
    ['Rate Limiting', 'Verify SMS sending is rate-limited per user', 'MEDIUM'],
    ['Error Handling', 'Test error responses do not expose API key details', 'HIGH'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], sms_items)

doc.add_heading('10.2 Face Recognition Data Privacy', level=2)
doc.add_paragraph(
    'Face recognition data is highly sensitive biometric information. Privacy testing validates that '
    'face data is handled according to privacy regulations and best practices.'
)
face_items = [
    ['Consent Collection', 'Verify explicit consent is collected for face data', 'HIGH'],
    ['Data Minimization', 'Test that only necessary face data is stored', 'HIGH'],
    ['Storage Encryption', 'Verify face embeddings are encrypted at rest', 'HIGH'],
    ['Access Controls', 'Test face data access is limited to attendance module', 'HIGH'],
    ['Deletion Rights', 'Verify face data can be deleted on user request', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], face_items)

doc.add_heading('10.3 Kafka Message Security', level=2)
doc.add_paragraph(
    'Apache Kafka is used for chat event streaming. Kafka security testing validates that messages are '
    'properly secured in transit and that topic access is properly controlled.'
)
kafka_items = [
    ['TLS Encryption', 'Verify Kafka communication uses TLS encryption', 'HIGH'],
    ['Authentication', 'Test Kafka broker authentication is enforced', 'HIGH'],
    ['Authorization', 'Verify topic-level authorization is configured', 'HIGH'],
    ['Message Encryption', 'Test message payload encryption if required', 'MEDIUM'],
    ['Access Logging', 'Verify Kafka access is logged for audit purposes', 'MEDIUM'],
]
add_styled_table(doc, ['Test Case', 'Description', 'Severity'], kafka_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. OWASP TOP 10 COMPLIANCE CHECKLIST
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. OWASP Top 10 Compliance Checklist', level=1)
doc.add_paragraph(
    'This section maps PIXOUS HR Portal security testing to the OWASP Top 10 2021 risk categories. '
    'Each category is evaluated against the application\'s security controls.'
)

owasp_items = [
    ['A01:2021', 'Broken Access Control', 'RBAC testing, privilege escalation, horizontal access', 'Testing Required'],
    ['A02:2021', 'Cryptographic Failures', 'Encryption at rest/transit, password hashing', 'Testing Required'],
    ['A03:2021', 'Injection', 'SQL injection, XSS, command injection', 'Testing Required'],
    ['A04:2021', 'Insecure Design', 'Threat modeling, security architecture review', 'Testing Required'],
    ['A05:2021', 'Security Misconfiguration', 'Default configs, unnecessary features, error handling', 'Testing Required'],
    ['A06:2021', 'Vulnerable Components', 'Dependency scanning, version management', 'Testing Required'],
    ['A07:2021', 'Auth Failures', 'Brute force, session management, credential storage', 'Testing Required'],
    ['A08:2021', 'Data Integrity Failures', 'CI/CD pipeline security, software updates', 'Testing Required'],
    ['A09:2021', 'Logging Failures', 'Audit logging, intrusion detection, log protection', 'Testing Required'],
    ['A10:2021', 'SSRF', 'Server-side request forgery prevention', 'Testing Required'],
]
add_styled_table(doc, ['Category', 'Risk', 'Testing Focus', 'Status'], owasp_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. SECURITY TEST CASES
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. Security Test Cases', level=1)
doc.add_paragraph(
    'This section provides detailed security test cases with unique identifiers for traceability. '
    'Each test case includes preconditions, test steps, expected results, and pass/fail criteria.'
)

tc_items = [
    ['SEC-AUTH-001', 'JWT Token Expiry Validation', 'Verify access tokens expire after 4 hours', 'CRITICAL', ''],
    ['SEC-AUTH-002', 'Refresh Token Rotation', 'Verify refresh tokens are rotated on use', 'HIGH', ''],
    ['SEC-AUTH-003', 'Account Lockout (5 attempts)', 'Verify account locks after 5 failed logins', 'HIGH', ''],
    ['SEC-AUTH-004', 'BCrypt Password Hashing', 'Verify passwords are BCrypt hashed', 'CRITICAL', ''],
    ['SEC-AUTH-005', 'Password Complexity Enforcement', 'Test minimum 8 char with complexity rules', 'MEDIUM', ''],
    ['SEC-AUTH-006', 'JWT Token Tampering', 'Attempt to modify JWT payload/signature', 'CRITICAL', ''],
    ['SEC-AUTH-007', 'Token Revocation on Logout', 'Verify all tokens revoked on logout', 'HIGH', ''],
    ['SEC-RBAC-001', 'Vertical Privilege Escalation', 'Employee attempts admin operations', 'CRITICAL', ''],
    ['SEC-RBAC-002', 'Horizontal Privilege Escalation', 'User A accesses User B data', 'CRITICAL', ''],
    ['SEC-RBAC-003', 'Permission Code Enforcement', 'Test each of 14+ permission codes', 'HIGH', ''],
    ['SEC-RBAC-004', 'Role Boundary Testing', 'Test cross-role data access attempts', 'HIGH', ''],
    ['SEC-API-001', 'SQL Injection - Login', 'SQL injection in login form fields', 'CRITICAL', ''],
    ['SEC-API-002', 'SQL Injection - Search', 'SQL injection in search/filter parameters', 'HIGH', ''],
    ['SEC-API-003', 'XSS - Stored', 'Submit script in user profile fields', 'CRITICAL', ''],
    ['SEC-API-004', 'XSS - Reflected', 'Inject script in URL parameters', 'HIGH', ''],
    ['SEC-API-005', 'CSRF Token Validation', 'Test CSRF protection on state changes', 'HIGH', ''],
    ['SEC-API-006', 'Rate Limiting - Login', 'Verify login rate limiting works', 'HIGH', ''],
    ['SEC-API-007', 'CORS Policy Validation', 'Test CORS headers and restrictions', 'MEDIUM', ''],
    ['SEC-DATA-001', 'Encryption at Rest', 'Verify MySQL data encryption enabled', 'HIGH', ''],
    ['SEC-DATA-002', 'Encryption in Transit', 'Verify TLS 1.2+ on all connections', 'HIGH', ''],
    ['SEC-DATA-003', 'PII Data Exposure', 'Check PII not in logs or errors', 'HIGH', ''],
    ['SEC-DATA-004', 'Salary Data Access', 'Test salary data role restrictions', 'CRITICAL', ''],
    ['SEC-DATA-005', 'File Upload Validation', 'Test malicious file upload prevention', 'HIGH', ''],
    ['SEC-INFRA-001', 'Docker Security', 'Verify container security configuration', 'HIGH', ''],
    ['SEC-INFRA-002', 'Network Exposure', 'Test internal services not public', 'CRITICAL', ''],
    ['SEC-INFRA-003', 'SSL Certificate Validation', 'Verify certificate chain and expiry', 'HIGH', ''],
    ['SEC-INFRA-004', '.env File Exposure', 'Verify .env not web-accessible', 'CRITICAL', ''],
    ['SEC-WS-001', 'STOMP Authentication', 'Verify WebSocket requires authentication', 'HIGH', ''],
    ['SEC-WS-002', 'Topic Authorization', 'Test topic subscription restrictions', 'HIGH', ''],
    ['SEC-MOB-001', 'Secure Storage', 'Verify expo-secure-store/flutter_secure_storage', 'HIGH', ''],
    ['SEC-MOB-002', 'Certificate Pinning', 'Test certificate pinning implementation', 'HIGH', ''],
    ['SEC-3RD-001', 'SMS API Key Security', 'Verify API keys not exposed', 'HIGH', ''],
    ['SEC-3RD-002', 'Face Data Privacy', 'Test biometric data protection', 'HIGH', ''],
    ['SEC-3RD-003', 'Kafka Message Security', 'Verify Kafka encryption and auth', 'HIGH', ''],
]
add_styled_table(doc, ['TC ID', 'Test Case Name', 'Description', 'Severity', 'Result'], tc_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. VULNERABILITY ASSESSMENT RESULTS TEMPLATE
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. Vulnerability Assessment Results Template', level=1)
doc.add_paragraph(
    'This section provides a template for recording vulnerability assessment results. Each finding '
    'should be documented with the following fields for proper tracking and remediation.'
)

vuln_template = [
    ['Finding ID', 'Unique identifier for the vulnerability (e.g., VLN-001)'],
    ['Title', 'Brief description of the vulnerability'],
    ['Severity', 'Critical / High / Medium / Low / Informational'],
    ['CVSS Score', 'Common Vulnerability Scoring System score (0-10)'],
    ['Affected Component', 'API endpoint, module, or infrastructure component'],
    ['Description', 'Detailed description of the vulnerability'],
    ['Steps to Reproduce', 'Step-by-step instructions to reproduce the issue'],
    ['Evidence', 'Screenshots, logs, or proof-of-concept details'],
    ['Impact', 'Potential impact if exploited'],
    ['Remediation', 'Recommended fix or mitigation strategy'],
    ['Status', 'Open / In Progress / Remediated / False Positive'],
    ['Date Found', 'Date the vulnerability was identified'],
    ['Date Remediated', 'Date the vulnerability was fixed'],
]
add_styled_table(doc, ['Field', 'Description'], vuln_template)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. RISK MATRIX
# ══════════════════════════════════════════════════════════════
doc.add_heading('14. Risk Matrix', level=1)
doc.add_paragraph(
    'The risk matrix below maps vulnerability likelihood against potential impact to determine overall '
    'risk severity. This helps prioritize remediation efforts for the most critical findings.'
)

risk_matrix = [
    ['Critical (5)', 'Likely + Catastrophic', 'Immediate remediation required', 'Exploit chain possible'],
    ['High (4)', 'Likely + Major OR Possible + Catastrophic', 'Remediate within 1 week', 'Active exploitation risk'],
    ['Medium (3)', 'Possible + Major OR Likely + Moderate', 'Remediate within 30 days', 'Potential for exploitation'],
    ['Low (2)', 'Possible + Moderate OR Unlikely + Major', 'Remediate within 90 days', 'Limited exploitation risk'],
    ['Informational (1)', 'Unlikely + Moderate or lower', 'Best practice improvement', 'No immediate risk'],
]
add_styled_table(doc, ['Risk Level', 'Criteria', 'Remediation SLA', 'Notes'], risk_matrix)

doc.add_heading('14.1 Likelihood Scale', level=2)
likelihood = [
    ['5 - Very High', 'Almost certain to be exploited'],
    ['4 - High', 'Likely to be exploited'],
    ['3 - Medium', 'Possible to be exploited'],
    ['2 - Low', 'Unlikely to be exploited'],
    ['1 - Very Low', 'Rarely exploitable'],
]
add_styled_table(doc, ['Rating', 'Description'], likelihood)

doc.add_heading('14.2 Impact Scale', level=2)
impact = [
    ['5 - Catastrophic', 'Complete system compromise, data breach, regulatory violation'],
    ['4 - Major', 'Significant data exposure, service disruption'],
    ['3 - Moderate', 'Limited data exposure, partial service impact'],
    ['2 - Minor', 'Minimal impact, cosmetic issues'],
    ['1 - Negligible', 'No meaningful impact on security posture'],
]
add_styled_table(doc, ['Rating', 'Description'], impact)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 15. SECURITY RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════
doc.add_heading('15. Security Recommendations', level=1)
doc.add_paragraph(
    'Based on the PIXOUS HR Portal security testing scope, the following recommendations are provided '
    'to enhance the overall security posture of the application.'
)

doc.add_heading('15.1 Authentication Recommendations', level=2)
auth_recs = [
    ['Implement MFA', 'Add multi-factor authentication for admin and HR roles', 'HIGH'],
    ['Password Rotation', 'Implement mandatory password rotation policy', 'MEDIUM'],
    ['Session Management', 'Implement session timeout and concurrent session limits', 'MEDIUM'],
    ['Login Monitoring', 'Add real-time login attempt monitoring and alerting', 'MEDIUM'],
]
add_styled_table(doc, ['Recommendation', 'Description', 'Priority'], auth_recs)

doc.add_heading('15.2 Authorization Recommendations', level=2)
authz_recs = [
    ['RBAC Audit', 'Conduct quarterly RBAC permission review', 'HIGH'],
    ['Least Privilege', 'Implement principle of least privilege across all roles', 'HIGH'],
    ['Audit Logging', 'Enhance audit logging for all permission-sensitive operations', 'HIGH'],
    ['Access Reviews', 'Implement periodic access reviews for high-privilege roles', 'MEDIUM'],
]
add_styled_table(doc, ['Recommendation', 'Description', 'Priority'], authz_recs)

doc.add_heading('15.3 Data Protection Recommendations', level=2)
data_recs = [
    ['Data Classification', 'Implement formal data classification policy', 'HIGH'],
    ['Encryption Enhancement', 'Add field-level encryption for additional PII fields', 'MEDIUM'],
    ['Data Retention', 'Implement automated data retention and deletion', 'MEDIUM'],
    ['Privacy Controls', 'Add user consent management for data collection', 'MEDIUM'],
]
add_styled_table(doc, ['Recommendation', 'Description', 'Priority'], data_recs)

doc.add_heading('15.4 Infrastructure Recommendations', level=2)
infra_recs = [
    ['Container Hardening', 'Implement Docker security best practices', 'HIGH'],
    ['Network Segmentation', 'Enhance network segmentation for internal services', 'HIGH'],
    ['Vulnerability Scanning', 'Implement automated dependency vulnerability scanning', 'HIGH'],
    ['Security Monitoring', 'Deploy application security monitoring and alerting', 'MEDIUM'],
]
add_styled_table(doc, ['Recommendation', 'Description', 'Priority'], infra_recs)

doc.add_heading('15.5 Process Recommendations', level=2)
process_recs = [
    ['Security Training', 'Conduct security awareness training for development team', 'HIGH'],
    ['Secure SDLC', 'Integrate security testing into CI/CD pipeline', 'HIGH'],
    ['Incident Response', 'Develop and test incident response procedures', 'HIGH'],
    ['Penetration Testing', 'Schedule annual penetration testing by external firm', 'MEDIUM'],
]
add_styled_table(doc, ['Recommendation', 'Description', 'Priority'], process_recs)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 16. APPENDIX: SECURITY SCANNING TOOLS
# ══════════════════════════════════════════════════════════════
doc.add_heading('16. Appendix: Security Scanning Tools', level=1)
doc.add_paragraph(
    'The following security scanning tools are recommended for comprehensive security testing of the '
    'PIXOUS HR Portal. Each tool serves a specific purpose in the security testing lifecycle.'
)

tools_items = [
    ['OWASP ZAP', 'Web Application Security Scanner', 'Dynamic application security testing (DAST)', 'Free'],
    ['Burp Suite', 'Web Vulnerability Scanner', 'Manual and automated web security testing', 'Commercial'],
    ['SonarQube', 'Code Quality & Security', 'Static application security testing (SAST)', 'Free/Commercial'],
    ['Snyk', 'Dependency Vulnerability Scanner', 'Open source vulnerability detection', 'Free/Commercial'],
    ['Trivy', 'Container Security Scanner', 'Docker image vulnerability scanning', 'Free'],
    ['SQLMap', 'SQL Injection Tool', 'Automated SQL injection testing', 'Free'],
    ['Nmap', 'Network Scanner', 'Port scanning and network discovery', 'Free'],
    ['Nikto', 'Web Server Scanner', 'Web server vulnerability scanning', 'Free'],
    ['Metasploit', 'Penetration Testing Framework', 'Exploitation and vulnerability validation', 'Free/Commercial'],
    ['Veracode', 'Application Security Platform', 'Comprehensive security testing', 'Commercial'],
]
add_styled_table(doc, ['Tool Name', 'Category', 'Purpose', 'License'], tools_items)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DOCUMENT REVISION HISTORY
# ══════════════════════════════════════════════════════════════
doc.add_heading('Document Revision History', level=1)
revision_items = [
    ['1.0', datetime.date.today().strftime('%B %d, %Y'), 'PIXOUS Technologies Security Team', 'Initial document creation'],
]
add_styled_table(doc, ['Version', 'Date', 'Author', 'Description'], revision_items)

doc.add_paragraph()
p = doc.add_paragraph('END OF DOCUMENT')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs:
    r.font.bold = True
    r.font.size = Pt(12)

# ── Save ──
doc.save('PIXOUS_HR_Portal_Security_Testing_Document.docx')
print("Document generated: PIXOUS_HR_Portal_Security_Testing_Document.docx")
